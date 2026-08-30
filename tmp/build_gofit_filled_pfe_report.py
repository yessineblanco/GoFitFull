import json
import os
import re
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from build_pfe_report_template import (
    setup_styles,
    add_heading,
    add_caption,
    add_basic_table,
    finalize,
)


OUT = "output/documents/PFE_GoFit_Polished_Academic_Draft.docx"
DIAGRAM_DIR = "output/documents/gofit_report_diagrams"
FEATURES_MD = "FEATURES.md"


def normalize_feature_text(value):
    text = str(value or "")
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = text.replace("**Description:**", "").replace("**Technical flow:**", "")
    text = text.replace("&amp;", "&")
    text = text.encode("ascii", "ignore").decode("ascii")
    text = re.sub(r"\s+", " ", text).strip(" -|")
    return text or "N/A"


def compact_feature_text(value, limit=170):
    text = normalize_feature_text(value)
    if len(text) <= limit:
        return text
    return text[: limit - 3].rstrip() + "..."


def combine_feature_fields(fields, names, limit=170):
    parts = []
    for name in names:
        value = normalize_feature_text(fields.get(name, ""))
        if value != "N/A" and value not in parts:
            parts.append(value)
    return compact_feature_text("; ".join(parts) if parts else "N/A", limit)


def parse_feature_inventory():
    if not os.path.exists(FEATURES_MD):
        return []
    with open(FEATURES_MD, "r", encoding="utf-8", errors="replace") as handle:
        lines = handle.read().splitlines()

    rows = []
    current_area = "Unclassified"
    for index, line in enumerate(lines):
        stripped = line.strip()
        if (stripped.startswith("## ") or stripped.startswith("### ")) and "Tech Stack per Feature Category" in normalize_feature_text(stripped):
            break
        if stripped.startswith("### "):
            current_area = normalize_feature_text(stripped[4:])
            continue
        if not stripped.startswith("#### "):
            continue

        title = normalize_feature_text(stripped[5:])
        table_index = index + 1
        while table_index < len(lines) and not lines[table_index].strip():
            table_index += 1
        if table_index >= len(lines) or lines[table_index].strip() != "| Field | Detail |":
            continue

        fields = {}
        row_index = table_index + 2
        while row_index < len(lines) and lines[row_index].strip().startswith("|"):
            parts = lines[row_index].strip().strip("|").split("|")
            if len(parts) >= 2:
                field = normalize_feature_text(parts[0])
                detail = normalize_feature_text("|".join(parts[1:]))
                fields[field] = detail
            row_index += 1

        while row_index < len(lines) and not lines[row_index].strip():
            row_index += 1

        description = "Implemented source feature documented in FEATURES.md."
        if row_index < len(lines) and lines[row_index].strip().startswith("**Description:**"):
            description = normalize_feature_text(lines[row_index].strip())

        rows.append(
            {
                "area": current_area,
                "title": title,
                "fields": fields,
                "description": description,
            }
        )
    return rows


def feature_group(area):
    lower = area.lower()
    if "mobile" in lower and "client" in lower:
        return "Mobile client"
    if "mobile" in lower and "coach" in lower:
        return "Mobile coach"
    if "admin" in lower:
        return "Admin panel"
    if "backend" in lower or "infrastructure" in lower:
        return "Backend and infrastructure"
    return "Other"


def feature_inventory_table_rows(features, group_name):
    selected = [row for row in features if feature_group(row["area"]) == group_name]
    table_rows = []
    for number, row in enumerate(selected, start=1):
        fields = row["fields"]
        evidence = combine_feature_fields(
            fields,
            ["Screen / Route", "Components", "Service", "Store"],
            165,
        )
        backend_and_libraries = combine_feature_fields(
            fields,
            ["Supabase Tables", "API Routes", "Libraries"],
            165,
        )
        table_rows.append(
            [
                str(number),
                compact_feature_text(row["title"], 95),
                evidence,
                backend_and_libraries,
                compact_feature_text(row["description"], 180),
            ]
        )
    return table_rows


def dependency_role(package_name):
    lower = package_name.lower()
    if lower.startswith("@radix-ui"):
        return "Accessible primitive component used in the admin UI."
    if lower.startswith("@react-navigation"):
        return "Mobile navigation stacks, tabs, and route transitions."
    if lower.startswith("@livekit"):
        return "Coach/client video-call and WebRTC integration."
    if lower.startswith("@supabase"):
        return "Supabase authentication, database, storage, realtime, SSR, or server access."
    if lower.startswith("@expo"):
        return "Expo-supported mobile UI, fonts, icons, or native capability."
    if lower.startswith("expo"):
        return "Expo native module or build/runtime capability used by the mobile application."
    if lower.startswith("react-native"):
        return "React Native ecosystem package for native UI, media, animation, or platform support."
    role_map = {
        "@aws-sdk/client-s3": "Optional S3-compatible object-storage integration.",
        "@hookform/resolvers": "Connects React Hook Form with schema validation.",
        "@react-native-async-storage/async-storage": "Local persisted preference and state storage.",
        "class-variance-authority": "Reusable component variant styling for the admin interface.",
        "clsx": "Conditional class-name composition.",
        "cmdk": "Command/search interface behavior.",
        "date-fns": "Date formatting, filtering, and scheduling logic.",
        "i18next": "Mobile localization engine.",
        "jpeg-js": "JPEG/image processing support for image-based flows.",
        "lucide-react": "Admin panel icon library.",
        "lucide-react-native": "Mobile icon library.",
        "next": "Next.js administration panel framework.",
        "next-themes": "Admin panel light/dark theme handling.",
        "react": "React component runtime.",
        "react-day-picker": "Admin date-picker/calendar interactions.",
        "react-dom": "React web rendering for the admin panel.",
        "react-hook-form": "Form state management and validation workflow.",
        "react-i18next": "React bindings for mobile localization.",
        "recharts": "Admin dashboard charts and analytics visualizations.",
        "sonner": "Toast notifications in the admin panel.",
        "tailwind-merge": "Tailwind class conflict resolution.",
        "tailwindcss-animate": "Animation utilities for Tailwind-based admin UI.",
        "webrtc-adapter": "WebRTC compatibility adapter for video-call behavior.",
        "zod": "Typed schema validation.",
        "zustand": "Mobile state management and persisted stores.",
    }
    return role_map.get(lower, "Runtime dependency used by the GoFit implementation.")


def dependency_rows(package_json_path):
    if not os.path.exists(package_json_path):
        return [["N/A", "N/A", f"Missing package file: {package_json_path}"]]
    with open(package_json_path, "r", encoding="utf-8") as handle:
        data = json.load(handle)
    rows = []
    for name, version in sorted(data.get("dependencies", {}).items()):
        rows.append([name, version, dependency_role(name)])
    return rows or [["N/A", "N/A", "No runtime dependencies listed."]]


def p(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.12
    para.add_run(text)
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(3)
    para.add_run(text)
    return para


def numbered(doc, text):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.space_after = Pt(3)
    para.add_run(text)
    return para


def _font(size=26, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for path in candidates:
        if os.path.exists(path):
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def make_diagram(filename, title, rows):
    os.makedirs(DIAGRAM_DIR, exist_ok=True)
    path = os.path.join(DIAGRAM_DIR, filename)
    width = 1450
    margin = 60
    title_h = 90
    row_h = 165
    height = title_h + row_h * len(rows) + 70
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = _font(32, True)
    box_font = _font(23, False)
    draw.text((margin, 28), title, fill=(31, 77, 120), font=title_font)

    centers_by_row = []
    for r_idx, row in enumerate(rows):
        y = title_h + r_idx * row_h + 25
        gap = 35
        box_w = int((width - 2 * margin - gap * (len(row) - 1)) / len(row))
        box_h = 95
        centers = []
        for c_idx, label in enumerate(row):
            x = margin + c_idx * (box_w + gap)
            draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=18, fill=(247, 249, 252), outline=(88, 124, 164), width=3)
            wrapped = textwrap.wrap(label, width=max(12, int(box_w / 18)))
            text_h = len(wrapped) * 26
            for i, line in enumerate(wrapped[:3]):
                bbox = draw.textbbox((0, 0), line, font=box_font)
                tx = x + (box_w - (bbox[2] - bbox[0])) / 2
                ty = y + (box_h - text_h) / 2 + i * 26
                draw.text((tx, ty), line, fill=(25, 42, 61), font=box_font)
            centers.append((x + box_w / 2, y + box_h / 2, x, y, box_w, box_h))
        for c_idx in range(len(centers) - 1):
            x1 = centers[c_idx][2] + centers[c_idx][4]
            y1 = centers[c_idx][1]
            x2 = centers[c_idx + 1][2]
            draw.line((x1 + 8, y1, x2 - 8, y1), fill=(46, 116, 181), width=4)
            draw.polygon([(x2 - 8, y1), (x2 - 24, y1 - 9), (x2 - 24, y1 + 9)], fill=(46, 116, 181))
        centers_by_row.append(centers)

    for r_idx in range(len(centers_by_row) - 1):
        upper = centers_by_row[r_idx][len(centers_by_row[r_idx]) // 2]
        lower = centers_by_row[r_idx + 1][len(centers_by_row[r_idx + 1]) // 2]
        x1, y1 = upper[0], upper[3] + upper[5]
        x2, y2 = lower[0], lower[3]
        draw.line((x1, y1 + 8, x2, y2 - 8), fill=(46, 116, 181), width=4)
        draw.polygon([(x2, y2 - 8), (x2 - 9, y2 - 24), (x2 + 9, y2 - 24)], fill=(46, 116, 181))
    image.save(path)
    return path


def add_diagram(doc, filename, caption, title, rows):
    path = make_diagram(filename, title, rows)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(path, width=Inches(6.35))
    add_caption(doc, caption)


def add_report_note(doc, title, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.line_spacing = 1.08
    r = para.add_run(f"{title}: ")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    para.add_run(text)


def add_code_block(doc, title, code):
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(4)
    label.paragraph_format.space_after = Pt(2)
    r = label.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.0
    run = para.add_run(code.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    return para


def add_todo(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(f"[TODO: {text}]")
    run.italic = True
    run.font.color.rgb = RGBColor(92, 92, 92)
    return para


def mermaid_code_for(kind, sprint=None):
    if kind == "global_use_case":
        return """
flowchart LR
  U[Mobile user] --> A((Manage account and profile))
  U --> W((Plan and execute workouts))
  U --> P((Track progress and measurements))
  U --> M((Browse coaches and book sessions))
  C[Coach] --> CP((Manage coach profile))
  C --> CL((Follow clients and programs))
  C --> B((Manage bookings and session packs))
  AD[Administrator] --> ADM((Manage users, coaches, content, analytics))
  SB[Supabase backend] --> AUTH[(Auth, database, storage, realtime, Edge Functions)]
  A --> AUTH
  W --> AUTH
  P --> AUTH
  M --> AUTH
  CP --> AUTH
  CL --> AUTH
  B --> AUTH
  ADM --> AUTH
"""
    if kind == "architecture":
        return """
flowchart TB
  Mobile[GoFit mobile application] --> Supabase[Supabase backend]
  Admin[GoFit administration panel] --> Supabase
  Supabase --> DB[(PostgreSQL + RLS)]
  Supabase --> Storage[(Storage buckets)]
  Supabase --> Realtime[Realtime channels]
  Supabase --> Edge[Edge Functions]
  Edge --> Push[Expo Push]
  Edge --> Video[LiveKit token service]
  Edge --> Groq[Groq LLM API]
  Mobile --> CV[On-device CV: MediaPipe / MoveNet]
  Automation[n8n automations] --> Supabase
  Automation --> Groq
"""
    if kind == "deployment":
        return """
flowchart LR
  Repo[Source repository] --> EAS[EAS mobile build]
  Repo --> Web[Next.js hosting]
  Repo --> Supa[Supabase project]
  EAS --> Android[Android build]
  EAS --> IOS[iOS build]
  Web --> AdminPanel[Administration panel]
  Supa --> Auth[Auth]
  Supa --> DB[(Database migrations)]
  Supa --> Edge[Edge Functions]
  Supa --> Storage[Storage]
  Env[Environment variables] --> EAS
  Env --> Web
  Env --> Supa
"""
    if sprint == 3 and kind == "use_case":
        return """
flowchart LR
  U[Mobile user] --> Login((Sign up / sign in))
  U --> Onboard((Complete onboarding))
  U --> Profile((Manage profile))
  U --> Build((Create workout))
  U --> Session((Execute workout session))
  U --> Timer((Use rest timer))
  Session --> Save[(workout_sessions)]
"""
    if sprint == 4 and kind == "use_case":
        return """
flowchart LR
  U[Mobile user] --> Library((Search exercise library))
  U --> Progress((View progress dashboard))
  U --> Measure((Log body measurements))
  U --> AI((Generate AI measurement draft))
  U --> Nutrition((Manage nutrition / saved meals))
  AI --> Review((Review and correct values))
  Review --> Save[(body_measurements)]
"""
    if sprint == 5 and kind == "use_case":
        return """
flowchart LR
  Client[Mobile user] --> Browse((Browse coaches))
  Client --> Book((Book coaching session))
  Client --> Chat((Chat with coach))
  Coach[Coach] --> Onboard((Complete coach onboarding))
  Coach --> Programs((Manage client programs))
  Admin[Administrator] --> Validate((Validate coach))
  Admin --> Manage((Manage platform content))
"""
    if sprint == 3 and kind == "class":
        return """
classDiagram
  class UserProfile { id; goal; preferences; onboarding_status }
  class Workout { id; name; difficulty; owner_id }
  class Exercise { id; name; muscle_group; equipment }
  class WorkoutExercise { workout_id; exercise_id; order; sets; reps }
  class WorkoutSession { id; workout_id; started_at; completed_at; summary }
  class TimerPreferences { rest_seconds; warnings_enabled; auto_advance }
  UserProfile "1" --> "*" Workout
  Workout "1" --> "*" WorkoutExercise
  Exercise "1" --> "*" WorkoutExercise
  Workout "1" --> "*" WorkoutSession
  UserProfile "1" --> "1" TimerPreferences
"""
    if sprint == 4 and kind == "class":
        return """
classDiagram
  class UserProfile { id; height_cm; goal; gender }
  class WorkoutSession { id; user_id; completed_at; duration }
  class BodyMeasurement { id; chest_cm; waist_cm; hip_cm; shoulder_cm; source }
  class ProgressPhoto { id; user_id; uri; capture_type }
  class NutritionEntry { id; user_id; meal_type; calories }
  class SavedMeal { id; user_id; name }
  class MeasurementFeatureVector { pose_score; scale_cm_per_px; segmentation_quality }
  UserProfile "1" --> "*" WorkoutSession
  UserProfile "1" --> "*" BodyMeasurement
  BodyMeasurement "1" --> "0..1" MeasurementFeatureVector
  UserProfile "1" --> "*" NutritionEntry
  UserProfile "1" --> "*" SavedMeal
"""
    if sprint == 5 and kind == "class":
        return """
classDiagram
  class CoachProfile { id; user_id; status; specialties }
  class Booking { id; coach_id; client_id; scheduled_at; status }
  class SessionPack { id; coach_id; title; sessions_count }
  class PurchasedPack { id; client_id; pack_id; remaining_sessions }
  class Conversation { id; coach_id; client_id }
  class Message { id; conversation_id; sender_id; body }
  class Program { id; coach_id; client_id; title }
  CoachProfile "1" --> "*" Booking
  SessionPack "1" --> "*" PurchasedPack
  Conversation "1" --> "*" Message
  CoachProfile "1" --> "*" Program
"""
    if sprint == 3 and kind == "sequence":
        return """
sequenceDiagram
  actor User
  participant Screen as Workout Session Screen
  participant Store as Workout Store
  participant Service as Workout Service
  participant DB as Supabase backend
  User->>Screen: Start workout
  Screen->>Store: Initialize active session
  User->>Screen: Complete set and rest
  Screen->>Store: Update session state
  User->>Screen: Finish workout
  Store->>Service: Persist summary
  Service->>DB: Insert workout_sessions row
  DB-->>Screen: Return saved session
"""
    if sprint == 4 and kind == "sequence":
        return """
sequenceDiagram
  actor User
  participant Camera as Body Measurement Screen
  participant Pose as MediaPipe / MoveNet
  participant Segmenter as Segmentation Service
  participant Estimator as Measurement Estimator
  participant DB as Supabase backend
  User->>Camera: Capture front and side photos
  Camera->>Pose: Detect body landmarks
  Camera->>Segmenter: Extract person mask
  Segmenter-->>Estimator: Widths, depths, quality values
  Estimator-->>Camera: Draft measurements and warnings
  User->>Camera: Review or correct values
  Camera->>DB: Save final measurement
"""
    if sprint == 5 and kind == "sequence":
        return """
sequenceDiagram
  actor Client
  participant Market as Marketplace Screen
  participant Service as Booking Service
  participant DB as Supabase backend
  participant Coach as Coach Dashboard
  Client->>Market: Select coach and session
  Market->>Service: Request booking
  Service->>DB: Insert booking and related notification
  DB-->>Market: Confirmation result
  DB-->>Coach: Realtime/dashboard update
  Coach-->>Client: Follow-up through chat or program
"""
    return ""


def detailed_use_cases_for_sprint(n):
    cases = {
        3: [
            {
                "caption": "Table 3.2: Detailed use case - Authenticate and complete onboarding",
                "rows": [
                    ["Use case", "Authenticate and complete onboarding"],
                    ["Main actor", "Mobile user"],
                    ["Preconditions", "The user has installed the GoFit mobile application and has access to an email/password or an existing account."],
                    ["Postconditions", "An authenticated session is stored, the user profile is created or updated, and the user is routed to the appropriate application area."],
                    ["Nominal scenario", "The user opens the application, signs up or signs in, completes required onboarding fields, saves preferences, and reaches the protected home screen."],
                    ["Alternative scenario", "If the user already completed onboarding, the application bypasses onboarding and opens the protected navigator."],
                    ["Exception scenario", "If credentials are invalid or the network request fails, the application keeps the user on the authentication screen and displays a recoverable error."],
                ],
            },
            {
                "caption": "Table 3.3: Detailed use case - Execute workout session",
                "rows": [
                    ["Use case", "Execute and save a workout session"],
                    ["Main actor", "Mobile user"],
                    ["Preconditions", "The user is authenticated and a workout template or custom workout exists."],
                    ["Postconditions", "The workout session summary is saved with duration, performed exercises, set data, and notes."],
                    ["Nominal scenario", "The user opens a workout, starts the session, completes sets, uses the rest timer between sets, finishes the workout, reviews the summary, and saves it."],
                    ["Alternative scenario", "The user edits set values, skips an exercise, pauses the timer, or exits before saving the final summary."],
                    ["Exception scenario", "If persistence fails, the application must keep the local session state and show an error so the user does not lose workout data."],
                ],
            },
        ],
        4: [
            {
                "caption": "Table 4.2: Detailed use case - View progress dashboard",
                "rows": [
                    ["Use case", "View progress dashboard"],
                    ["Main actor", "Mobile user"],
                    ["Preconditions", "The user is authenticated and may have workout sessions, measurements, or nutrition entries."],
                    ["Postconditions", "The user obtains a readable view of training evolution, measurement history, or empty-state guidance."],
                    ["Nominal scenario", "The user opens the progress screen, selects a metric or period, and the application loads chart data from workout and measurement services."],
                    ["Alternative scenario", "If no historical data exists, the application displays an empty state that explains how to create the first entry."],
                    ["Exception scenario", "If the Supabase query fails, the application must display an error state and preserve navigation stability."],
                ],
            },
            {
                "caption": "Table 4.3: Detailed use case - Generate AI-assisted body measurement draft",
                "rows": [
                    ["Use case", "Generate AI-assisted body measurement draft"],
                    ["Main actor", "Mobile user"],
                    ["Preconditions", "The user is authenticated, has provided height information, and captures front and side photos under acceptable conditions."],
                    ["Postconditions", "Draft measurements, confidence values, quality warnings, and editable fields are presented before saving."],
                    ["Nominal scenario", "The app analyzes the photos with pose detection and segmentation, builds a feature vector, estimates draft values, and asks the user to review them."],
                    ["Alternative scenario", "If segmentation quality is weak, the estimator uses safer statistical fallback values and displays quality warnings."],
                    ["Exception scenario", "If the native AI module is unavailable or the capture is unusable, the app blocks automatic estimation and invites manual entry or retake."],
                ],
            },
        ],
        5: [
            {
                "caption": "Table 5.2: Detailed use case - Book coaching session",
                "rows": [
                    ["Use case", "Book a coaching session"],
                    ["Main actor", "Mobile user"],
                    ["Preconditions", "The user is authenticated, at least one coach profile is available, and a valid slot or session pack exists."],
                    ["Postconditions", "A booking record is created and becomes visible to the client and coach according to access rules."],
                    ["Nominal scenario", "The user browses coaches, opens a profile, chooses a service or slot, confirms the booking, and receives confirmation."],
                    ["Alternative scenario", "The user uses an existing purchased pack instead of creating a standalone booking."],
                    ["Exception scenario", "If the slot is unavailable, payment/session pack is missing, or the request violates RLS policies, the booking is rejected with a clear message."],
                ],
            },
            {
                "caption": "Table 5.3: Detailed use case - Validate and manage coach from administration panel",
                "rows": [
                    ["Use case", "Validate and manage coach account"],
                    ["Main actor", "Administrator"],
                    ["Preconditions", "The administrator is authenticated with admin privileges and a coach profile exists in pending or active status."],
                    ["Postconditions", "The coach status, visibility, or profile data is updated and recorded in the administration interface."],
                    ["Nominal scenario", "The administrator opens the coach management page, reviews profile information, validates or updates the status, and saves the decision."],
                    ["Alternative scenario", "The administrator requests missing information or leaves the coach in pending status."],
                    ["Exception scenario", "If the connected user is not an administrator, middleware and server-side checks must block access to the page or action."],
                ],
            },
        ],
    }
    return cases[n]


def add_detailed_use_cases(doc, n):
    for case in detailed_use_cases_for_sprint(n):
        add_basic_table(doc, case["caption"], ["Field", "Description"], case["rows"], [1.45, 5.05])


def professional_test_rows(n):
    evidence = "TODO: attach screenshot or terminal output"
    rows = {
        3: [
            ["S1-T01", "Authentication", "Valid user signs in", "Test account exists", "Open app; enter credentials; submit", "Protected navigator opens", "No attached execution evidence yet", "To be validated", evidence],
            ["S1-T02", "Onboarding", "New user saves profile", "Authenticated user without onboarding", "Fill goals, units, preferences; save", "Profile is stored and onboarding flag updates", "No attached execution evidence yet", "To be validated", evidence],
            ["S1-T03", "Workout session", "User completes workout", "Workout template exists", "Start workout; complete sets; finish summary", "Session is saved in workout history", "No attached execution evidence yet", "To be validated", evidence],
            ["S1-T04", "Rest timer", "Timer controls behave correctly", "Active workout session", "Start, pause, resume, warning, complete", "Timer state and feedback remain consistent", "No attached execution evidence yet", "Manual validation required", evidence],
        ],
        4: [
            ["S2-T01", "Exercise library", "Search and open exercise", "Exercise data exists", "Search; filter; open detail", "Correct exercise detail is displayed", "No attached execution evidence yet", "To be validated", evidence],
            ["S2-T02", "Progress dashboard", "Display charts and empty state", "User with/without history", "Open progress; switch metric and period", "Charts or empty state render correctly", "No attached execution evidence yet", "To be validated", evidence],
            ["S2-T03", "AI measurement", "Generate draft measurement", "Front/side photos and height available", "Capture photos; run analysis; review values", "Draft values and warnings are shown before save", "Android debug observations exist, final evidence missing", "Partially validated", evidence],
            ["S2-T04", "Notifications", "Configure reminder permission", "Device/emulator supports notifications", "Request permission; save reminder preference", "Preference persists and no crash occurs", "No attached execution evidence yet", "Manual validation required", evidence],
            ["S2-T05", "Nutrition/saved meals", "Create or reuse saved meal", "Authenticated user", "Create saved meal; edit item; reopen list", "Meal entry persists in history", "No attached execution evidence yet", "To be validated", evidence],
        ],
        5: [
            ["S3-T01", "Coach onboarding", "Coach submits profile", "Coach account exists", "Fill profile; upload/support documents; submit", "Coach appears in pending/review state", "No attached execution evidence yet", "To be validated", evidence],
            ["S3-T02", "Marketplace booking", "Client books coach", "Coach and slot/pack exist", "Open marketplace; select coach; confirm booking", "Booking is visible to client and coach", "No attached execution evidence yet", "To be validated", evidence],
            ["S3-T03", "Chat", "Client and coach exchange messages", "Conversation exists", "Send message from one side; open other side", "Message appears in conversation", "No attached execution evidence yet", "Manual validation required", evidence],
            ["S3-T04", "Admin access", "Non-admin blocked from admin panel", "Non-admin account and admin route exist", "Attempt to open admin route", "Access is denied or redirected", "No attached execution evidence yet", "To be validated", evidence],
            ["S3-T05", "AI briefing automation", "Coach briefing generated for booking", "n8n env and booking context exist", "Run workflow; inspect ai_session_notes and notification", "One briefing note and notification are created", "Workflow logic documented, final run evidence missing", "Partially validated", evidence],
        ],
    }
    return rows[n]


def add_professional_tests(doc, n, table_number):
    add_basic_table(
        doc,
        f"Table {table_number}: Sprint validation matrix",
        ["Test ID", "Feature", "Scenario", "Preconditions", "Steps", "Expected", "Actual", "Status", "Evidence"],
        professional_test_rows(n),
        [0.48, 0.72, 0.88, 0.85, 1.05, 0.95, 0.9, 0.76, 0.91],
    )


def sprint_conclusion(n):
    conclusions = {
        3: "Sprint 1 established the operational base of the mobile application. Authentication, onboarding, protected navigation, workout templates, active workout execution, persistence, and rest timer behavior became the foundation for later progress tracking. The next sprint builds on this foundation by transforming workout data into progress, measurement, notification, and nutrition-oriented features.",
        4: "Sprint 2 transformed the mobile foundation into a progress-oriented experience. It connected exercise browsing, history, progress dashboards, body measurements, nutrition foundations, reminders, and AI-assisted measurement research. The sprint also clarified the limits of computer-vision estimation and the need for manual review. The next sprint extends GoFit from individual tracking to coach marketplace and administration workflows.",
        5: "Sprint 3 completed the platform dimension of GoFit by adding coach marketplace, coach-client follow-up, administration, and LLM-assisted preparation workflows. This sprint connects users, coaches, administrators, backend services, and automations into a broader operational system. The following chapter therefore focuses on deployment, architecture, security, limitations, and closing considerations.",
    }
    return conclusions[n]


def cover(doc):
    top = doc.add_paragraph()
    top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = top.add_run("University / Institute Name\nDepartment / Program")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(31, 77, 120)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(90)
    r = title.add_run("Graduation Internship Report")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(46, 116, 181)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("GoFit - Complete Fitness Mobile Application, Coach Marketplace, and Administration Platform")
    r.bold = True
    r.font.size = Pt(17)

    add_diagram(
        doc,
        "figure_0_1_cover_identity.png",
        "Figure 0.1: University and GoFit visual identity area",
        "Report Identity",
        [["University Logo", "GoFit", "Host Organization Logo"]],
    )
    add_basic_table(
        doc,
        "Table 0.1: Report identification",
        ["Field", "Value"],
        [
            ["Prepared by", "[TODO: insert official student full name]"],
            ["Academic supervisor", "[TODO: insert official academic supervisor name]"],
            ["Professional supervisor", "[TODO: insert official professional supervisor name]"],
            ["Host organization", "[TODO: insert official host organization details]"],
            ["Academic year", "2025 - 2026"],
        ],
        [2.1, 4.2],
    )
    doc.add_page_break()


def front_matter(doc):
    add_heading(doc, "Validation / Signature Page", 1)
    p(doc, "This page is reserved for the official validation of the graduation internship report. It should include the signatures of the academic supervisor, professional supervisor, jury members, and any required institutional stamp.")
    add_basic_table(
        doc,
        "Table 0.2: Validation signatures",
        ["Role", "Name", "Signature", "Date"],
        [
            ["Academic supervisor", "[TODO: official name]", "[TODO: signature]", "[TODO: date]"],
            ["Professional supervisor", "[TODO: official name]", "[TODO: signature]", "[TODO: date]"],
            ["Jury president", "[TODO: official name]", "[TODO: signature]", "[TODO: date]"],
            ["Reviewer / examiner", "[TODO: official name]", "[TODO: signature]", "[TODO: date]"],
        ],
        [1.8, 1.8, 1.5, 1.2],
    )
    doc.add_page_break()

    add_heading(doc, "Acknowledgements", 1)
    p(doc, "I would like to express my sincere gratitude to my academic and professional supervisors for their guidance, feedback, and availability throughout this graduation internship. Their support helped me transform the GoFit project from a broad fitness platform idea into a structured mobile, web, and backend solution.")
    p(doc, "I also thank the teaching staff and all people who contributed directly or indirectly to this work. Their advice, technical observations, and encouragement helped me improve the quality of the analysis, implementation, and documentation.")
    doc.add_page_break()

    add_heading(doc, "Dedication", 1)
    p(doc, "I dedicate this work to my family, friends, mentors, and everyone who supported me during my academic journey and internship period. Their confidence and encouragement were essential during the demanding phases of design, development, testing, and report writing.")
    doc.add_page_break()

    add_heading(doc, "Table of Contents", 1)
    p(doc, "[TODO: generate the final Word table of contents after page numbering and final layout are validated.]")
    doc.add_page_break()
    add_heading(doc, "List of Figures", 1)
    p(doc, "[TODO: generate the final Word list of figures after official diagrams and real screenshots are inserted.]")
    doc.add_page_break()
    add_heading(doc, "List of Tables", 1)
    p(doc, "[TODO: generate the final Word list of tables after the final formatting pass.]")
    doc.add_page_break()


def introduction(doc):
    add_heading(doc, "General Introduction", 1)
    p(doc, "Digital fitness has become an important part of modern health and training habits. Users increasingly expect mobile applications to help them plan workouts, follow progress, manage nutrition, receive reminders, and access professional coaching from a single experience. At the same time, coaches need tools to organize clients, programs, bookings, communication, and session follow-up without relying on several disconnected platforms.")
    p(doc, "The GoFit project responds to this context by proposing a unified fitness platform. It combines an Expo/React Native mobile application, a Supabase backend, and a Next.js administration panel. The mobile application supports authentication, onboarding, workout planning, exercise library browsing, active workout sessions, rest timer management, progress monitoring, nutrition-related features, coach marketplace flows, bookings, chat, programs, and user preferences. The administration panel provides operational control over users, coaches, exercises, workouts, analytics, settings, and activity logs.")
    p(doc, "The project also includes an AI dimension. GoFit is not only a CRUD application; it integrates AI-assisted body measurement, adaptive workout recommendations, AI coach briefing preparation, and automation workflows. The body-measurement feature combines on-device pose estimation, segmentation, feature extraction, statistical estimation, and manual validation. The recommendation and briefing features use a language-model workflow through Groq and controlled backend contexts. In addition, AI coding tools, including Codex, were used during project construction to accelerate analysis, documentation, code review, debugging support, report generation, and architecture reasoning. MCP-style connected tools and plugins helped keep the work grounded in current project files, official documentation, local command outputs, document tooling, and backend/automation context. These tools did not replace engineering validation; they were used as supervised assistants within a human-controlled development workflow.")
    p(doc, "The main objective of the internship was to design, implement, and document a substantial software foundation for GoFit. The work included requirements analysis, Agile planning, database modeling, mobile and web development, backend integration, security considerations, testing, and deployment preparation.")
    p(doc, "The scientific and technical interest of the project lies in the combination of several software-engineering concerns: a cross-platform mobile interface, a web administration panel, a secure backend-as-a-service architecture, realtime communication, media and notification integrations, AI-assisted user features, and a documented Agile delivery process. The final system is therefore studied as a complete digital product rather than as an isolated mobile application.")
    p(doc, "This report is organized into six chapters. Chapter 1 presents the project context, problem statement, existing limitations, proposed solution, methodology, and modeling language. Chapter 2 describes Sprint 0, including actors, requirements, product backlog, Scrum organization, and global use case modeling. Chapters 3, 4, and 5 present the main development sprints: mobile foundation and workout core, progress and nutrition features, then coach marketplace and administration. Chapter 6 presents the working environment, architecture, technologies, deployment approach, and final validation. The report ends with a general conclusion, bibliography, and appendices.")
    doc.add_page_break()


def chapter1(doc):
    add_heading(doc, "1 Project Study", 1)
    add_heading(doc, "1.1 Project Context", 2)
    p(doc, "GoFit is positioned in the digital fitness and coaching domain. The project targets users who want to organize training, monitor progress, and access fitness support through a mobile application. It also targets coaches who need a more structured way to present their services, follow clients, schedule sessions, and communicate. Finally, it includes an administration layer for managing content and platform operations.")
    p(doc, "The technical context is a multi-part product: an Expo/React Native mobile application for iOS and Android, a Next.js web administration panel, and a Supabase backend providing PostgreSQL, authentication, storage, and security policies. This architecture allows the project to separate mobile user experience, administrative workflows, and backend persistence.")

    add_heading(doc, "1.2 Host Organization Presentation", 2)
    p(doc, "The host organization or project environment supported the development of GoFit as a graduation internship project. The work was conducted with regular supervision, progressive validation, and a Scrum-inspired organization. The project owner and supervisors helped prioritize the scope around workout planning, progress tracking, coach marketplace flows, administration, and deployment preparation.")
    add_heading(doc, "1.2.1 Presentation of the Company", 3)
    p(doc, "[TODO: complete official host organization details: legal name, business area, services, target clients, structure, and supervising department.] In the current project material, GoFit is presented as the software product developed during the internship: a fitness platform dedicated to workout tracking, progress monitoring, coach marketplace workflows, AI-assisted support, and operational administration.")
    add_heading(doc, "1.2.2 Company Logo", 3)
    add_diagram(
        doc,
        "figure_1_1_gofit_identity.png",
        "Figure 1.1: GoFit and host organization identity",
        "GoFit Identity",
        [["GoFit Fitness Platform", "Mobile App", "Coach Marketplace", "Admin Panel"]],
    )

    add_heading(doc, "1.3 Problem Statement", 2)
    p(doc, "Many fitness users rely on several disconnected tools: one application for workout planning, another for nutrition, a spreadsheet for body measurements, messaging applications for coach communication, and separate calendars or payment tools for coaching sessions. This fragmentation reduces engagement, creates duplicated data entry, and makes long-term progress harder to understand.")
    p(doc, "For coaches, the issue is also operational. Managing client profiles, programs, bookings, messages, session packs, and progress follow-up across separate services increases administrative workload and reduces the quality of personalized support. Administrators also need a centralized interface to manage users, coaches, exercises, workouts, analytics, and platform content.")
    p(doc, "A second difficulty concerns personalization. Fitness applications often ask users to follow predefined programs without considering recent activity, readiness, goals, progress, or coach context. Coaches may personalize plans manually, but this becomes time-consuming as the number of clients grows. GoFit therefore investigates AI-assisted recommendations and AI-generated coach briefings to support decision-making while keeping the coach and user in control.")
    p(doc, "A third difficulty concerns body-progress tracking. Users often want visual and numerical feedback about body evolution, but manual tape measurements are irregular and progress photos are hard to interpret objectively. Computer vision can help extract pose and silhouette indicators, but it also introduces technical constraints: photo quality, camera angle, clothing, segmentation quality, privacy, model performance, and the fact that consumer photos cannot produce perfect circumference measurements without validation.")

    add_heading(doc, "1.4 Study of Existing Conditions", 2)
    p(doc, "Existing fitness solutions often focus on only one part of the user journey. Workout applications generally provide templates and logging, nutrition applications focus on meals and calories, and coaching platforms often specialize in booking or messaging. Although these tools may be strong individually, they do not always provide a coherent end-to-end experience for users, coaches, and administrators.")
    p(doc, "Manual processes also remain common. Coaches may use spreadsheets, messaging apps, calendar tools, and external payment links. Users may manually enter measurements and training notes without automated reminders, dashboards, or coach visibility. These conditions motivated the design of a more integrated GoFit platform.")
    add_basic_table(
        doc,
        "Table 1.3: Existing-condition analysis",
        ["Area", "Common existing approach", "Observed limitation", "GoFit response"],
        [
            ["Workout tracking", "Standalone apps with fixed programs and session logs.", "Limited connection with coach follow-up, bookings, and personalized marketplace services.", "Integrated workout library, custom workouts, active sessions, rest timer, history, and coach program context."],
            ["Coach management", "Spreadsheets, messaging apps, calendars, and manual notes.", "Client information is scattered and difficult to operationalize.", "Coach dashboard, clients, notes, programs, bookings, packs, chat, and AI session briefing."],
            ["Progress tracking", "Manual body measurements, isolated progress photos, and simple charts.", "Users may forget to measure, enter inconsistent values, or fail to interpret visual changes.", "Progress screens, body-measurement workflow, quality checks, draft AI values, and manual correction."],
            ["Administration", "Direct database access or disconnected back-office tools.", "Risky operations, weak visibility, and no clear operational dashboard.", "Next.js admin panel with protected routes, user/coach/content management, settings, and analytics."],
            ["AI assistance", "Generic prompts or no AI support.", "Recommendations may be disconnected from real user data and unsafe if not scoped.", "Backend-controlled prompts using profile, workout, coach, booking, and note context with human review."],
        ],
        [1.2, 1.8, 1.8, 1.9],
    )

    add_heading(doc, "1.5 Limits and Inconveniences of Existing Solution", 2)
    bullet(doc, "Fragmented experience across workout, nutrition, progress, booking, payment, and messaging tools.")
    bullet(doc, "Limited coach-client connectivity and weak operational visibility for coaches.")
    bullet(doc, "Manual progress tracking and reduced accuracy for body measurements.")
    bullet(doc, "Lack of centralized administration for users, coaches, exercises, workouts, analytics, and settings.")
    bullet(doc, "Difficulty maintaining consistent data security and access control when several tools are used.")

    add_heading(doc, "1.6 Proposed Solution", 2)
    p(doc, "The proposed solution is GoFit, a complete fitness platform composed of three main parts. The first part is a cross-platform mobile application built with Expo and React Native. It supports authentication, onboarding, user profile management, workout planning, exercise browsing, active workout sessions, rest timers, progress monitoring, nutrition-related features, coach marketplace flows, bookings, chat, and personalization settings.")
    p(doc, "The second part is a Next.js administration panel. It allows administrators to manage users, coaches, exercises, workouts, transactions, settings, notifications, and activity logs. The third part is a Supabase backend that provides authentication, PostgreSQL database, storage, realtime capabilities, and row-level security.")
    p(doc, "The solution also includes AI-assisted services. On the client side, body measurement uses pose detection, segmentation, feature extraction, and statistical estimation to generate editable draft measurements. On the recommendation side, GoFit uses controlled language-model prompts to propose workout recommendations and prepare session briefings for coaches. On the operational side, n8n workflows prepare automated coach notes and administrative digests by combining Supabase data with Groq-generated summaries.")
    add_basic_table(
        doc,
        "Table 1.4: Proposed-solution modules",
        ["Layer", "Main responsibility", "Representative components"],
        [
            ["Mobile client layer", "Daily user and coach experience.", "Authentication, onboarding, workouts, sessions, rest timer, progress, nutrition, coach marketplace, bookings, chat, programs, video screen."],
            ["AI and automation layer", "Assist measurement, recommendation, briefing, and operational follow-up.", "MediaPipe, MoveNet, segmentation service, Groq LLM prompts, n8n AI session prep, admin operations digest."],
            ["Administration layer", "Platform monitoring and data management.", "Dashboard, users, coaches, exercises, workouts, transactions, notifications, settings, analytics, activity logs."],
            ["Backend layer", "Persistent data, access control, realtime communication, and privileged operations.", "Supabase Auth, PostgreSQL, RLS, Storage, Realtime, Edge Functions, migrations, service-role server operations."],
            ["External integration layer", "Specialized external services that should remain isolated behind controlled interfaces.", "Expo Push, LiveKit video, optional media/storage providers, Groq API, hosting/build platforms."],
        ],
        [1.35, 2.3, 2.85],
    )

    add_heading(doc, "1.7 Adopted Development Methodology and Modeling Language", 2)
    p(doc, "The project was organized using an Agile approach inspired by Scrum. This choice was appropriate because the product scope includes several connected modules and required regular adaptation after technical discoveries, UI/UX decisions, and backend constraints.")
    add_heading(doc, "1.7.1 Comparison Between Traditional and Agile Methodologies", 3)
    add_basic_table(
        doc,
        "Table 1.1: Comparison between traditional and agile methodologies",
        ["Criterion", "Traditional methodology", "Agile methodology"],
        [
            ["Planning", "Most planning is completed before implementation starts.", "Planning is refined continuously sprint by sprint."],
            ["Change management", "Changes are costly and often delayed.", "Changes can be integrated into the backlog."],
            ["Delivery", "The product is delivered near the end of the project.", "Working increments are delivered regularly."],
            ["Fit for GoFit", "Less suitable because mobile, backend, and admin needs evolved during implementation.", "Suitable because GoFit modules can be implemented and validated incrementally."],
        ],
        [1.4, 2.45, 2.45],
    )
    add_heading(doc, "1.7.2 Comparison Between Scrum and Kanban", 3)
    add_basic_table(
        doc,
        "Table 1.2: Comparison between Scrum and Kanban",
        ["Criterion", "Scrum", "Kanban"],
        [
            ["Cadence", "Time-boxed sprints with planning and review.", "Continuous flow without fixed sprint boundaries."],
            ["Roles", "Product Owner, Scrum Master, and Development Team.", "No mandatory roles."],
            ["Artifacts", "Product backlog, sprint backlog, increment.", "Visual board and work-in-progress limits."],
            ["Choice for GoFit", "Selected because the internship had planned increments and deadlines.", "Useful for task visualization but less structured for academic reporting."],
        ],
        [1.4, 2.45, 2.45],
    )
    add_heading(doc, "1.7.3 Scrum", 3)
    p(doc, "Scrum was used to structure the work into fixed iterations. The product backlog listed major GoFit features, each sprint selected a coherent subset, and each increment targeted a working part of the product. Sprint reviews and retrospectives supported feedback and improvement.")
    add_diagram(
        doc,
        "figure_1_2_scrum_flow.png",
        "Figure 1.2: Scrum flow used for GoFit",
        "Scrum Flow Used for GoFit",
        [
            ["Product Backlog", "Sprint Planning", "Sprint Backlog"],
            ["Development Sprint", "Daily Follow-up", "Increment"],
            ["Sprint Review", "Retrospective", "Backlog Refinement"],
        ],
    )
    add_heading(doc, "1.7.4 Modeling Language / UML", 3)
    p(doc, "UML was selected to model requirements and system design. Use case diagrams describe actors and system functions, class diagrams describe structural relationships, sequence diagrams describe interactions, and deployment diagrams describe the runtime environment.")
    add_heading(doc, "1.8 Conclusion", 2)
    p(doc, "This chapter presented the GoFit project, its context, problem statement, existing limitations, proposed solution, methodology, and modeling language. The next chapter presents Sprint 0, where requirements and project organization are formalized.")
    doc.add_page_break()


def chapter2(doc):
    add_heading(doc, "2 Sprint 0", 1)
    p(doc, "Sprint 0 prepared the project before feature implementation. It clarified actors, requirements, backlog organization, sprint planning, and the global functional perimeter of GoFit.")
    add_heading(doc, "2.1 Needs Assessment", 2)
    add_heading(doc, "2.1.1 Identification of Actors", 3)
    add_basic_table(
        doc,
        "Table 2.1: Identification of actors",
        ["Actor", "Type", "Description"],
        [
            ["Mobile user / client", "Primary", "Uses the mobile application to manage profile, workouts, nutrition, progress, bookings, packs, programs, and communication."],
            ["Coach", "Primary", "Creates a coach profile, manages clients, programs, session packs, availability, bookings, check-ins, and chat."],
            ["Administrator", "Primary", "Uses the web admin panel to manage users, coaches, exercises, workouts, transactions, notifications, settings, and analytics."],
            ["Supabase", "External system", "Provides authentication, PostgreSQL database, storage, realtime features, edge functions, and row-level security."],
            ["Payment/video/push services", "External services", "Support planned or implemented communication, notification, and marketplace capabilities."],
        ],
        [1.5, 1.2, 3.8],
    )
    add_heading(doc, "2.1.2 Functional Requirements", 3)
    add_basic_table(
        doc,
        "Table 2.2: Functional requirements",
        ["ID", "Requirement", "Actor", "Priority"],
        [
            ["FR-01", "Register, sign in, reset password, and manage authenticated sessions.", "Mobile user / Coach / Admin", "High"],
            ["FR-02", "Complete onboarding and maintain personal profile, goals, units, preferences, and measurements.", "Mobile user", "High"],
            ["FR-03", "Browse exercise library, view details, create custom workouts, and start workout sessions.", "Mobile user", "High"],
            ["FR-04", "Track active workout progress with sets, reps, weights, rest timer, and workout summary.", "Mobile user", "High"],
            ["FR-05", "View progress statistics, body measurements, consistency, and workout history.", "Mobile user", "High"],
            ["FR-06", "Use marketplace features to find coaches, book sessions, buy packs, receive programs, and chat.", "Mobile user / Coach", "Medium"],
            ["FR-07", "Allow coaches to onboard, manage clients, create programs, define availability, and follow check-ins.", "Coach", "Medium"],
            ["FR-08", "Allow administrators to manage users, coaches, exercises, workouts, transactions, notifications, settings, and analytics.", "Administrator", "High"],
        ],
        [0.6, 3.55, 1.55, 0.8],
    )
    add_heading(doc, "2.1.3 Non-Functional Requirements", 3)
    add_basic_table(
        doc,
        "Table 2.3: Non-functional requirements",
        ["ID", "Category", "Requirement"],
        [
            ["NFR-01", "Security", "Authentication must be handled through Supabase Auth and data access must respect RLS policies."],
            ["NFR-02", "Performance", "Mobile screens, lists, and admin pages must remain responsive during repeated daily use."],
            ["NFR-03", "Maintainability", "The codebase must separate screens, components, services, stores, navigation, and database migrations."],
            ["NFR-04", "Usability", "Workout, progress, coach, and admin flows must include clear navigation, loading states, and error messages."],
            ["NFR-05", "Scalability", "The database model must support native workouts, custom workouts, sessions, coaches, clients, messages, and admin growth."],
        ],
        [0.7, 1.1, 4.7],
    )
    add_report_note(
        doc,
        "Requirement interpretation",
        "The non-functional requirements are especially important for GoFit because the platform manipulates personal fitness data, progress photos, coach-client communication, and privileged administration actions. Security, privacy, reliability, and transparency therefore have the same importance as the visible mobile screens."
    )
    add_basic_table(
        doc,
        "Table 2.4: AI and privacy constraints",
        ["Constraint", "Reason", "Design decision"],
        [
            ["Human review of AI outputs", "AI-generated body measurements and coach summaries can be wrong or incomplete.", "Measurements are shown as draft values; coach briefings are assistance notes, not automatic decisions."],
            ["No medical diagnosis", "Fitness data and photos should not be interpreted as clinical information.", "The report and application position AI as progress support only."],
            ["Scoped prompts", "Language models can hallucinate if asked open-ended questions.", "Prompts use only selected profile, workout, booking, note, and check-in data."],
            ["Service-role protection", "Automations and admin workflows need privileged access.", "Service-role keys remain in backend/n8n/server contexts, never in the mobile app."],
            ["On-device CV preference", "Photos are sensitive user data.", "Pose and segmentation are designed primarily for local/native processing where feasible."],
        ],
        [1.55, 2.2, 2.75],
    )
    add_heading(doc, "2.2 Project Structure and Management with Scrum", 2)
    add_heading(doc, "2.2.1 Scrum Team", 3)
    add_basic_table(
        doc,
        "Table 2.5: Scrum team",
        ["Role", "Assigned person", "Responsibility"],
        [
            ["Product Owner", "Supervisor / project owner", "Prioritizes needs and validates the business value of GoFit features."],
            ["Scrum Master", "Student / supervisor", "Organizes sprint rhythm, removes blockers, and tracks progress."],
            ["Development Team", "Student / team members", "Designs, implements, tests, and documents the mobile app, backend, and admin panel."],
            ["Stakeholders", "Academic and professional supervisors", "Review increments and provide feedback."],
        ],
        [1.4, 1.7, 3.4],
    )
    add_heading(doc, "2.2.2 Product Backlog", 3)
    add_basic_table(
        doc,
        "Table 2.6: Product backlog",
        ["ID", "Epic / User story", "Priority", "Planned sprint"],
        [
            ["PB-01", "As a user, I want secure authentication and onboarding.", "High", "Sprint 1"],
            ["PB-02", "As a user, I want to create, schedule, and execute workouts.", "High", "Sprint 1"],
            ["PB-03", "As a user, I want an exercise library and workout history.", "High", "Sprint 1"],
            ["PB-04", "As a user, I want to track progress, measurements, and nutrition.", "High", "Sprint 2"],
            ["PB-05", "As a coach, I want onboarding, profile, client management, programs, bookings, and chat.", "Medium", "Sprint 3"],
            ["PB-06", "As an administrator, I want dashboard, user, coach, exercise, workout, transaction, notification, and settings management.", "High", "Sprint 3"],
            ["PB-07", "As the platform owner, I want deployment, monitoring, and release preparation.", "High", "Closing phase"],
        ],
        [0.65, 4.1, 0.75, 1.0],
    )
    p(doc, "The GoFit backlog was not limited to a simple workout tracker. It was organized as a multi-module platform that covers the client experience, the coach experience, the administration experience, and the backend services required to connect them. The project inventory identifies more than two hundred functional items distributed across mobile client features, coach features, admin-panel features, backend functions, database migrations, realtime communication, storage, notification, and AI-assisted progress tracking.")
    add_basic_table(
        doc,
        "Table 2.7: Detailed feature modules of GoFit",
        ["Module", "Main features", "Implementation elements"],
        [
            ["Authentication and onboarding", "Registration, login, session persistence, password reset, protected navigation, onboarding screens, profile creation, goals, units, preferences, and role-aware flows.", "Supabase Auth, Expo Secure Store, auth store, protected navigators, user_profiles table, onboarding state."],
            ["Workout planning and execution", "Exercise library, custom workouts, native workouts, day splits, workout builder, session tracking, sets, reps, weights, duration, workout summary, and history.", "Exercises, workouts, workout_exercises, workout_sessions, workout services, Zustand stores, active-session screens."],
            ["Enhanced rest timer", "Start, pause, resume, warning intervals, haptic feedback, audio cues, auto-advance settings, persistent timer preferences, and background-aware behavior.", "Timer preferences, mobile UI controls, audio/haptic APIs, workout session integration."],
            ["Progress and measurements", "Progress charts, workout statistics, consistency, body measurements, progress photos, AI-assisted draft measurements, manual correction, and validation notes.", "Body measurement service, MediaPipe native module, MoveNet fallback, segmentation service, measurement tables."],
            ["Nutrition and saved meals", "Nutrition logging foundations, saved meal management, meal item editing, reusable food entries, and history views where enabled.", "Nutrition screens, saved meal services, local and Supabase-backed persistence."],
            ["Coach marketplace", "Coach discovery, coach profiles, onboarding, certification/CV upload, pending validation, session packs, bookings, availability, client list, check-ins, programs, chat, video-session preparation, wallet, and settings.", "Coach profile tables, marketplace screens, booking services, realtime chat, storage uploads, LiveKit token function where applicable."],
            ["Administration panel", "Dashboard statistics, user management, coach validation, exercise/workout management, transactions, notifications, settings, activity logs, search, import/export, analytics, and protected admin routes.", "Next.js App Router, server-side Supabase admin client, middleware, shadcn-style components, admin routes."],
            ["Backend and integrations", "Authentication, PostgreSQL database, row-level security, storage, realtime messaging, push notifications, video-token generation, edge functions, migrations, and deployment configuration.", "Supabase Auth, RLS policies, Edge Functions, Storage buckets, Realtime channels, Expo Push, environment variables."],
        ],
        [1.45, 3.3, 2.2],
    )
    add_basic_table(
        doc,
        "Table 2.8: Feature coverage by product area",
        ["Product area", "Approximate feature count", "Examples"],
        [
            ["Mobile client", "116", "Authentication, onboarding, workouts, active sessions, rest timer, progress, measurements, nutrition, marketplace access, bookings, programs, chat."],
            ["Mobile coach", "35", "Coach onboarding, profile, dashboard, clients, client progress, notes, programs, calendar, availability, wallet, settings."],
            ["Admin panel", "55", "Dashboard, users, coaches, exercises, workouts, transactions, notifications, settings, analytics, imports/exports, activity logs."],
            ["Backend / Edge Functions", "43", "Database migrations, RLS, storage, realtime chat, push notification relay, video-token generation, AI measurement function support."],
        ],
        [1.6, 1.3, 3.6],
    )
    p(doc, "To avoid documenting only planned features, the report also uses a source-tree audit. The implementation evidence was taken from the actual repository areas: mobile screens, mobile services, Zustand stores, admin routes, Supabase Edge Functions, SQL migrations, and exported n8n workflow JSON files. This makes the report closer to the real GoFit product rather than only a theoretical project description.")
    add_basic_table(
        doc,
        "Table 2.9: Source-based implementation coverage",
        ["Repository area", "Implemented coverage identified", "Report sections affected"],
        [
            ["GoFitMobile client screens", "Authentication, onboarding, home, workout library, exercise selection, workout detail, workout builder, active session, workout summary, calendar/timeline, progress statistics, records, consistency, progress photos, body measurement, profile, settings, goals, units, themes, text size, health sync, habits, notifications, marketplace, coach detail, booking, packs, programs, nutrition, barcode food lookup, and check-ins.", "Sprints 1 and 2, mobile realization, AI/computer-vision sections, final validation."],
            ["GoFitMobile coach screens", "Coach welcome, login, signup, onboarding, CV upload, certifications, pending review, profile preview, dashboard, client list, client detail, client progress, notes, check-ins, calendar, availability, profile, settings, wallet, conversations, chat, pack creation, session packs, programs, program builder, and video-call screen.", "Sprint 3 coach marketplace, coach management, communication, video, pack and program sections."],
            ["Mobile services and stores", "Services for auth, profile, workouts, sessions, statistics, plans, exercises, body measurements, segmentation, measurement logging, photos, nutrition, health sync, habits, readiness, notifications, marketplace, bookings, coach profiles, client management, check-ins, chat, calendar, AI session notes, programs, packs, wallet, video calls, recommendations, and milestones. Stores manage auth, booking, calendar, chat, coach/client state, health, language, marketplace, onboarding, packs, programs, sessions, text size, theme, timer, UI, workouts, nutrition, daily coach, and wallet state.", "Architecture, maintainability, non-functional requirements, realization, and testing sections."],
            ["Admin panel", "Protected login, dashboard, workout/exercise CRUD, users, user details, coach management, certifications, transactions, settings, notifications, activity logs, BI snapshot/export/saved views/scheduled digests, imports, bulk operations, health endpoints, and route access checks.", "Sprint 3 administration, backend security, deployment, and final QA."],
            ["Supabase Edge Functions", "Push-notification relay, AI workout recommendation, video-token generation, food-barcode lookup, AI session notes, and shared CORS/runtime support.", "Backend services, AI features, notifications, video calls, nutrition, and deployment."],
            ["Database migrations", "Authentication support tables, workouts, sessions, native workout structures, nutrition, saved meals, barcode attribution, progress photos, health data, recovery/readiness, rest timer preferences, notification preferences, marketplace, coach profiles, bookings, packs, wallet transactions, chat, check-ins, AI templates/session notes, BI views, admin settings, admin notifications, audit logs, and RLS/security optimization.", "Data model, security, backend, feature coverage, architecture, and deployment."],
            ["n8n workflow exports", "AI Session Prep v1, Booking Reminders v1, Check-in Reminders v1, Coach Daily Digest v1, and Admin Ops Digest v1.", "Sprint 3 automation, Chapter 6 n8n automation layer, appendices, validation checklist."],
            ["Testing and quality docs", "Mobile Jest tests for services/stores, admin route and API validation notes, Supabase advisor remediation notes, body-measurement validation protocol, and report-generation audits.", "Sprint tests, final validation, remaining work, and academic transparency."],
        ],
        [1.55, 3.55, 2.1],
    )
    add_heading(doc, "2.2.3 Sprint Planning", 3)
    add_basic_table(
        doc,
        "Table 2.10: Sprint planning used in this report",
        ["Report sprint", "Calendar source", "Main goal", "Main deliverables"],
        [
            ["Sprint 1", "26 Jan - 9 Mar 2026", "Foundation and workout core", "Requirements, UI/UX, setup, authentication, profile, workout planner, rest timer."],
            ["Sprint 2", "9 Mar - 20 Apr 2026", "Advanced mobile tracking", "Calendar, exercise library, progress tracking, notifications, body measurement groundwork."],
            ["Sprint 3", "20 Apr - 29 Jun 2026", "Coach marketplace and administration", "Coach onboarding, marketplace, packs, bookings, chat, programs, admin panel and tests."],
        ],
        [1.0, 1.3, 1.6, 2.6],
    )
    add_diagram(
        doc,
        "figure_2_1_gantt_summary.png",
        "Figure 2.1: Gantt chart summary of GoFit development",
        "GoFit Development Timeline",
        [
            ["Foundation\nJan-Feb", "Auth and Setup\nFeb", "Workout Core\nFeb-Mar"],
            ["Progress and AI\nMar-Apr", "Coach Marketplace\nApr-Jun", "Admin Panel\nJun"],
            ["Tests and Optimization\nJun-Jul", "Deployment\nJul"],
        ],
    )
    add_heading(doc, "2.2.4 Global Use Case Diagram", 3)
    add_diagram(
        doc,
        "figure_2_2_global_use_case.png",
        "Figure 2.2: Global use case diagram",
        "Global Use Case View",
        [
            ["Mobile User", "GoFit Mobile App", "Workouts / Progress / Nutrition"],
            ["Coach", "Coach Marketplace", "Clients / Programs / Bookings"],
            ["Administrator", "Admin Panel", "Users / Coaches / Content / Analytics"],
            ["Supabase", "Auth / Database / Storage", "Secure Platform Data"],
        ],
    )
    add_code_block(doc, "Mermaid global use case diagram draft", mermaid_code_for("global_use_case"))
    add_heading(doc, "2.3 Project Management", 2)
    p(doc, "The project was managed through a backlog and sprint-based organization. Technical work was divided between the mobile app, web admin panel, database migrations, backend functions, documentation, testing, and deployment preparation. The repository structure itself supports this organization through separate folders for GoFitMobile, admin-panel, database, and docs.")
    p(doc, "The project management approach also included technical documentation as a project artifact, not as an afterthought. The repository contains architecture notes, feature inventories, sprint breakdowns, troubleshooting reports, database structure documentation, admin-panel notes, body-measurement research, validation protocols, and automation workflow descriptions. This documentation helped connect implementation choices to report sections and made it possible to justify complex AI and backend decisions.")
    add_basic_table(
        doc,
        "Table 2.11: Project documentation artifacts",
        ["Artifact type", "Purpose in the project", "Examples of information captured"],
        [
            ["Architecture documentation", "Explains the global system structure and deployment view.", "Mobile app, admin panel, Supabase backend, Edge Functions, external services, diagrams."],
            ["Feature inventory", "Lists implemented screens, services, routes, functions, and database artifacts.", "Client features, coach features, admin features, backend functions, migrations, integrations."],
            ["Sprint and Gantt documentation", "Connects implementation work to Agile planning.", "Sprint dates, objectives, deliverables, dependencies, release preparation."],
            ["AI research notes", "Records model decisions, tests, limitations, and validation needs.", "MoveNet, MediaPipe Pose Landmarker, segmentation models, feature vectors, statistical estimator."],
            ["Automation documentation", "Explains scheduled and event-driven backend workflows.", "n8n AI session preparation, admin operations digest, notification creation."],
            ["Academic/report drafts", "Transforms engineering material into PFE reporting structure.", "Problem statement, methodology, report plan, diagram drafts, bibliography entries."],
        ],
        [1.45, 2.4, 2.65],
    )
    add_heading(doc, "Conclusion", 2)
    p(doc, "Sprint 0 established the functional and technical base of the project. It clarified actors, requirements, backlog priorities, sprint grouping, and the global use case perimeter. The next chapter presents Sprint 1, focused on the mobile foundation and core workout features.")
    doc.add_page_break()


def add_ai_computer_vision_details(doc):
    add_heading(doc, "4.3.3 AI Body Measurement and Computer Vision Pipeline", 3)
    p(doc, "One of the most distinctive parts of GoFit is the AI-assisted body measurement feature. Its objective is to help users follow physical progress from photos, without replacing manual tape measurements and without claiming medical or tailoring-grade precision. The feature is designed as an assisted workflow: the system analyzes front and side photos, generates draft values for chest, waist, hip, and shoulder, explains capture-quality issues, then lets the user review and correct the values before saving them as progress data.")
    p(doc, "The pipeline starts on the mobile device. The user captures or selects front and side images. The application runs pose detection to locate key body landmarks, segmentation to isolate the person silhouette when available, feature extraction to transform image geometry into measurable indicators, and a measurement estimator to produce draft values. The result is presented with confidence, warnings, and manual editing fields. This design is important because computer vision can support progress tracking, but a single two-dimensional image cannot directly recover exact body circumferences in all conditions.")
    add_diagram(
        doc,
        "figure_4_6_ai_cv_pipeline.png",
        "Figure 4.7: AI and computer vision pipeline for body measurements",
        "GoFit AI Body Measurement Pipeline",
        [
            ["Front and Side Photos", "Pose Detection", "Segmentation"],
            ["Feature Vector", "Statistical Estimator", "Draft Measurements"],
            ["Quality Checks", "Manual Review", "Saved Progress"],
        ],
    )

    add_heading(doc, "4.3.4 Models and Artifacts Used", 3)
    p(doc, "The project evolved from a keypoint-only prototype toward a more robust pose-and-segmentation architecture. The first implementation used MoveNet Lightning through react-native-fast-tflite because it is lightweight and suitable for on-device inference. Later research and implementation introduced MediaPipe Tasks Vision because its pose landmarker provides richer landmarks and a native Android path for faster and more reliable analysis.")
    add_basic_table(
        doc,
        "Table 4.5: AI and computer vision models used in GoFit",
        ["Model / artifact", "Role in GoFit", "Input / output", "Status and reason"],
        [
            ["MoveNet Lightning (.tflite)", "Initial on-device pose detection and fallback path.", "Resized 192 x 192 RGB image; 17 body keypoints with confidence scores.", "Implemented through react-native-fast-tflite. Useful as a lightweight baseline, but limited for body measurement because 17 keypoints do not describe torso outline or body depth."],
            ["MediaPipe Pose Landmarker Full (pose_landmarker_full.task)", "Main candidate for richer body-pose analysis.", "Image input; 33 normalized landmarks, optional world landmarks, visibility/presence scores, pose count, and inference time.", "Implemented in the local Expo native module on Android. It improves body coverage over MoveNet and supports better pose-quality checks."],
            ["MediaPipe Pose Landmarker Heavy", "Potential higher-accuracy candidate.", "Same landmark structure as the Full model, with heavier computation.", "Not adopted by default. It is reserved for comparison only if Full is accurate enough and device performance remains acceptable."],
            ["selfie_multiclass_256x256.tflite", "Experimental segmentation research artifact.", "256 x 256 float image input; 256 x 256 x 6 class output.", "Investigated but not treated as the final silhouette model because the body mask was unstable in mirror/direct capture scenarios."],
            ["MediaPipe Image Segmenter (selfie_segmenter.tflite)", "Person-mask extraction for segmentation-assisted measurement.", "Image input; 256 x 256 confidence masks and category mask for background/person.", "Integrated in the native MediaPipe module. It provides a cleaner person mask than the earlier multiclass experiment, but it still needs validation and gating."],
            ["Statistical depth estimator", "Converts extracted features into draft circumferences.", "Feature vector using height, front widths, side depths, ratios, pose scores, and mask quality indicators.", "Used as a pragmatic baseline. It is not a trained medical model; it produces draft progress estimates and must remain editable by the user."],
            ["Future regression/PCA estimator", "Target research direction for stronger measurement prediction.", "Labeled dataset of body features and manual/tape measurements.", "Planned improvement based on statistical body-measurement literature. It would replace rigid formulas with learned predictors once enough validation data exists."],
        ],
        [1.55, 1.75, 1.65, 1.95],
    )

    add_heading(doc, "4.3.5 Pose Detection: MoveNet and MediaPipe", 3)
    p(doc, "Pose detection is responsible for finding anatomical landmarks that anchor the measurement process. In the first GoFit prototype, MoveNet Lightning detects 17 keypoints such as nose, shoulders, elbows, wrists, hips, knees, and ankles. These points are enough to estimate body orientation and basic proportions, but they are not enough to measure chest, waist, and hip circumference directly. For example, a shoulder keypoint tells where the shoulder is, but it does not describe the width of the torso at chest level or the depth of the body from the side view.")
    p(doc, "The MediaPipe Pose Landmarker Full model improves this stage by producing 33 landmarks and richer metadata. The local module returns image dimensions, pose count, normalized landmarks, world landmarks when available, and inference time. The GoFit Android implementation loads the bundled pose_landmarker_full.task asset, decodes the image URI, applies EXIF rotation, runs MediaPipe detection, and returns a typed result to the TypeScript layer. This makes the measurement workflow more reliable because the application can inspect landmark visibility, check whether core points are present, and reject only captures that are genuinely unusable.")
    p(doc, "The project keeps MoveNet as a fallback because it is already integrated in the mobile JavaScript service and can continue to support devices or builds where the native MediaPipe module is not available. This fallback strategy reduces deployment risk while the native module reaches parity across Android and iOS.")

    add_heading(doc, "4.3.6 Segmentation and Body Mask Extraction", 3)
    p(doc, "Segmentation is the second computer-vision component. Its role is to separate the person from the background and provide body-outline information. This is necessary because circumference estimation cannot be based only on joint coordinates. The pose model can locate shoulder, hip, and torso reference lines; the segmentation mask can then estimate front-view widths and side-view depths around the chest, waist, and hip levels.")
    p(doc, "The project first examined selfie_multiclass_256x256.tflite. That model exposes a 256 x 256 x 6 output rather than a simple binary body mask, which required class selection, argmax maps, connected-component cleanup, and debugging overlays. Real tests showed that the resulting mask could be unstable, especially with mirror photos, phone occlusion, and direct capture conditions. For this reason, the report classifies it as an experimental research artifact instead of a final production model.")
    p(doc, "The later implementation uses MediaPipe Image Segmenter with selfie_segmenter.tflite in the local native module. The Android bridge returns confidence masks, category masks, labels, quality scores, and inference timing. The TypeScript segmentation service decodes the masks, thresholds the person class, builds a cleaned binary mask, extracts sampled class grids for debugging, anchors scan lines using pose landmarks, and measures chest, waist, and hip widths/depths. This gives GoFit a more explainable measurement pipeline because each generated number can be traced back to pose landmarks, mask rows, and quality gates.")

    add_diagram(
        doc,
        "figure_4_7_model_flow.png",
        "Figure 4.8: Pose and segmentation model flow",
        "Model Flow for AI Measurements",
        [
            ["MoveNet 17 Keypoints", "Fallback Pose Path", "Basic Quality Checks"],
            ["MediaPipe 33 Landmarks", "Primary Pose Path", "Landmark Visibility"],
            ["MediaPipe Segmenter", "Person Mask", "Widths and Depths"],
            ["Estimator", "Confidence and Warnings", "Editable Result"],
        ],
    )

    add_heading(doc, "4.3.7 Feature Vector and Measurement Estimation", 3)
    p(doc, "The measurement service builds a structured feature vector before calculating the final draft values. This is a strong architectural choice because the application does not mix raw pixels, pose confidence, segmentation quality, and final centimeter values in a single opaque formula. Instead, it records the intermediate indicators that explain why a measurement was accepted, warned, or rejected.")
    add_basic_table(
        doc,
        "Table 4.6: Body measurement feature vector",
        ["Feature group", "Examples", "Purpose"],
        [
            ["User scale", "Declared height in centimeters, person height in pixels, scale in cm/px, height-span fraction.", "Converts image measurements into approximate real-world dimensions and detects bad framing."],
            ["Pose quality", "Front and side pose model source, mean landmark score, visible core keypoints, body center, top and bottom normalized positions.", "Checks whether the user is visible, centered, and sufficiently detected."],
            ["Front geometry", "Front shoulder width, hip width, estimated waist width, chest/waist/hip mask widths.", "Measures visible body widths from the front photo."],
            ["Side geometry", "Side shoulder width, side-to-front ratio, chest depth, abdomen depth, side mask depths.", "Approximates body depth, which is essential for circumference estimation."],
            ["Segmentation quality", "Raw coverage, clean coverage, body class index, mask bounds, fragmentation, class grid, person threshold.", "Detects unstable masks and prevents unreliable segmentation from overriding safer estimates."],
            ["Estimator outputs", "Draft chest, waist, hip, shoulder, depth source, confidence, failed checks, quality issues.", "Produces editable progress values and explains uncertainty to the user."],
        ],
        [1.5, 3.0, 2.0],
    )
    p(doc, "The estimator currently combines geometric measurements with sex-aware statistical depth assumptions. In practical terms, the service uses the user's height and body proportions to estimate a scale, then uses front widths and side/depth information to produce draft circumferences. When segmentation depth passes quality and sanity gates, it can contribute to the estimate; otherwise the service falls back to statistical priors. This prevents a bad mask from producing extremely wrong values.")
    p(doc, "Research notes in the project also identify a stronger long-term direction: use a labeled dataset of photos and manual measurements to train regression, ridge regression, gradient-boosted trees, PCA-based, or body-shape models. Such a model would learn the relationship between extracted features and real measurements instead of relying mainly on hand-built formulas. For the PFE scope, the current implementation is correctly presented as an AI-assisted progress estimator with validation and manual correction.")

    add_heading(doc, "4.3.8 Validation Results and Limitations", 3)
    p(doc, "Device tests confirmed that MediaPipe can run inside the Android development client and that the pose model detects one person with high landmark visibility in normal capture conditions. Example observations showed mirror-front captures around 0.99 pose score, side captures around 0.97 to 0.98 in good conditions, and inference times that varied by device and image conditions. Direct capture also worked but may show different inference times and body framing behavior.")
    p(doc, "The main limitation is not pose detection; it is reliable circumference estimation from consumer photos. Early segmentation experiments produced thin or unstable masks in some scans, and earlier formulas generated unstable chest, waist, and hip values. Later statistical-depth experiments were more stable across repeated scans, but the project documentation correctly notes that those values are still draft estimates, not ground truth. Therefore, the application keeps a manual review screen and should store corrected values when the user has measured with a tape.")
    add_basic_table(
        doc,
        "Table 4.7: AI measurement validation protocol",
        ["Validation item", "Expected practice", "Target / interpretation"],
        [
            ["Reference measurements", "Measure chest, waist, hip, and shoulder with a tape before comparing with the AI result.", "The tape value is the reference value for validation."],
            ["Capture conditions", "Use good lighting, full body visible, fitted clothes, front photo, side photo at approximately 90 degrees, and consistent distance.", "Bad framing, loose clothing, phone occlusion, and angled mirrors must be flagged."],
            ["AI result", "Record the draft AI value before manual correction.", "Used to calculate absolute error and identify model weaknesses."],
            ["Manual correction", "Allow the user to edit the values before saving.", "Saved corrected values should match the user's real measurement."],
            ["Fitness-progress target", "Use approximate progress targets rather than tailoring or medical targets.", "Chest error ideally below 8 cm; waist and hip below 6 cm; shoulder below 5 cm."],
            ["Confidence", "Confidence must decrease when pose, segmentation, scale, or repeated-scan stability is weak.", "The report should avoid claiming 100 percent confidence for unstable captures."],
        ],
        [1.45, 3.0, 2.05],
    )
    add_basic_table(
        doc,
        "Table 4.8: AI implementation status",
        ["Component", "Current status", "Remaining work"],
        [
            ["MoveNet pose fallback", "Implemented in TypeScript service through react-native-fast-tflite and bundled model asset.", "Keep as fallback and compare against MediaPipe outputs."],
            ["MediaPipe pose Android module", "Implemented with bundled pose_landmarker_full.task, EXIF rotation handling, 33 landmarks, world landmarks, pose count, and inference time.", "Complete parity and production testing across iOS and Android release builds."],
            ["MediaPipe image segmentation", "Implemented in native module with selfie_segmenter.tflite, confidence/category masks, labels, and quality scores.", "Improve segmentation quality gates and verify on more body types, clothes, lighting, and camera setups."],
            ["Feature vector", "Implemented with pose, scale, geometry, segmentation, confidence, and debug fields.", "Export validation datasets and compare AI estimates with tape measurements."],
            ["Estimator", "Statistical baseline and segmentation-gated depth source available.", "Train or calibrate a stronger model once enough labeled examples exist."],
            ["User workflow", "Draft measurement, quality warnings, and manual review approach defined.", "Finalize UX copy, screenshots, and submission-ready validation evidence."],
        ],
        [1.6, 2.55, 2.4],
    )
    p(doc, "In conclusion, the AI module is a meaningful technical contribution because it combines on-device inference, native module integration, pose landmarks, segmentation masks, typed feature vectors, statistical estimation, and validation methodology. At the same time, the report must be transparent: GoFit uses AI to assist progress tracking, not to replace professional body measurement.")


def add_sprint1_implementation_details(doc):
    add_heading(doc, "3.3.1 Mobile Foundation and Navigation", 3)
    p(doc, "The first implementation challenge was to create a stable mobile foundation that could support both client and coach experiences. The application was structured around authentication screens, onboarding screens, protected app navigators, profile flows, and domain-specific feature stacks. This foundation is important because the rest of the product depends on role-aware routing: a visitor should not reach protected screens, an incomplete profile should be redirected to onboarding, and coach-specific flows should remain separated from standard client flows.")
    p(doc, "State management was organized with stores so that authentication state, profile state, workout state, timer preferences, and user settings could be reused across screens without duplicating logic. Service files were used to isolate communication with Supabase and local modules. This separation makes the mobile code easier to test and maintain because UI components do not directly contain all persistence logic.")
    add_basic_table(
        doc,
        "Table 3.4: Sprint 1 implementation components",
        ["Component", "Role", "Implementation detail"],
        [
            ["Authentication flow", "Protects access to user data and feature screens.", "Supabase Auth, session persistence, auth store, protected navigation, login/signup/reset screens."],
            ["Onboarding flow", "Collects first profile information and directs users to the right experience.", "Profile setup, goals, preferences, role-aware continuation, user_profiles persistence."],
            ["Navigation structure", "Organizes the application into coherent areas.", "Auth navigator, onboarding navigator, client app navigator, coach app navigator, modal/detail screens."],
            ["Service layer", "Centralizes data access.", "Supabase client calls, typed methods, error handling, loading states, reusable domain services."],
            ["UI component base", "Creates reusable screens and controls.", "Buttons, cards, lists, forms, empty states, loading indicators, timer controls, workout cells."],
        ],
        [1.5, 2.1, 2.9],
    )
    add_heading(doc, "3.3.2 Workout Core and Rest Timer", 3)
    p(doc, "The workout core was implemented around the distinction between a workout template and an executed workout session. A workout template defines the exercises, order, difficulty, target muscle groups, and expected configuration. A workout session records what the user actually performed: start time, completion time, sets, reps, weights, notes, duration, and summary data. This separation is essential because the same workout can be performed many times while producing different session histories.")
    p(doc, "The rest timer was treated as a real workout tool rather than a decorative countdown. It includes start, pause, resume, warning cues, haptic feedback, audio cues, auto-advance options, and persisted preferences. This improves usability during training because users often interact with the timer under fatigue and need clear feedback without navigating away from the session.")
    add_basic_table(
        doc,
        "Table 3.5: Workout-session data flow",
        ["Step", "User action", "System behavior"],
        [
            ["1", "User selects or creates a workout.", "The app loads workout metadata, exercise list, and saved configuration."],
            ["2", "User starts a session.", "A session state is initialized with start time, current exercise, set counters, and timer preferences."],
            ["3", "User completes sets.", "The app records repetitions, weights, rest periods, and exercise completion state."],
            ["4", "Rest timer runs between sets.", "Timer UI, haptics, audio cues, warning intervals, and auto-advance settings are applied."],
            ["5", "User finishes the workout.", "Summary values are calculated and stored in workout_sessions for progress screens."],
        ],
        [0.6, 2.0, 3.9],
    )


def add_sprint2_progress_details(doc):
    add_heading(doc, "4.3.1 Progress, Nutrition, and Reminder Features", 3)
    p(doc, "Sprint 2 connected the workout foundation to user progress. A fitness application becomes more valuable when it transforms repeated sessions into understandable feedback. For this reason, the sprint added progress screens, history views, statistics, body-measurement foundations, nutrition-related flows, and reminder infrastructure. These features help users move from isolated workout execution to long-term self-monitoring.")
    p(doc, "Progress tracking uses workout session data, body-measurement data, and optional nutrition entries to display user evolution. Empty states are important in this sprint because new users may not have enough data yet. The interface should explain what the user can do next without showing broken charts or misleading statistics.")
    add_basic_table(
        doc,
        "Table 4.4: Sprint 2 user-facing progress modules",
        ["Module", "Objective", "Important states"],
        [
            ["Exercise library", "Allow users to browse and understand available exercises.", "Search result, filter result, detail screen, no-result state, loading state."],
            ["Workout calendar/history", "Help users see planned and completed training.", "Scheduled workout, completed workout, missed day, empty calendar, history detail."],
            ["Progress dashboard", "Summarize evolution over time.", "Charts with data, first-use empty state, metric selector, date-range filtering."],
            ["Body measurements", "Record and compare physical changes.", "Manual entry, AI draft, review/correction, quality warning, saved measurement."],
            ["Nutrition/saved meals", "Support nutrition tracking foundations.", "Saved meal list, meal detail, item edit, history, reusable entry."],
            ["Notifications", "Encourage consistency and communication.", "Permission request, reminder settings, notification inbox, push-token registration."],
        ],
        [1.4, 2.3, 2.8],
    )
    add_heading(doc, "4.3.2 Data Interpretation and User Feedback", 3)
    p(doc, "A key design point in Sprint 2 is that GoFit should not only store values; it should help the user interpret them. For example, workout history becomes useful when the app can show consistency, completed sessions, training load, and trends. Body measurements become useful when the user can compare corrected values over time and understand why an AI draft may be uncertain. Nutrition entries become useful when they are reusable and connected to habits rather than isolated forms.")
    p(doc, "The sprint therefore uses feedback patterns such as charts, summaries, empty states, warnings, and review screens. This is especially important for AI-assisted measurements. Instead of hiding uncertainty, the application should show quality issues such as poor pose, incomplete body visibility, weak segmentation, bad framing, or inconsistent scale between front and side photos.")
    add_diagram(
        doc,
        "figure_4_6_progress_feedback_loop.png",
        "Figure 4.6: Progress feedback loop",
        "Progress Feedback Loop",
        [
            ["Workout / Nutrition / Photo", "Data Processing", "Metric History"],
            ["Charts and Summaries", "User Interpretation", "Habit Adjustment"],
            ["Coach Context", "AI Assistance", "Manual Validation"],
        ],
    )


def add_sprint3_platform_details(doc):
    add_heading(doc, "5.3.1 Coach Marketplace and Client Follow-up", 3)
    p(doc, "Sprint 3 expanded GoFit from an individual fitness tracker into a marketplace and coaching platform. This required a different type of workflow because the application now connects two human roles: the client and the coach. The client needs to discover coaches, view profiles, book sessions, purchase packs, access programs, and communicate. The coach needs to onboard, submit profile and document information, wait for validation when required, manage availability, consult clients, create programs, read notes, and follow progress.")
    p(doc, "The marketplace is not only a list of coaches. It depends on trust, availability, transaction status, communication, and follow-up. For this reason, the sprint includes coach profiles, session packs, bookings, purchased packs, conversations, messages, programs, notifications, and wallet-related views. Each of these elements must be connected to permissions so that clients can see their own bookings and coaches can see only their own clients.")
    add_basic_table(
        doc,
        "Table 5.4: Coach marketplace feature detail",
        ["Feature", "Client-side value", "Coach-side value"],
        [
            ["Coach discovery", "Find coaches by profile, expertise, and services.", "Increase visibility and attract clients."],
            ["Coach onboarding", "Improves trust by requiring complete coach data.", "Allows coach to create a professional presence."],
            ["Session packs", "Buy or use grouped coaching sessions.", "Sell structured services and track remaining sessions."],
            ["Bookings", "Reserve a coaching slot and follow status.", "Organize calendar and session workload."],
            ["Programs", "Receive personalized training plans.", "Create and assign plans to clients."],
            ["Chat and notes", "Communicate with the coach and receive follow-up.", "Keep client context and preparation notes together."],
            ["Video call screen", "Join remote coaching sessions.", "Connect with clients through LiveKit-backed sessions."],
        ],
        [1.45, 2.35, 2.35],
    )
    add_heading(doc, "5.3.2 Administration and Operational Control", 3)
    p(doc, "The administration panel is necessary because the platform contains user-generated data, coach applications, exercise content, workouts, transactions, notifications, and analytics. Without a back-office interface, the operator would need to manipulate database rows directly, which is risky and unsuitable for a production platform. The admin panel provides safer and more structured operational control.")
    p(doc, "The admin panel uses Next.js App Router and protected admin access. Server-side Supabase utilities are used for privileged operations, while middleware and admin-role checks prevent normal users from opening back-office pages. The design includes dashboard statistics, user management, coach validation, exercise/workout management, transaction review, notification management, settings, logs, analytics, search, and import/export foundations.")
    add_basic_table(
        doc,
        "Table 5.5: Admin-panel operational modules",
        ["Admin area", "Managed data", "Purpose"],
        [
            ["Dashboard", "Users, exercises, workouts, sessions, transactions, coach status.", "Provide a fast overview of platform health."],
            ["Users", "Auth users and profile data.", "Search, review, update, and monitor user accounts."],
            ["Coaches", "Coach profiles, status, documents, marketplace visibility.", "Validate coaches and maintain service quality."],
            ["Exercises and workouts", "Exercise library, workout templates, metadata.", "Maintain content consistency for mobile users."],
            ["Transactions and packs", "Purchases, wallet/pack status, booking-related financial records.", "Support marketplace operations and revenue visibility."],
            ["Notifications and logs", "User notifications, admin activity, operational events.", "Trace actions and communicate important updates."],
            ["Settings and analytics", "Platform configuration and business indicators.", "Prepare the product for scaling and decision-making."],
        ],
        [1.5, 2.45, 2.2],
    )
    add_heading(doc, "5.3.3 Admin Analytics and Business Intelligence Features", 3)
    p(doc, "The administration panel also includes a business intelligence dimension. This work is more than a simple dashboard counter: it transforms operational data from users, workouts, bookings, coaches, session packs, finance records, and health/readiness signals into indicators that an administrator can use for monitoring and decision-making. The BI area is implemented inside the main dashboard so that operational control and performance interpretation stay in the same protected back-office interface.")
    p(doc, "The first analytics layer provides platform visibility: registered users, daily/weekly/monthly active users, user growth, popular exercises, workout completion rates, activity heatmaps, session activity, and recent activity. The advanced BI layer adds deeper decision-support views: finance by currency and coach/package scope, user lifecycle and activation signals, workout cohort retention, coach operations, client health risk queues, threshold alerts, saved BI views, snapshot notifications, scheduled digests, and CSV exports.")
    add_basic_table(
        doc,
        "Table 5.6: Admin analytics and BI dashboard features",
        ["BI feature", "Source / implementation", "Administrative value"],
        [
            ["Engagement metrics", "admin-panel/lib/analytics.ts; EngagementMetricsCards; workout_sessions and user_profiles.", "Tracks total users, DAU, WAU, and MAU to measure platform activity."],
            ["User growth chart", "UserGrowthChart; user_profiles.created_at grouped by date.", "Shows acquisition trend and cumulative user growth."],
            ["Popular exercises", "PopularExercisesCard; exercises joined with workout_exercises usage counts.", "Identifies content that is most used by members."],
            ["Workout completion and session activity", "WorkoutCompletionCard and SessionActivityChart; workout_sessions.", "Detects completion rate, training volume, and possible workout drop-off."],
            ["Activity heatmap and recent activity", "ActivityHeatmap and RecentActivityFeed; completed sessions and activity records.", "Shows when users are active and what happened recently on the platform."],
            ["Advanced BI overview", "admin-panel/app/dashboard/page.tsx with bi-finance, bi-user-lifecycle, bi-coach-ops, and bi-client-health services.", "Combines finance, lifecycle, coach operations, and client health into one decision panel."],
            ["Threshold alerts", "BIThresholdAlertsCard with client-health rows.", "Highlights users with risk signals so administrators can notice operational issues earlier."],
            ["Saved BI views", "AdvancedBISavedViews and /api/bi/saved-views.", "Allows administrators to save BI filter/range configurations and digest cadence."],
            ["BI snapshots and scheduled digests", "AdvancedBISnapshotButton, /api/bi/snapshot, and /api/bi/scheduled-digests.", "Creates admin notification summaries on demand or according to saved-view cadence."],
            ["CSV exports", "/api/bi/export with finance, lifecycle, cohort, coach-ops, and client-health slices.", "Lets administrators extract BI data for external reporting or academic evidence."],
        ],
        [1.55, 2.7, 2.25],
    )
    add_basic_table(
        doc,
        "Table 5.7: Advanced BI data domains and limits",
        ["Domain", "Database / service evidence", "Metrics covered", "Known limit in v1"],
        [
            ["Finance", "bi_finance_daily view; admin-panel/lib/bi-finance.ts; FinanceCurrencyDetailCard.", "Gross pack sales, pack count, ledger signals, currency summaries, coach/package filters.", "Net revenue, refunds, payout reconciliation, and full accounting remain deferred."],
            ["User lifecycle", "bi_user_lifecycle_daily view; admin-panel/lib/bi-user-lifecycle.ts; LifecycleActivationDetailCard.", "Signup, first workout activation, first booking activation, activity windows, DAU/WAU/MAU.", "Coach filters do not scope lifecycle cohorts in v1; lifecycle may include non-client profiles where relevant."],
            ["Workout cohort retention", "bi-user-lifecycle cohort helper; RetentionCohortCard.", "Signup-to-workout retention by cohort period.", "Cohorts depend on available historical workout data."],
            ["Coach operations", "bi_coach_ops_daily view; admin-panel/lib/bi-coach-ops.ts; CoachOpsDetailCard and CoachPerformanceTable.", "Booking activity, completion rate, cancellation/no-show signals, coach activity, package context.", "Coach filters scope coach operations, but not every dashboard card is coach-scoped."],
            ["Client health", "bi_client_health_daily view; admin-panel/lib/bi-client-health.ts; ClientHealthRiskQueue and ClientHealthTrendDetailCard.", "Inactive clients, nutrition gaps, body-measurement gaps, pack/session risk signals, aggregated risk counts.", "Risk indicators are operational signals, not medical diagnosis."],
            ["Security and access", "fix_bi_views_security_invoker.sql; admin middleware; server-side Supabase utilities.", "Admin-only dashboard and BI API access with protected routes.", "Final QA must still attach admin/non-admin access screenshots and API test outputs."],
        ],
        [1.2, 2.05, 2.25, 1.65],
    )
    p(doc, "This BI layer is important academically because it shows that GoFit is not only a mobile user interface. The project also studies how operational data can be modeled, aggregated, protected, visualized, exported, and summarized for administrators. The current implementation remains intentionally transparent about v1 limitations so that the report does not exaggerate financial reconciliation, medical interpretation, or filter completeness.")
    add_heading(doc, "5.3.4 In-App AI Workout Recommendation and Personalization", 3)
    p(doc, "The in-app AI workout recommendation feature deserves separate treatment because it is a user-facing personalization feature, not a computer-vision feature and not an n8n scheduled automation. Its goal is to generate one practical custom workout from the user's real GoFit context, then let the user decide whether to save it into the workout library. In the mobile application, this feature appears as the Adaptive workout card implemented in GoFitMobile/src/components/home/RecommendedWorkouts.tsx.")
    p(doc, "The mobile flow starts from the home screen. When the user taps the Adaptive workout card, the component calls workoutRecommendationService.generateAIWorkout from GoFitMobile/src/services/workoutRecommendations.ts. This service invokes the Supabase Edge Function named ai-workout-recommendation. If the response is valid, the card displays the generated workout name, focus/reason text, number of exercises, and an adaptation label such as readiness level, days since last workout, coach companion context, and volume adjustment. If the user accepts the suggestion, the feature saves it as a custom workout with workoutService.createCustomWorkout, reloads the workout store, triggers success haptic feedback, and navigates to the workout detail screen.")
    p(doc, "On the backend, the ai-workout-recommendation Edge Function authenticates the caller through the Authorization header before collecting context. It reads the user's profile, recent completed workout sessions, the latest daily readiness snapshot, recent health data, the latest assigned program, any active purchased coach pack, and the exercise catalog. The function computes an adaptive context before calling the language model. This context includes readiness level, readiness score, days since last workout, volume adjustment, intensity guidance, recovery signals, coach-program state, and explicit constraints.")
    p(doc, "The AI model used for this feature is Groq's OpenAI-compatible chat completion endpoint with the model llama-3.3-70b-versatile. The function requests JSON output with a low temperature and instructs the model to create one workout only, use only exercises from the provided catalog, avoid invented exercise identifiers, reduce volume when readiness is low, propose moderate catch-up work after several inactive days, and treat the result as optional companion work when a coach program or active pack exists. This prompt design is important because it grounds the LLM output in GoFit data instead of allowing generic fitness advice.")
    p(doc, "The feature also contains post-generation validation. The returned JSON is parsed, every exercise ID is matched against the real exercise catalog, invalid exercises are discarded, and the final recommendation is rejected if fewer than three valid exercises remain. The maximum exercise count is also adapted: reduced-volume sessions are limited more strictly than normal sessions. The response sent back to the mobile app includes the final workout name, difficulty set to Custom, focus, reason, image, validated exercise list, and adaptation metadata. This means the mobile app does not blindly trust the LLM response; it receives a cleaned and structured result.")
    add_basic_table(
        doc,
        "Table 5.8: In-app AI workout recommendation workflow",
        ["Step", "Implementation evidence", "Description"],
        [
            ["User entry point", "RecommendedWorkouts.tsx", "Displays the Adaptive workout card on the home screen with loading, error, generated-result, and save states."],
            ["Mobile service call", "workoutRecommendations.ts", "Calls supabase.functions.invoke('ai-workout-recommendation') and rejects unusable responses without exercises."],
            ["Backend function", "supabase/functions/ai-workout-recommendation/index.ts", "Authenticates the user, gathers context, builds the prompt, calls Groq, validates the output, and returns a typed recommendation."],
            ["Model call", "Groq chat completions; llama-3.3-70b-versatile; response_format json_object", "Generates one custom workout using controlled JSON output and a constrained exercise catalog."],
            ["Save action", "workoutService.createCustomWorkout; useWorkoutsStore.loadWorkouts", "Stores the accepted AI recommendation as a normal custom workout so it can be opened and reused like other workouts."],
            ["User feedback", "Expo Haptics; toastManager; navigation to WorkoutDetail", "Provides tactile confirmation, error feedback, and direct access to the generated workout after saving."],
        ],
        [1.3, 2.35, 2.85],
    )
    add_basic_table(
        doc,
        "Table 5.9: Personalization inputs and adaptation rules",
        ["Input / rule", "Source table or code", "Role in recommendation"],
        [
            ["User profile", "user_profiles: goal, activity_level, age, gender", "Anchors the workout to the user's declared objective and fitness context."],
            ["Recent training history", "workout_sessions joined with workouts; last 8 completed sessions", "Avoids recommending workouts without considering what the user recently trained and when."],
            ["Readiness snapshot", "daily_readiness; computeReadiness in readiness.ts", "Uses score, level, and recommendation text to reduce, maintain, or increase training volume."],
            ["Recovery and health signals", "health_data: steps, active calories, sleep, resting heart rate, HRV", "Supports recovery-aware constraints such as lowering intensity after short sleep."],
            ["Coach context", "custom_programs and purchased_packs", "Prevents the AI suggestion from replacing coach-assigned programming; frames it as companion work when relevant."],
            ["Exercise catalog", "exercises table; first 120 catalog rows", "Forces the model to select only existing exercises and allows backend validation of each returned ID."],
            ["Adaptation logic", "computeAdaptiveContext in ai-workout-recommendation/index.ts", "Creates volumeAdjustment, intensityGuidance, constraints, recovery, and coachContext before the LLM call."],
            ["Output validation", "exerciseMap filtering; min 3 valid exercises; max 5 or 8 depending on volume", "Blocks invented or unusable model output before it reaches the mobile UI."],
        ],
        [1.5, 2.25, 2.75],
    )
    add_diagram(
        doc,
        "figure_5_6_ai_workout_recommendation_flow.png",
        "Figure 5.6: In-app AI workout recommendation flow",
        "Adaptive AI Workout Recommendation Flow",
        [
            ["Home Adaptive Card", "Mobile Recommendation Service", "Supabase Edge Function"],
            ["User Context", "Adaptive Context Builder", "Groq JSON Generation"],
            ["Catalog Validation", "Recommendation Payload", "Save as Custom Workout"],
        ],
    )
    p(doc, "This feature is therefore best described as a supervised personalization workflow. The model proposes a workout, but GoFit controls the context, validates the result, keeps secrets on the backend, shows the reasoning and adaptation label to the user, and stores the workout only after the user chooses to save it. The current implementation is appropriate for a PFE scope because it demonstrates LLM integration, personalization, prompt control, response validation, mobile interaction, and human-in-the-loop acceptance without claiming that AI replaces a professional coach.")

    add_heading(doc, "5.3.5 LLM-Based Coach Briefings and Automations", 3)
    p(doc, "Sprint 3 also includes language-model assistance for coach preparation and operational automation. Unlike the in-app workout recommendation, these features do not directly create a workout for the mobile user. They summarize context, prepare notes, and create notifications so that coaches and administrators can act faster.")
    p(doc, "The AI coach briefing flow prepares concise pre-session notes for a coach. It gathers booking information, recent completed sessions, private coach notes, recent check-ins, and client profile data. The model is instructed to summarize only the provided information, avoid invented medical facts, and keep the output short. The resulting note helps the coach prepare a session while still requiring human judgment.")
    add_basic_table(
        doc,
        "Table 5.10: LLM and automation features",
        ["Feature", "Input context", "Generated output", "Safety/control"],
        [
            ["AI coach/session briefing", "Booking, client profile, recent completed sessions, private coach notes, recent check-ins.", "Concise pre-session briefing for coach preparation.", "Coach-only context, authentication checks, instruction not to invent injuries or diagnoses."],
            ["n8n AI Session Prep v1", "Confirmed bookings in the next 24 hours, client profile, recent sessions, coach notes, and check-ins retrieved from Supabase.", "Stored ai_session_notes row and coach notification with data.kind = ai_session_ready.", "Unexpired-note guard avoids duplicate briefings; service-role key remains in n8n; Groq key is used only by this workflow."],
            ["n8n Booking Reminders v1", "Confirmed bookings starting within the next 60 minutes and mapped coach/client recipients.", "Client and coach notification rows with data.kind = booking_starting_soon.", "Duplicate guard uses booking, recipient role, and reminder window."],
            ["n8n Check-in Reminders v1", "Enabled check-in schedules due in UTC and current-day check-in responses.", "Client notification row with data.kind = check_in_due.", "Skips schedules not yet due, clients who already responded, and duplicate reminders."],
            ["n8n Coach Daily Digest v1", "Coach profiles, check-in schedules/responses, upcoming bookings, active packs, and BI client-health data.", "Coach notification row with missed check-ins, low wellness signals, upcoming sessions, and inactive clients.", "Skips coaches with no useful signal and avoids duplicate digest notifications for the same UTC date."],
            ["n8n Admin Ops Digest v1", "Admin users, new users, pending coaches, booking outcomes, and pack purchases from the last 24 hours.", "Admin notification row containing operational counts and a deterministic digest link.", "Skips duplicate admin digests and keeps automation failure count deterministic in v1."],
        ],
        [1.55, 2.2, 1.7, 2.05],
    )
    add_diagram(
        doc,
        "figure_5_6_llm_automation_flow.png",
        "Figure 5.7: LLM and automation flow",
        "LLM Assistance and Automation Flow",
        [
            ["Supabase Data", "Prompt Builder", "Groq LLM"],
            ["Recommendation / Briefing", "Saved Note", "Notification"],
            ["User or Coach Review", "Human Decision", "Follow-up Action"],
        ],
    )


def add_ai_development_tooling_details(doc):
    add_heading(doc, "6.2.8 AI-Assisted Development Methodology", 3)
    p(doc, "GoFit was not developed only with traditional manual coding tools. The project also used AI-assisted software-engineering tools during analysis, implementation support, debugging, documentation, and report construction. This aspect must be documented because the internship concerns both the final GoFit platform and the way the platform was produced. The use of AI assistants changed the workflow: instead of only searching manually through files and documentation, the developer could ask targeted questions, compare implementation alternatives, generate structured drafts, and verify project evidence faster.")
    p(doc, "Codex was used as the main AI coding assistant. It worked inside the local GoFit workspace and had access to the repository files, terminal commands, generated documentation, and selected project tools. Codex helped inspect the mobile application, admin panel, database migrations, Supabase functions, n8n workflows, AI research notes, and academic documents. Its role was not to replace the developer; its role was to accelerate understanding, propose changes, identify missing documentation, summarize repository evidence, and generate drafts that were then reviewed and corrected.")
    p(doc, "A typical Codex-assisted workflow followed four steps. First, the developer described the objective, such as documenting the AI body-measurement pipeline or improving a sprint section. Second, Codex inspected the repository using command-line search and file reads. Third, Codex proposed or generated a concrete artifact, such as a report section, table, diagram draft, test plan, or code change. Fourth, the result was checked by running scripts, reading diffs, verifying generated files, and comparing claims against the repository. This human-in-the-loop process is important because AI-generated content can be incomplete, outdated, or too generic if it is not grounded in real project files.")
    add_basic_table(
        doc,
        "Table 6.5: Codex usage during the GoFit project",
        ["Project phase", "How Codex was used", "Human validation"],
        [
            ["Requirements and scope analysis", "Summarized existing documents, identified actors, extracted feature lists, and transformed project notes into functional and non-functional requirements.", "Requirements were compared with project files, feature inventory, and supervisor expectations."],
            ["Architecture and design", "Helped connect mobile, admin, Supabase, Edge Functions, n8n, AI models, and external services into coherent architecture explanations and diagrams.", "Architecture claims were checked against repository folders, database migrations, and implementation notes."],
            ["Implementation support", "Suggested implementation plans, reviewed code organization, identified related files, and helped reason about bugs or integration risks.", "The developer decided what to apply, ran commands, inspected diffs, and avoided uncontrolled changes."],
            ["AI feature documentation", "Connected body-measurement code, MediaPipe module details, segmentation research, validation notes, and LLM automation workflows.", "Model names, limitations, and statuses were checked against the source files and troubleshooting documents."],
            ["Testing and QA planning", "Generated test matrices, validation checklists, and structural document audits.", "Tests and audits were executed locally when possible; limitations were documented when tools were unavailable."],
            ["PFE report writing", "Generated the report structure, expanded chapter content, produced tables, created diagram drafts, and rebuilt the DOCX artifact.", "The report was rebuilt, audited structurally, and revised based on missing project-process details."],
        ],
        [1.55, 3.2, 1.9],
    )

    add_heading(doc, "6.2.9 MCPs and Connected Tooling Used to Stay Up to Date", 3)
    p(doc, "The project also used connected tools and MCP-style integrations. MCP, or Model Context Protocol, is useful because it allows an AI assistant to interact with controlled external contexts instead of relying only on static model memory. In practice, this means the assistant can inspect local files, run commands, consult project documentation, use document-generation tools, and connect to specialized services when available. This was important for GoFit because the technology stack evolves quickly: Expo, React Native, Next.js, Supabase, AI tooling, and deployment workflows can change between academic years.")
    p(doc, "The MCP and plugin-based workflow helped keep the work closer to the current technical state. For example, official documentation and local repository files were preferred over memory when describing frameworks, Supabase behavior, OpenAI/Codex usage, and document generation. Local commands were used to inspect the real project instead of inventing features. Document tooling was used to generate and audit the Word report. Supabase and automation-related context was used to describe backend workflows, database artifacts, Edge Functions, and n8n flows more accurately.")
    add_basic_table(
        doc,
        "Table 6.6: MCP and connected tools used in the project workflow",
        ["Tool / connector type", "Purpose", "Example contribution to the report or project"],
        [
            ["Local filesystem and terminal tools", "Inspect repository files, search code, run scripts, check generated artifacts, and verify counts.", "Used to read GoFitMobile, admin-panel, database, docs, and n8n workflow files before writing report sections."],
            ["Official documentation / web lookup", "Verify current technology information when framework behavior or product details may have changed.", "Used as a principle for up-to-date descriptions of fast-moving tools such as Expo, Supabase, Next.js, OpenAI/Codex, and AI APIs."],
            ["Documents tooling", "Create, edit, audit, and regenerate DOCX deliverables.", "Used to generate the PFE report, add accessibility metadata, mark table headers, and validate DOCX structure."],
            ["Supabase-related tooling", "Reason about backend services, migrations, RLS, Edge Functions, database tables, and security boundaries.", "Supported sections on authentication, PostgreSQL, storage, realtime, RLS, service-role isolation, and backend architecture."],
            ["Browser / app inspection tools", "Check local or web interfaces when visual or runtime verification is required.", "Useful for validating frontend behavior, screenshots, documentation pages, and current technology references."],
            ["Automation tooling", "Inspect or execute n8n workflows and Docker-based automation environments when needed.", "Supported documentation of AI session preparation, admin digest automation, and service-role workflow execution."],
            ["PDF/document readers", "Use the uploaded example report for structural inspiration without copying its content.", "Helped reproduce the academic organization, front matter, chapter order, tables, figures, and report style."],
        ],
        [1.75, 2.55, 2.25],
    )
    add_report_note(
        doc,
        "MCP methodology",
        "MCPs were not treated as automatic truth sources. They were used to access controlled context: local files, official documentation, generated documents, database-related material, and automation artifacts. Final claims still required human review and consistency checks."
    )

    add_heading(doc, "6.2.10 Complete Toolchain Used to Build GoFit", 3)
    p(doc, "The following table summarizes the main tools used to create GoFit. It includes product technologies, development tools, AI assistants, automation services, and documentation tools. This list is useful because the PFE evaluates not only the final application but also the engineering environment used to produce it.")
    add_basic_table(
        doc,
        "Table 6.7: Complete development and production toolchain",
        ["Category", "Tools", "Use in the project"],
        [
            ["Mobile development", "Expo SDK, React Native, TypeScript, React Navigation, Zustand, Expo Camera, Expo Secure Store.", "Build the cross-platform mobile app, navigation flows, state management, authentication persistence, camera/photo capture, and client/coach screens."],
            ["Web administration", "Next.js App Router, TypeScript, Tailwind CSS, shadcn-style UI components.", "Build the admin panel, protected routes, dashboards, management pages, forms, analytics, and operational interfaces."],
            ["Backend and database", "Supabase Auth, PostgreSQL, RLS, Storage, Realtime, Edge Functions, SQL migrations.", "Provide authentication, persistent data, security policies, media storage, realtime communication, and privileged backend operations."],
            ["AI and computer vision", "MoveNet Lightning, TensorFlow Lite, MediaPipe Pose Landmarker, MediaPipe Image Segmenter, local Expo native module.", "Implement AI-assisted body measurement using pose detection, segmentation, feature extraction, and statistical estimation."],
            ["LLM and automation", "Groq API, n8n workflows, Supabase service-role REST calls.", "Generate adaptive workout recommendations, coach session briefings, AI session notes, and admin operations digests."],
            ["Communication and external services", "Expo Push, LiveKit, optional hosting/storage providers.", "Support notifications, video-call token flows, media/storage extension, and remote coaching features."],
            ["Development assistance", "Codex, MCP/tool connectors, terminal commands, repository search, document-generation scripts.", "Accelerate analysis, implementation support, debugging, report writing, current-doc lookup, and artifact generation."],
            ["Quality and documentation", "Git, Python docx scripts, accessibility audit, generated diagrams, project documentation files.", "Version work, generate the DOCX report, audit structure, maintain architecture/feature/sprint notes, and prepare submission artifacts."],
            ["Deployment preparation", "EAS Build, Next.js-compatible hosting, Supabase project configuration, environment variables.", "Prepare mobile builds, web hosting, backend deployment, secret separation, and production release steps."],
        ],
        [1.5, 2.4, 2.55],
    )

    add_diagram(
        doc,
        "figure_6_3_ai_tooling_workflow.png",
        "Figure 6.3: AI-assisted development and MCP workflow",
        "AI-Assisted Development Workflow",
        [
            ["Developer Goal", "Codex Assistant", "Repository + MCP Context"],
            ["File Inspection", "Current Docs / Tools", "Implementation or Report Draft"],
            ["Human Review", "Tests / Audits / Diffs", "Validated Artifact"],
        ],
    )
    add_report_note(
        doc,
        "Academic transparency",
        "The report should clearly state that AI assistants supported development and documentation, but they did not replace human implementation responsibility. All generated code, diagrams, report text, and technical claims must be checked against the repository and corrected when inaccurate."
    )


def add_n8n_automation_details(doc):
    add_heading(doc, "6.2.11 n8n Automation Layer", 3)
    p(doc, "GoFit includes an automation layer implemented with n8n workflow exports located in docs/automation/n8n/workflows. These workflows are part of the product because they perform scheduled operational work that would otherwise require manual monitoring by coaches or administrators. The n8n layer is separate from the mobile application and the admin panel: it runs in a trusted automation environment, calls Supabase through REST endpoints, and creates persistent rows in the database.")
    p(doc, "The v1 automation design uses self-hosted n8n Community Edition as the practical free option. Runtime secrets are injected through environment variables. SUPABASE_URL identifies the backend project, SUPABASE_SERVICE_ROLE_KEY allows controlled server-side access, and GROQ_API_KEY is required only by the AI Session Prep workflow. These secrets must never be committed to the repository, placed in mobile code, or exposed in client-side admin bundles.")
    p(doc, "Each workflow uses a schedule trigger followed by a tested Code node. The exported canvases also include sticky-note cards that document the schedule window, query logic, duplicate guard, notification creation, and run summary. This structure keeps the executable path small while still making the workflow understandable for maintenance and academic documentation.")
    add_basic_table(
        doc,
        "Table 6.8: n8n workflows implemented in GoFit",
        ["Workflow", "Trigger", "Data read", "Actions performed", "Guard / control"],
        [
            ["AI Session Prep v1", "Every 60 minutes", "Confirmed bookings in the next 24 hours, coach profile, client profile, recent completed sessions, private coach notes, and recent check-ins.", "Calls Groq from n8n, inserts an ai_session_notes row, and creates a coach notification with data.kind = ai_session_ready.", "Skips unexpired existing notes for the same coach/client pair and repairs missing coach notification when needed."],
            ["Booking Reminders v1", "Every 15 minutes", "Confirmed bookings starting within the next 60 minutes plus coach user mapping.", "Creates one client notification and one coach notification with data.kind = booking_starting_soon.", "Skips duplicate reminders for the same booking, recipient role, and 60-minute window."],
            ["Check-in Reminders v1", "Every 4 hours", "Enabled check_in_schedules, current UTC date/time, and check_in_responses for the day.", "Creates a client notification with data.kind = check_in_due.", "Skips schedules not due, clients who already responded, and existing reminders for the same schedule/date."],
            ["Coach Daily Digest v1", "Daily", "Coach profiles, check-in schedules and responses, upcoming confirmed bookings, active packs, booking history, and BI client-health daily rows.", "Creates a coach notification summarizing missed check-ins, low wellness signals, upcoming sessions, and inactive clients.", "Skips coaches with no useful signal and skips duplicate digest notifications for the same coach and UTC date."],
            ["Admin Ops Digest v1", "Daily", "Admin user profiles, new users, pending coaches, booking outcomes, and pack purchases from the last 24 hours.", "Creates an admin_notifications row per administrator with operational counts and a deterministic digest link.", "Skips duplicate admin digests for the same admin, title, and date link; automation_failures_count remains 0 in v1 because failures are not yet persisted."],
        ],
        [1.25, 1.05, 2.45, 2.45, 2.35],
    )
    add_basic_table(
        doc,
        "Table 6.9: n8n security, limits, and evidence requirements",
        ["Topic", "Design decision", "Reason / evidence to attach"],
        [
            ["Secret management", "SUPABASE_SERVICE_ROLE_KEY and GROQ_API_KEY stay inside the n8n runtime environment.", "Prevents privileged keys from being bundled into the mobile app, admin client, .env.example, or workflow JSON."],
            ["Service-role access", "Workflows use service-role access only from trusted automation infrastructure.", "Needed for scheduled background work while keeping normal app access protected by RLS."],
            ["AI boundary", "AI Session Prep calls Groq directly from n8n instead of calling the ai-session-notes Edge Function.", "The Edge Function is designed for authenticated coach actions; the automation flow is a scheduled server-side job."],
            ["Notification strategy", "v1 workflows create in-app notification rows rather than sending push notifications directly.", "Push delivery needs token, retry, throttling, and failure handling before being automated safely."],
            ["Client privacy", "Clients do not read ai_session_notes through normal app access.", "The table is coach-facing and protected by coach-only authenticated access rules."],
            ["Time handling", "Daily and due-date workflows use UTC in v1.", "Timezone-aware scheduling can be added after user/coach timezone preferences are finalized."],
            ["Idempotency", "Each workflow checks existing rows before inserting new notifications or notes.", "Manual verification must show that a second run does not create duplicates."],
            ["Report evidence", "Final report should attach workflow canvas screenshots, run summaries, and Supabase rows for each workflow.", "Evidence proves that workflows are implemented and not only planned."],
        ],
        [1.35, 2.65, 2.75],
    )
    add_diagram(
        doc,
        "figure_6_4_n8n_automation_layer.png",
        "Figure 6.4: n8n automation layer",
        "GoFit n8n Automation Layer",
        [
            ["n8n Schedule Trigger", "Code Node", "Supabase REST"],
            ["Groq API", "AI Session Notes", "Notifications"],
            ["Coach Digest", "Admin Digest", "Manual Verification"],
        ],
    )
    add_code_block(
        doc,
        "Mermaid n8n automation diagram draft",
        """flowchart LR
    A[n8n schedule trigger] --> B[Workflow code node]
    B --> C[Supabase REST API]
    B --> D[Groq API for AI Session Prep]
    C --> E[(PostgreSQL tables)]
    E --> F[ai_session_notes]
    E --> G[notifications]
    E --> H[admin_notifications]
    G --> I[Mobile coach/client inbox]
    H --> J[Admin panel]
    B --> K[Run summary counts]
""",
    )


def add_library_and_interaction_details(doc):
    add_heading(doc, "6.2.12 Libraries, Dependencies, and Device Feedback", 3)
    p(doc, "A complete PFE report should not describe GoFit only through screens and database tables. The implementation also depends on a large set of libraries that provide navigation, state management, forms, validation, media capture, haptic feedback, animations, charts, video calls, localization, notifications, artificial intelligence, and administration-table behavior. The following tables summarize the main libraries identified from the mobile and admin package files and explain how they support the product.")
    add_basic_table(
        doc,
        "Table 6.10: Mobile libraries and implementation roles",
        ["Library / family", "Used for", "Examples in GoFit"],
        [
            ["Expo SDK and React Native", "Cross-platform mobile runtime, native APIs, Android/iOS builds, and shared TypeScript UI.", "Main GoFitMobile application, Expo dev client, camera access, file/document picker, image picker, splash screen, status bar, and EAS build preparation."],
            ["React Navigation", "Stack, tab, and role-aware navigation between client, coach, authentication, onboarding, and detail screens.", "Auth flow, onboarding flow, client app tabs, coach app tabs, profile/settings screens, marketplace details, booking, programs, and workout session routes."],
            ["Zustand and persisted state", "Lightweight state management for reusable mobile state.", "Auth/session state, workout state, timer preferences, bookings, calendar, chat, coach/client data, health, language, marketplace, packs, programs, theme, text size, nutrition, daily coach, and wallet state."],
            ["Supabase JavaScript client", "Authenticated access to Supabase Auth, PostgreSQL, Storage, Realtime, and Edge Functions.", "User profiles, workouts, sessions, coach marketplace, bookings, packs, chat, notifications, AI session notes, barcode lookup, video-token generation, and push notification relay."],
            ["Expo Secure Store and Async Storage", "Secure session persistence and local preference persistence.", "Auth tokens, user session continuity, persisted preferences, and local app state that must survive app restarts."],
            ["Expo Notifications and push helpers", "Local notification scheduling, notification permissions, push-token registration, and notification inbox integration.", "Workout reminders, weekly progress reports, achievement/test notifications, push-token registration, notification preferences, and transactional notification rows."],
            ["Expo Haptics", "Physical feedback for taps, selections, timer events, success states, warning states, and error states.", "Auth buttons, onboarding controls, home cards, rest timer, workout session actions, body measurement, nutrition logging, booking, chat, coach programs, coach packs, settings, and profile actions."],
            ["Expo Camera, Image Picker, Image Manipulator, Image, File System", "Photo capture, photo selection, image processing, and file access.", "Progress photos, body-measurement capture, coach CV/certification upload flows, and image-based AI preparation."],
            ["Media and native AI libraries", "On-device or native computer-vision support.", "react-native-fast-tflite, MediaPipe native module, MoveNet/TFLite fallback, segmentation analysis, pose landmarks, and measurement estimator support."],
            ["LiveKit and WebRTC libraries", "Realtime video-call support between coaches and clients.", "LiveKit React Native, Expo plugin, WebRTC bridge, video token function, and video-call screen."],
            ["Forms, validation, localization, and UI helpers", "Typed input handling, validation, translation, icons, charts, dates, gestures, animations, and SVG rendering.", "React Hook Form, Zod, i18next/react-i18next, lucide-react-native, react-native-chart-kit, date-fns, gesture handler, reanimated, vector icons, SVG, safe-area, screens, and worklets."],
        ],
        [1.65, 2.25, 2.8],
    )
    add_basic_table(
        doc,
        "Table 6.11: Admin panel libraries and implementation roles",
        ["Library / family", "Used for", "Examples in GoFit admin"],
        [
            ["Next.js and React", "Administration web application, App Router pages, route handlers, server-side access, and protected dashboard views.", "Login, dashboard, users, coaches, exercises, workouts, transactions, notifications, settings, activity logs, BI routes, and health endpoints."],
            ["Supabase SSR and Supabase JS", "Server/client Supabase access with session-aware and admin/server-side operations.", "Protected admin routes, admin data loading, API route handlers, role checks, database access, and privileged management actions."],
            ["Tailwind CSS and shadcn/Radix UI patterns", "Consistent admin interface components and accessible controls.", "Dialogs, dropdowns, checkboxes, popovers, progress, scroll areas, select menus, switches, toasts, alert dialogs, and management forms."],
            ["React Hook Form, Zod, and Hookform resolvers", "Form state, validation, typed payloads, and error messages.", "Admin create/edit forms, user/coach management validation, exercise/workout forms, settings, and API request validation."],
            ["TanStack React Table", "Dense sortable/filterable tabular administration views.", "Users, coaches, exercises, workouts, transactions, notifications, and BI/listing screens."],
            ["Recharts, date-fns, Lucide React", "Charts, date formatting, and consistent admin iconography.", "Dashboard metrics, BI snapshots, date filters, operations digests, navigation icons, action buttons, and status indicators."],
            ["Sonner, cmdk, next-themes, class utilities", "Toasts, command/search interactions, theme handling, and conditional styling.", "Admin feedback messages, search/command patterns, light/dark support, and reusable component variants."],
            ["AWS S3 client", "Optional object-storage integration where hosting/storage extension is required.", "Admin or backend-facing media/file workflows when external storage is configured."],
        ],
        [1.65, 2.3, 2.75],
    )
    add_basic_table(
        doc,
        "Table 6.12: Haptics and tactile feedback coverage",
        ["Feature area", "Feedback type", "Implementation detail"],
        [
            ["Rest timer and workout session", "Light, medium, heavy impact, success feedback, and countdown feedback.", "The timer preference includes haptics_enabled; useRestTimer, EnhancedRestTimer, TimerModal, and WorkoutSessionScreen call haptic helpers for timer actions, countdown thresholds, set completion, and workout milestones."],
            ["Authentication and onboarding", "Light impact and selection feedback.", "Welcome buttons, social buttons, checkboxes, onboarding navigation, weight selector, and personal details flows use Expo Haptics to make taps and selections feel responsive."],
            ["Home dashboard and quick actions", "Light or medium impact, plus success feedback where relevant.", "Home cards for articles, health, check-ins, nutrition, daily coach loop, recommended workouts, quick actions, programs, top workouts, and top trainers use haptics for card presses and action confirmations."],
            ["Progress and body measurement", "Light/medium impact, warning feedback, and success feedback.", "Progress photos, photo detail, workout statistics, body-measurement capture, AI scan warnings, and successful measurement saves use tactile confirmation."],
            ["Nutrition and check-ins", "Light/medium impact, selection feedback, and success notification feedback.", "Nutrition dashboard, food adding, nutrition goals, barcode/food actions, and check-in submission use haptics to distinguish selection, confirmation, and completion events."],
            ["Profile and settings", "Light impact, selection feedback, success feedback, and error feedback.", "Text size, units, language, account information, height/weight editing, habits, notification settings, bookings/reviews, health sync, and preferences use haptics to support settings changes and save/error states."],
            ["Coach marketplace and booking", "Light/medium impact and warning feedback.", "Marketplace browsing, coach detail, booking slot selection, client views, calendar interactions, availability, session packs, and booking decisions provide tactile response on important actions."],
            ["Coach app communication and management", "Light/medium impact and success/error feedback.", "Coach dashboard navigation, client list/detail, notes, check-ins, calendar, profile/settings, chat, conversations, programs, program builder, packs, certifications, and CV upload use haptics for action acknowledgement."],
            ["Design principle", "Feedback is assistive, not decorative.", "Light feedback is used for ordinary taps/selections, medium/heavy feedback for stronger actions or timer thresholds, and notification feedback for success, warning, or error states. Haptics should remain configurable where the feature exposes preferences."],
        ],
        [1.55, 1.7, 3.45],
    )
    add_heading(doc, "6.2.13 Exact Runtime Dependency Inventory", 3)
    p(doc, "The previous tables explain the main technology families. For traceability, the following tables list the exact runtime dependencies declared by the mobile application and the administration panel package files. This makes the report auditable against the project source instead of depending only on a summarized technology paragraph.")
    add_basic_table(
        doc,
        "Table 6.13: Mobile runtime dependency inventory",
        ["Package", "Version", "Role in GoFit"],
        dependency_rows("GoFitMobile/package.json"),
        [2.35, 1.0, 3.25],
    )
    add_basic_table(
        doc,
        "Table 6.14: Admin panel runtime dependency inventory",
        ["Package", "Version", "Role in GoFit admin"],
        dependency_rows("admin-panel/package.json"),
        [2.35, 1.0, 3.25],
    )


def add_chapter6_platform_architecture_details(doc):
    add_heading(doc, "6.2.5 GoFit Mobile Application Architecture", 3)
    p(doc, "The mobile application is structured as a layered client application. Screens and reusable components are responsible for presentation, navigation manages access to feature areas, Zustand stores hold cross-screen state, service modules centralize data access, and infrastructure clients communicate with Supabase or local native modules. This separation reduces coupling between user interface code and persistence logic.")
    add_basic_table(
        doc,
        "Table 6.3: Mobile architecture responsibilities",
        ["Layer", "Responsibility", "Examples in GoFit"],
        [
            ["Presentation", "Display screens, forms, lists, charts, timers, and feedback states.", "Authentication screens, workout screens, progress screens, coach marketplace screens."],
            ["Navigation", "Route users according to authentication, onboarding, role, and feature context.", "Auth navigator, onboarding flow, client app area, coach app area, detail screens."],
            ["State management", "Store reusable client state and synchronize UI updates.", "Auth store, workout store, profile state, timer preferences, progress state."],
            ["Domain services", "Encapsulate API calls, persistence logic, and feature-specific operations.", "Workout service, progress service, body measurement service, coach/booking services."],
            ["Native/AI modules", "Execute device-specific or performance-sensitive functionality.", "Camera capture, MediaPipe native module, MoveNet fallback, segmentation analysis."],
        ],
        [1.35, 2.5, 2.65],
    )
    add_heading(doc, "6.2.6 GoFit Administration Panel Architecture", 3)
    p(doc, "The administration panel provides a controlled back-office interface. It is designed separately from the mobile application because administrators need operational views, management tables, filters, analytics, settings, and privileged actions that should not be exposed in the client application. The panel uses protected routing and server-side Supabase access to keep administrative operations behind explicit authorization checks.")
    p(doc, "The admin panel should be validated with at least two account types: an administrator account that can access management screens and a normal account that must be blocked or redirected. This distinction is important because the panel can manage sensitive platform data such as coach validation status, user records, exercise content, transactions, and notifications.")
    add_heading(doc, "6.2.6.1 Admin BI Data Architecture", 3)
    p(doc, "The BI architecture follows a protected data-to-dashboard pipeline. PostgreSQL views such as bi_finance_daily, bi_user_lifecycle_daily, bi_coach_ops_daily, and bi_client_health_daily prepare canonical daily metrics. TypeScript service modules in the admin panel read those views and normalize the results into dashboard contracts. React components then display summary cards, trend panels, risk queues, cohort tables, finance details, and threshold alerts. API routes under /api/bi provide exports, saved views, snapshots, and scheduled digest execution.")
    p(doc, "The BI layer is deliberately separated from normal CRUD pages. CRUD pages manage individual records, while BI pages aggregate and interpret platform behavior. This separation reduces duplication and makes it clearer which parts of the panel are operational management and which parts are decision-support analytics. The v1 BI scope is also explicit: finance is gross-pack-sales oriented, lifecycle and cohort metrics depend on available history, risk queues are operational signals, and all BI access must remain admin-protected.")
    add_heading(doc, "6.2.7 Supabase Backend, Data Security, and Realtime Services", 3)
    p(doc, "The Supabase backend centralizes authentication, PostgreSQL data, storage, realtime communication, and Edge Functions. Row-level security is a central design decision: each user should access only their own client data, each coach should access only authorized client data, and administrators should access management data through controlled server-side contexts. Service-role keys must remain in trusted backend, admin, or automation environments and must never be bundled into the mobile application.")
    add_basic_table(
        doc,
        "Table 6.4: Backend and security components",
        ["Component", "Role", "Security / validation concern"],
        [
            ["Supabase Auth", "Manages identity, sessions, password reset, and protected access.", "Check session handling, role mapping, and logout/session expiration behavior."],
            ["PostgreSQL database", "Stores users, workouts, sessions, measurements, coaches, bookings, packs, messages, programs, and admin data.", "Validate foreign keys, migrations, indexes, and data consistency."],
            ["RLS policies", "Restrict row access according to user, coach, and administrator roles.", "Test user-scoped, coach-scoped, and admin-scoped queries with different accounts."],
            ["Storage", "Stores user-uploaded files such as coach documents, media, or progress-related assets when enabled.", "Check bucket policies, private/public access, and file ownership."],
            ["Realtime/chat", "Supports live communication and status updates.", "Validate conversation membership, message visibility, and realtime channel permissions."],
            ["Edge Functions", "Execute privileged operations such as push notification relay, video-token generation, and AI/backend workflows.", "Keep secrets server-side and verify authorization inside each function."],
            ["n8n automations", "Run scheduled or event-driven workflows such as AI session preparation and admin operations digest.", "Use service-role access only in controlled environments and log workflow results."],
        ],
        [1.35, 2.35, 2.8],
    )


def sprint_chapter(doc, n, title, goal, backlog, diagrams, realisation, tests):
    add_heading(doc, f"{n} {title}", 1)
    p(doc, goal)
    p(doc, f"This chapter documents {title.lower()} as an engineering increment rather than as a simple list of implemented screens. It presents the sprint objective, selected backlog, structural model, refined use cases, sequence flow, implementation choices, validation matrix, and burn-down interpretation.")
    add_heading(doc, f"{n}.1 Needs Assessment", 2)
    p(doc, "The sprint needs were selected from the product backlog according to dependency order, user value, and technical feasibility.")
    add_report_note(doc, "Sprint objective", goal)
    add_heading(doc, f"{n}.1.1 Sprint Backlog", 3)
    add_basic_table(doc, f"Table {n}.1: {title} backlog", ["ID", "Task", "Priority", "Result"], backlog, [0.7, 3.4, 0.8, 1.6])
    add_heading(doc, f"{n}.1.2 Sprint Class Diagram", 3)
    p(doc, diagrams["class"])
    add_diagram(
        doc,
        f"figure_{n}_1_class.png",
        f"Figure {n}.1: Class diagram of {title}",
        f"{title} - Main Structural Classes",
        diagrams["class_rows"],
    )
    add_code_block(doc, f"Mermaid class diagram draft for {title}", mermaid_code_for("class", n))
    add_heading(doc, f"{n}.2 Functional Specification of Requirements", 2)
    p(doc, "This section specifies the sprint use cases and explains how the selected actors interact with the GoFit system.")
    add_heading(doc, f"{n}.2.1 Refined Use Case Diagram of {title}", 3)
    p(doc, diagrams["usecase"])
    add_diagram(
        doc,
        f"figure_{n}_2_usecase.png",
        f"Figure {n}.2: Refined use case diagram of {title}",
        f"{title} - Refined Use Cases",
        diagrams["usecase_rows"],
    )
    add_code_block(doc, f"Mermaid refined use case diagram draft for {title}", mermaid_code_for("use_case", n))
    add_heading(doc, f"{n}.2.2 Use Case Analysis", 3)
    add_detailed_use_cases(doc, n)
    p(doc, diagrams["sequence"])
    add_diagram(
        doc,
        f"figure_{n}_3_sequence.png",
        f"Figure {n}.3: Sequence diagram of key use case",
        f"{title} - Sequence Flow",
        diagrams["sequence_rows"],
    )
    add_code_block(doc, f"Mermaid sequence diagram draft for {title}", mermaid_code_for("sequence", n))
    add_heading(doc, f"{n}.3 Realisation", 2)
    for item in realisation:
        p(doc, item)
    if n == 3:
        add_todo(doc, "Insert screenshot of mobile authentication screen")
        add_todo(doc, "Insert screenshot of workout session interface")
    elif n == 4:
        add_todo(doc, "Insert screenshot of progress dashboard")
        add_todo(doc, "Insert screenshot of body measurement review screen")
    elif n == 5:
        add_todo(doc, "Insert screenshot of coach marketplace screen")
        add_todo(doc, "Insert screenshot of admin dashboard")
    add_diagram(
        doc,
        f"figure_{n}_4_realisation_summary.png",
        f"Figure {n}.4: Realisation summary of {title}",
        f"{title} - Implemented Views",
        [["Mobile Screens", "Services", "Supabase Data"], ["Stores", "Validation", "User Interface"], ["TODO Screenshots", "Feature Demo", "Test Evidence"]],
    )
    if n == 3:
        add_sprint1_implementation_details(doc)
    if n == 4:
        add_sprint2_progress_details(doc)
        add_ai_computer_vision_details(doc)
    if n == 5:
        add_sprint3_platform_details(doc)
    if n == 3:
        test_table_number = "3.6"
        burndown_figure_number = "3.5"
    elif n == 4:
        test_table_number = "4.9"
        burndown_figure_number = "4.9"
    elif n == 5:
        test_table_number = "5.7"
        burndown_figure_number = "5.7"
    else:
        test_table_number = f"{n}.3"
        burndown_figure_number = f"{n}.5"
    add_heading(doc, f"{n}.4 {'Sprint Tests' if n == 3 else 'Tests'}", 2)
    p(doc, "The following validation matrix avoids declaring success without evidence. When a screenshot, terminal output, or exported test report is missing, the evidence cell explicitly marks the remaining artifact to attach before final submission.")
    add_professional_tests(doc, n, test_table_number)
    add_heading(doc, f"{n}.5 Sprint Burn Down Chart", 2)
    p(doc, "The burn-down chart is represented as a sprint-management artifact. In the final report, this figure should be replaced by the exported chart from the project management board if such evidence is available. Until then, it explains the expected evolution from selected backlog to validated increment.")
    add_diagram(
        doc,
        f"figure_{n}_5_burndown.png",
        f"Figure {burndown_figure_number}: Burn down chart for {title}",
        f"{title} - Burn Down Summary",
        [["Sprint Start", "Backlog Selected", "Tasks In Progress"], ["Mid Sprint", "Validation", "Remaining Work Reduced"], ["Sprint End", "Tests", "Increment Delivered"]],
    )
    add_heading(doc, "Conclusion", 2)
    p(doc, sprint_conclusion(n))
    doc.add_page_break()


def chapters3_to5(doc):
    sprint_chapter(
        doc,
        3,
        "Sprint 1",
        "Sprint 1 focused on the GoFit foundation: project setup, authentication, onboarding, profile management, workout planning, workout session execution, and the enhanced rest timer.",
        [
            ["S1-01", "Set up Expo/React Native, Next.js, Supabase, repository structure, and environment variables.", "High", "Completed foundation"],
            ["S1-02", "Implement authentication, onboarding, profile management, and protected navigation.", "High", "Mobile auth flow ready"],
            ["S1-03", "Implement workout templates, exercise configuration, workout sessions, and rest timer.", "High", "Workout core functional"],
            ["S1-04", "Persist workout data through Supabase services and Zustand stores.", "High", "Mobile/backend connection validated"],
        ],
        {
            "class": "Include classes/entities such as UserProfile, Workout, WorkoutExercise, Exercise, WorkoutSession, TimerPreferences, AuthService, WorkoutService, and WorkoutStore.",
            "class_rows": [["UserProfile", "AuthService", "AuthStore"], ["Workout", "WorkoutExercise", "Exercise"], ["WorkoutSession", "TimerPreferences", "WorkoutService"]],
            "usecase": "Actors: Mobile user and Supabase. Use cases: sign up, log in, complete onboarding, manage profile, browse workouts, create workout, start session, track sets, use rest timer.",
            "usecase_rows": [["Mobile User", "Sign Up / Log In", "Complete Onboarding"], ["Manage Profile", "Browse Workouts", "Create Workout"], ["Start Session", "Track Sets", "Use Rest Timer"], ["Supabase", "Persist Profile", "Persist Session"]],
            "sequence": "Show login or workout-session flow: user action -> mobile screen -> Zustand store -> service layer -> Supabase -> UI update.",
            "sequence_rows": [["User Action", "Mobile Screen", "Zustand Store"], ["Service Layer", "Supabase Query", "Database Row"], ["Response", "Store Update", "UI Refresh"]],
            "usecase_table": [
                ["Use case", "Start and complete a workout session"],
                ["Primary actor", "Mobile user"],
                ["Preconditions", "The user is authenticated and a workout template exists."],
                ["Main scenario", "The user selects a workout, starts a session, completes sets, uses the rest timer, and saves the workout summary."],
                ["Alternative scenarios", "The user pauses, exits, or edits set data before completion."],
                ["Postconditions", "A workout_sessions row stores duration, exercises_completed, calories, and notes."],
            ],
        },
        [
            "The mobile foundation was implemented with Expo SDK 54 and React Native. Navigation separates authentication, onboarding, client app, and coach app flows.",
            "Authentication uses Supabase Auth. The session is mirrored in the auth store, and protected navigation redirects users according to authentication and onboarding state.",
            "The workout model separates templates from execution logs. Workouts and exercises are configured before execution, while workout_sessions store actual completed activity.",
            "The enhanced rest timer includes animated display, audio and haptic feedback, warning intervals, auto-advance configuration, and persisted preferences.",
        ],
        [
            ["T1", "Authentication", "Valid credentials open the correct navigator.", "Legacy source row not rendered"],
            ["T2", "Workout creation", "A custom workout is saved with ordered exercises.", "Legacy source row not rendered"],
            ["T3", "Workout session", "Sets, reps, weights, rest time, and summary are stored.", "Legacy source row not rendered"],
            ["T4", "Rest timer", "Timer starts, pauses, resumes, warns, and completes.", "Legacy source row not rendered"],
        ],
    )

    sprint_chapter(
        doc,
        4,
        "Sprint 2",
        "Sprint 2 extended the mobile experience with exercise library improvements, calendar planning, progress tracking, notifications, nutrition-related tracking, and body measurement groundwork.",
        [
            ["S2-01", "Implement exercise library search, detail views, images, and filtering.", "High", "Library usable"],
            ["S2-02", "Implement calendar, workout history, and progress statistics.", "High", "Tracking views ready"],
            ["S2-03", "Add notification settings and reminders.", "Medium", "Reminder flow integrated"],
            ["S2-04", "Prepare body measurement and progress-photo features.", "Medium", "AI measurement groundwork documented"],
            ["S2-05", "Implement nutrition or saved meal flows where applicable.", "Medium", "Nutrition module draft integrated"],
        ],
        {
            "class": "Include Exercise, Workout, WorkoutSession, BodyMeasurement, ProgressPhoto, NotificationPreference, NutritionEntry, SavedMeal, and related services/stores.",
            "class_rows": [["Exercise", "Workout", "WorkoutSession"], ["BodyMeasurement", "ProgressPhoto", "ProgressService"], ["NotificationPreference", "NutritionEntry", "SavedMeal"]],
            "usecase": "Actors: Mobile user and Supabase. Use cases: search exercise, view exercise detail, plan workout on calendar, view progress, log measurements, configure notifications, track nutrition.",
            "usecase_rows": [["Mobile User", "Search Exercises", "View Details"], ["Plan Calendar", "View Progress", "Log Measurements"], ["Configure Reminders", "Track Nutrition", "View History"], ["Supabase", "Store Metrics", "Return Charts Data"]],
            "sequence": "Show progress tracking flow: user opens progress screen -> store requests stats -> service queries workout_sessions/body_measurements -> chart renders data.",
            "sequence_rows": [["Open Progress", "Progress Store", "Progress Service"], ["Query Sessions", "Query Measurements", "Build Dataset"], ["Return Data", "Render Charts", "Show Empty State if Needed"]],
            "usecase_table": [
                ["Use case", "View progress dashboard"],
                ["Primary actor", "Mobile user"],
                ["Preconditions", "The user is authenticated and has workout, measurement, or nutrition history."],
                ["Main scenario", "The user opens progress screens, selects a metric, and views charts or history."],
                ["Alternative scenarios", "If no data exists, the system displays an empty state and invites the user to log data."],
                ["Postconditions", "The user obtains a clear view of fitness evolution over time."],
            ],
        },
        [
            "The exercise library provides access to workouts, exercises, detail screens, and workout builder flows. Exercise metadata is loaded through service layers and cached by stores.",
            "Progress screens present workout statistics, body measurements, consistency, and record details. These views help users interpret their evolution rather than only log isolated sessions.",
            "Notification features support reminders and inbox-style communication. Push-token registration and notification preferences are part of the infrastructure.",
            "The body measurement work explores AI-supported progress tracking. It is positioned as a fitness-progress feature, not a medical or tailoring-grade measurement system.",
        ],
        [
            ["T1", "Exercise library", "Search and detail screens return expected exercise data.", "Legacy source row not rendered"],
            ["T2", "Progress charts", "Charts handle populated and empty states.", "Legacy source row not rendered"],
            ["T3", "Notifications", "Permissions and reminders behave correctly on device.", "Manual device test required"],
            ["T4", "Nutrition/progress", "Entries persist and appear in history.", "Legacy source row not rendered"],
        ],
    )

    sprint_chapter(
        doc,
        5,
        "Sprint 3",
        "Sprint 3 focused on coach marketplace features and administration. It added coach onboarding, coach/client flows, bookings, session packs, chat, programs, and a web admin panel for operational management.",
        [
            ["S3-01", "Implement coach authentication, onboarding, CV/certification upload, and pending review flow.", "High", "Coach onboarding ready"],
            ["S3-02", "Implement marketplace, coach detail, session booking, session packs, chat, and programs.", "High", "Marketplace flows implemented"],
            ["S3-03", "Implement coach dashboard, clients, calendar, profile, wallet, and settings.", "Medium", "Coach app area structured"],
            ["S3-04", "Implement admin panel dashboard, users, coaches, exercises, workouts, transactions, notifications, settings, and analytics.", "High", "Admin platform operational"],
        ],
        {
            "class": "Include CoachProfile, Client, SessionPack, PurchasedPack, Booking, Conversation, Message, Program, AdminUser, Exercise, Workout, Transaction, Notification, and AuditLog.",
            "class_rows": [["CoachProfile", "Client", "Program"], ["SessionPack", "PurchasedPack", "Booking"], ["Conversation", "Message", "Notification"], ["AdminUser", "Exercise", "Workout"], ["Transaction", "AuditLog", "Settings"]],
            "usecase": "Actors: Mobile user, coach, administrator, Supabase. Use cases: browse coaches, book session, buy pack, chat, create program, manage clients, validate coach, manage content.",
            "usecase_rows": [["Client", "Browse Coaches", "Book Session"], ["Coach", "Manage Clients", "Create Program"], ["Admin", "Validate Coach", "Manage Content"], ["Supabase", "Bookings / Messages", "Audit / Analytics"]],
            "sequence": "Show booking or chat flow: client selects coach -> creates booking/conversation -> Supabase persists data -> coach dashboard updates.",
            "sequence_rows": [["Client Selects Coach", "Marketplace Screen", "Booking Service"], ["Supabase Insert", "Booking / Conversation", "Realtime Update"], ["Coach Dashboard", "Client Confirmation", "Notification"]],
            "usecase_table": [
                ["Use case", "Book a coaching session"],
                ["Primary actor", "Mobile user"],
                ["Preconditions", "The user is authenticated, a coach profile is available, and a session slot or pack exists."],
                ["Main scenario", "The user opens marketplace, selects a coach, chooses a slot or pack, confirms booking, and receives confirmation."],
                ["Alternative scenarios", "The slot is unavailable, payment/session pack is missing, or the booking is cancelled."],
                ["Postconditions", "The booking is stored and visible to both client and coach."],
            ],
        },
        [
            "Coach-related screens include coach authentication, onboarding, profile preview, pending validation, dashboard, clients list, client detail, client progress, notes, programs, calendar, chat, wallet, and settings.",
            "Marketplace flows allow users to browse coaches, view coach profiles, book sessions, manage packs, access programs, and communicate through chat.",
            "The admin panel uses Next.js App Router, server-side data access, protected middleware, and Supabase admin clients. It includes dashboard stats, user management, coach management, exercise/workout CRUD foundations, transactions, notifications, activity logs, settings, search, import/export, and analytics components.",
            "Security is handled through authentication, admin role checks, service-role usage on server-side routes where intended, and Supabase RLS for user-scoped data.",
        ],
        [
            ["T1", "Coach onboarding", "Coach profile data and documents can be submitted and reviewed.", "Legacy source row not rendered"],
            ["T2", "Marketplace", "Clients can browse coaches and navigate to booking/detail flows.", "Legacy source row not rendered"],
            ["T3", "Chat/bookings/packs", "Conversation and booking records persist correctly.", "Legacy source row not rendered"],
            ["T4", "Admin panel", "Protected admin pages load data and restrict non-admin users.", "Legacy source row not rendered"],
        ],
    )


def chapter6(doc):
    add_heading(doc, "6 Deployment and Closing Phase", 1)
    p(doc, "This chapter presents the working environment, architecture, technology choices, deployment strategy, and final validation of GoFit.")
    add_heading(doc, "6.1 Working Environment", 2)
    add_heading(doc, "6.1.1 Material Environment", 3)
    add_basic_table(
        doc,
        "Table 6.1: Hardware/material environment",
        ["Resource", "Specification", "Purpose"],
        [
            ["Development machine", "Windows 11 workstation with modern multi-core CPU and sufficient RAM", "Mobile, web, backend, and documentation work"],
            ["Android emulator / device", "Model and version to confirm", "Mobile testing"],
            ["iOS simulator / device", "Model and version to confirm", "Mobile testing when available"],
            ["Cloud services", "Supabase, optional Vercel/EAS/Expo services", "Backend, admin deployment, and mobile builds"],
        ],
        [1.7, 2.6, 2.2],
    )
    add_heading(doc, "6.1.2 Software Environment", 3)
    add_basic_table(
        doc,
        "Table 6.2: Software environment",
        ["Software / Tool", "Version / Stack", "Use"],
        [
            ["Expo / React Native", "Expo SDK 54, React Native 0.81", "Cross-platform mobile app"],
            ["TypeScript", "Project TypeScript configuration", "Typed mobile and web development"],
            ["Next.js", "Next.js admin panel", "Web administration interface"],
            ["Supabase", "PostgreSQL, Auth, Storage, Realtime, Edge Functions", "Backend services and security"],
            ["Zustand", "Persisted stores", "Mobile state management"],
            ["React Navigation", "Stack and tab navigation", "Mobile routing"],
            ["Tailwind CSS / shadcn-style UI", "Admin UI stack", "Web interface styling"],
            ["Git", "Repository versioning", "Source control and collaboration"],
        ],
        [2.1, 2.0, 2.4],
    )
    add_heading(doc, "6.2 Architecture and Technological Choices", 2)
    add_heading(doc, "6.2.1 Physical Architecture Diagram", 3)
    add_diagram(
        doc,
        "figure_6_1_physical_architecture.png",
        "Figure 6.1: Physical architecture diagram",
        "GoFit Physical Architecture",
        [
            ["iOS / Android Device", "GoFit Mobile App", "Supabase APIs"],
            ["Desktop Browser", "Next.js Admin Panel", "Supabase Admin Access"],
            ["Supabase Cloud", "PostgreSQL / Auth / Storage", "Realtime / Edge Functions"],
            ["External Services", "Push / Video / Payment", "Users and Coaches"],
        ],
    )
    add_code_block(doc, "Mermaid physical architecture diagram draft", mermaid_code_for("architecture"))
    add_heading(doc, "6.2.2 Technologies", 3)
    p(doc, "Expo and React Native were selected to target Android and iOS with a shared TypeScript codebase. Supabase was selected because it provides authentication, PostgreSQL, storage, realtime features, and security policies without requiring a custom backend server for every feature. Next.js was selected for the administration panel because it supports server-side rendering, route handlers, and secure server-side access to privileged Supabase operations.")
    add_heading(doc, "6.2.3 Architecture Pattern", 3)
    p(doc, "The mobile application follows a layered structure: presentation screens and reusable components, navigation, Zustand stores, domain services, and infrastructure clients such as Supabase. The admin panel follows a web application structure based on Next.js routes, server/client components, API route handlers, and Supabase admin utilities. The backend follows a database-centered BaaS architecture with migrations, row-level security, views, functions, storage, and edge functions.")
    add_heading(doc, "6.2.4 Logical Architecture Diagram", 3)
    add_diagram(
        doc,
        "figure_6_2_logical_architecture.png",
        "Figure 6.2: Logical architecture diagram",
        "GoFit Logical Architecture",
        [
            ["Mobile Screens", "Navigation", "Zustand Stores"],
            ["Mobile Services", "Supabase Client", "RLS Policies"],
            ["Admin Pages", "Route Handlers", "Admin Supabase Client"],
            ["PostgreSQL", "Storage", "Realtime / Edge Functions"],
        ],
    )
    add_code_block(doc, "Mermaid logical architecture diagram draft", mermaid_code_for("architecture"))
    add_chapter6_platform_architecture_details(doc)
    add_ai_development_tooling_details(doc)
    add_n8n_automation_details(doc)
    add_library_and_interaction_details(doc)
    add_heading(doc, "6.3 Deployment", 2)
    add_heading(doc, "6.3.1 Deployment Diagram", 3)
    add_diagram(
        doc,
        "figure_6_3_deployment.png",
        "Figure 6.5: Deployment diagram",
        "GoFit Deployment View",
        [
            ["Source Code", "EAS Build", "Android / iOS App"],
            ["Source Code", "Next.js Hosting", "Admin Web App"],
            ["Supabase Project", "Database Migrations", "Auth / Storage / Edge"],
            ["Environment Variables", "Monitoring", "Production Users"],
        ],
    )
    add_code_block(doc, "Mermaid deployment diagram draft", mermaid_code_for("deployment"))
    add_heading(doc, "6.3.2 CI/CD or Hosting", 3)
    p(doc, "The mobile application can be built using Expo/EAS for Android and iOS. The web administration panel can be hosted on a Next.js-compatible platform such as Vercel. Supabase hosts the database, authentication, storage, realtime features, and edge functions. The n8n automation layer can be hosted as a self-hosted container with a persistent volume and environment-injected secrets. Environment variables must be separated by environment, and service-role keys must remain server-side only.")
    add_heading(doc, "6.3.3 Test Results", 3)
    add_basic_table(
        doc,
        "Table 6.15: Final validation matrix",
        ["Test ID", "Feature", "Scenario", "Preconditions", "Steps", "Expected", "Actual", "Status", "Evidence"],
        [
            ["F-T01", "End-to-end mobile flow", "Client signs in and completes core workout flow", "Test user and workout data available", "Sign in; open workout; complete session; view history", "Session appears in history and progress data", "No attached execution evidence yet", "To be validated", "TODO: attach screenshots or device recording"],
            ["F-T02", "Database/RLS", "User cannot access another user's data", "Two test accounts and RLS policies enabled", "Query scoped data from each account", "Only authorized rows are returned", "No attached execution evidence yet", "To be validated", "TODO: attach SQL or Supabase policy test output"],
            ["F-T03", "Admin panel", "Non-admin is blocked and admin can manage data", "Admin and non-admin accounts exist", "Open admin routes with both accounts", "Admin allowed; non-admin denied", "No attached execution evidence yet", "To be validated", "TODO: attach browser screenshots"],
            ["F-T04", "AI measurement", "Photo analysis returns draft values or warnings", "Android dev build and test captures available", "Capture front/side photos; run analysis; review result", "Draft values are editable and warnings are shown when quality is weak", "Debug observations documented; final evidence missing", "Partially validated", "TODO: attach device screenshot and debug output"],
            ["F-T05", "n8n automations", "All five workflows run manually without duplicates", "n8n container configured with Supabase and Groq environment variables", "Run AI Session Prep, Booking Reminders, Check-in Reminders, Coach Daily Digest, and Admin Ops Digest twice", "First run creates expected rows; second run reports skips or no duplicates", "Workflow exports documented; final run screenshots missing", "To be validated", "TODO: attach n8n run summaries and Supabase row evidence"],
            ["F-T06", "Deployment", "Mobile/admin/backend/automation deployment readiness", "Production env variables and build config prepared", "Run EAS build, admin deployment, Supabase checks, and n8n workflow import", "Builds and hosted services are reachable; workflows are imported and scheduled", "No attached execution evidence yet", "Manual validation required", "TODO: attach build/deployment output"],
        ],
        [0.48, 0.72, 0.88, 0.85, 1.05, 0.95, 0.9, 0.76, 0.91],
    )
    doc.add_page_break()


def add_complete_feature_inventory_appendix(doc):
    features = parse_feature_inventory()
    add_heading(doc, "Appendix C: Complete Source Feature Inventory", 2)
    p(doc, "This appendix is generated from the project feature inventory and is used as a coverage ledger for the report. It does not replace the narrative chapters; instead, it proves that the report accounts for the implemented mobile client, mobile coach, administration, backend, AI, automation, and infrastructure work found in the repository.")
    p(doc, "The source inventory declares 249 implemented feature blocks. The generated appendix keeps one additional shared mobile application-shell row because it documents bootstrapping, global providers, notifications, deep links, and role-aware routing used by both client and coach areas. Therefore, the appendix contains 250 coverage rows: 249 declared feature blocks plus 1 shared foundation row.")

    groups = [
        ("Mobile client", "Mobile client application features, including the shared mobile shell"),
        ("Mobile coach", "Mobile coach application features"),
        ("Admin panel", "Administration panel features"),
        ("Backend and infrastructure", "Backend, database, Edge Function, and infrastructure features"),
    ]
    summary_rows = []
    for group_name, label in groups:
        count = len([row for row in features if feature_group(row["area"]) == group_name])
        summary_rows.append([label, str(count), "Parsed from FEATURES.md feature blocks."])
    summary_rows.append(["Declared source feature blocks", "249", "Published in FEATURES.md summary table."])
    summary_rows.append(["Coverage rows retained in appendix", str(len(features)), "Includes the shared mobile application-shell row."])

    add_basic_table(
        doc,
        "Table A.3: Feature inventory summary",
        ["Area", "Feature blocks", "Source"],
        summary_rows,
        [2.4, 1.1, 3.0],
    )

    table_specs = [
        ("Table A.4: Mobile client feature inventory", "Mobile client"),
        ("Table A.5: Mobile coach feature inventory", "Mobile coach"),
        ("Table A.6: Admin panel feature inventory", "Admin panel"),
        ("Table A.7: Backend and infrastructure feature inventory", "Backend and infrastructure"),
    ]
    for caption, group_name in table_specs:
        rows = feature_inventory_table_rows(features, group_name)
        add_basic_table(
            doc,
            caption,
            ["No.", "Feature", "Implementation evidence", "Libraries / backend", "Purpose"],
            rows or [["0", "No parsed feature block", "N/A", "N/A", "Check FEATURES.md parsing."]],
            [0.45, 1.35, 1.75, 1.75, 1.9],
        )


def final_pages(doc):
    add_heading(doc, "General Conclusion", 1)
    p(doc, "The GoFit graduation internship project produced a substantial software foundation for a modern fitness platform. It combines a mobile application for clients and coaches, a web administration panel, and a Supabase backend. The repository contains implementations or foundations for authentication, onboarding, profile management, workout planning, exercise library, active workout sessions, rest timer, progress tracking, nutrition-related flows, coach marketplace features, bookings, chat, programs, notifications, analytics, administration, and deployment preparation.")
    p(doc, "Beyond the visible interfaces, the project demonstrates several engineering competencies. The mobile application required cross-platform architecture, state management, navigation design, camera and native-module integration, and device-oriented testing. The backend required database normalization, migrations, row-level security, storage, realtime communication, Edge Functions, and service-role separation. The web administration panel required protected server-side access, operational dashboards, CRUD foundations, analytics, and admin-specific security.")
    p(doc, "The AI dimension of the project is also significant. The body-measurement feature required a study of computer-vision models, including MoveNet, MediaPipe Pose Landmarker, MediaPipe Image Segmenter, segmentation masks, feature vectors, statistical estimation, confidence handling, and validation protocols. The language-model features required controlled prompts, backend/n8n integration, Groq usage, and careful limits so that generated recommendations and session briefings remain assistant outputs rather than automatic decisions. The automation layer includes five n8n workflows: AI Session Prep, Booking Reminders, Check-in Reminders, Coach Daily Digest, and Admin Ops Digest.")
    p(doc, "The project also reflects the evolution of modern software development practice. AI tools such as Codex were used during development and reporting to accelerate repository analysis, code navigation, documentation synthesis, and report generation. MCP-style connected tools made the assistant more useful by grounding work in local files, official documentation, document-generation utilities, Supabase-related context, automation artifacts, and command outputs. This use was supervised and verified through local files, structural checks, diffs, and project-specific evidence. The report therefore documents both AI as a GoFit product capability and AI as a controlled development-support tool.")
    p(doc, "The main remaining work before production release concerns final QA and institutional finalization. The application should be tested on real Android and iOS devices, the iOS MediaPipe path should be finalized if the AI measurement feature is shipped on both platforms, store builds should be generated, production environment variables should be reviewed, payment/video flows should be hardened if activated, and final screenshots should replace diagram drafts where required by the university.")
    add_heading(doc, "Bibliography", 1)
    add_basic_table(
        doc,
        "Table B.1: Bibliography",
        ["Reference", "Source", "Use"],
        [
            ["[1]", "Expo documentation", "Mobile application development and EAS build process"],
            ["[2]", "React Native documentation", "Cross-platform mobile UI and APIs"],
            ["[3]", "Supabase documentation", "Authentication, database, storage, realtime, RLS, and edge functions"],
            ["[4]", "Next.js documentation", "Administration panel, App Router, server components, route handlers"],
            ["[5]", "Scrum Guide", "Scrum roles, artifacts, events, and iterative methodology"],
            ["[6]", "GoFit repository documentation", "Project-specific architecture, features, database, and sprint planning"],
            ["[7]", "MediaPipe documentation", "Pose Landmarker, Image Segmenter, native vision tasks, landmark and segmentation concepts"],
            ["[8]", "TensorFlow Lite / MoveNet documentation", "On-device pose-estimation baseline and model format"],
            ["[9]", "Groq API documentation", "Language-model integration for recommendation and briefing flows"],
            ["[10]", "n8n documentation", "Workflow automation, schedule triggers, HTTP/API orchestration, code nodes"],
            ["[11]", "OpenAI Codex / AI coding assistant usage notes", "AI-assisted development, repository analysis, documentation support, and supervised code-generation workflow"],
            ["[12]", "LiveKit documentation", "Video-call token generation and realtime room connection concepts"],
            ["[13]", "Model Context Protocol and connector/tooling documentation", "Controlled access to current project context, external tools, official documentation, and local artifact generation"],
            ["[14]", "Git documentation", "Version control, change tracking, branch management, and diff review"],
        ],
        [0.8, 2.3, 3.4],
    )
    add_heading(doc, "Appendices", 1)
    add_heading(doc, "Appendix A: Repository Evidence Map", 2)
    add_basic_table(
        doc,
        "Table A.1: Repository evidence used for the report",
        ["Path / area", "Content", "How it supports the report"],
        [
            ["GoFitMobile", "Expo/React Native mobile application, screens, services, stores, navigation, local native MediaPipe module, assets, model files.", "Supports mobile implementation chapters, body-measurement AI section, workout/session details, coach/client flows."],
            ["admin-panel", "Next.js administration interface, protected admin pages, dashboard, management screens, UI components.", "Supports Sprint 3 administration section and deployment architecture."],
            ["database", "SQL migrations, schema design, RLS policies, functions, marketplace helpers, measurement tables.", "Supports backend, security, data model, and deployment sections."],
            ["supabase/functions", "Edge Functions such as push notification, video token generation, and body-measurement related paths.", "Supports backend integration, privileged operations, and external-service sections."],
            ["docs", "Architecture notes, feature inventory, sprint breakdown, troubleshooting, AI research, validation protocol, report drafts.", "Supports academic structure, methodology, feature scope, and AI transparency."],
            ["docs/automation/n8n", "n8n workflow JSON files and automation notes.", "Supports AI Session Prep, Booking Reminders, Check-in Reminders, Coach Daily Digest, Admin Ops Digest, automation security, and verification sections."],
        ],
        [1.55, 2.4, 2.55],
    )
    add_heading(doc, "Appendix B: Complete Feature Coverage From Source Audit", 2)
    p(doc, "This appendix exists because GoFit is larger than the first report draft suggested. The project contains a client mobile app, a coach mobile app, an administration panel, Supabase backend functions and migrations, AI modules, and n8n operational workflows. The following checklist should be used when adding final screenshots or proof before submission.")
    add_basic_table(
        doc,
        "Table A.2: Source-audited feature coverage checklist",
        ["Area", "Feature families to cover", "Evidence to attach"],
        [
            ["Client mobile app", "Auth, onboarding, profile, goals, units, workout library, exercise detail, workout builder, active session, rest timer, summary, calendar/timeline, statistics, records, progress photos, body measurements, nutrition, barcode lookup, health sync, habits, readiness, notifications, marketplace, coach detail, bookings, packs, programs, chat, video entry points, settings, haptics, and device feedback.", "Screenshots of main flows; screen list or navigation capture; service/store references; haptic interaction examples."],
            ["Coach mobile app", "Coach onboarding, CV/certification upload, pending validation, profile preview, dashboard, clients, progress, notes, check-ins, calendar, availability, chat, wallet, packs, programs, program builder, profile, settings, video session screen, and coach-side haptic feedback.", "Screenshots of coach flow; Supabase rows for coach profile, bookings, packs, and messages; haptic interaction examples."],
            ["Admin panel", "Login, protected dashboard, users, user details, coaches, certifications, exercises, workouts, transactions, notifications, settings, activity logs, BI APIs, exports, saved views, scheduled digests, imports, bulk operations, and health route.", "Admin screenshots; protected route test; route/API inventory."],
            ["Backend and database", "Auth, user profiles, workouts, sessions, native workouts, nutrition, saved meals, progress photos, health/readiness, notifications, marketplace, bookings, packs, wallet, chat, check-ins, AI templates/session notes, BI views, admin settings, admin notifications, audit logs, storage, realtime, RLS, and security optimization.", "Migration list; schema screenshots; RLS/security notes; Supabase function list."],
            ["AI and computer vision", "MediaPipe Pose Landmarker, MediaPipe Image Segmenter, MoveNet/TFLite fallback, feature extraction, statistical measurement estimator, confidence scoring, photo quality warnings, manual correction, AI workout recommendation, AI session notes, and Groq-based briefing generation.", "Model files/module references; debug output; validation photos; prompt/function screenshots."],
            ["n8n automation", "AI Session Prep, Booking Reminders, Check-in Reminders, Coach Daily Digest, Admin Ops Digest, duplicate guards, run summaries, notification creation, and secret separation.", "Workflow canvas screenshots; run summaries; Supabase notification/note/admin_notification rows."],
            ["Development process", "Codex-assisted repository analysis, MCP/tool usage, official documentation lookup, document generation, dependency/library inventory, accessibility audit, Git versioning, and report regeneration.", "Command outputs, package.json dependency evidence, report-generation script, accessibility audit, and final DOCX metadata."],
        ],
        [1.35, 3.75, 1.95],
    )
    add_complete_feature_inventory_appendix(doc)
    add_heading(doc, "Appendix D: n8n Workflow Verification Checklist", 2)
    add_basic_table(
        doc,
        "Table A.8: n8n workflow verification checklist",
        ["Workflow", "Manual verification steps", "Expected evidence"],
        [
            ["AI Session Prep v1", "Create/find a confirmed booking within the next 24 hours; run workflow manually; confirm ai_session_notes row; confirm coach notification; run again.", "n8n run summary with generated/skipped counts, ai_session_notes row, notification row with data.kind = ai_session_ready, second run without duplicate note."],
            ["Booking Reminders v1", "Create/find confirmed booking starting within 60 minutes; run workflow manually; run again.", "Client and coach notifications with data.kind = booking_starting_soon; second run skips duplicates for same booking/recipient/window."],
            ["Check-in Reminders v1", "Create/find enabled schedule due today in UTC without a response; run workflow manually; run again.", "Client notification with data.kind = check_in_due; second run without duplicate reminder."],
            ["Coach Daily Digest v1", "Create/find coach with at least one digest signal; run workflow manually; run again.", "Coach notification with data.kind = coach_daily_digest and counts for missed check-ins, low wellness, upcoming sessions, or inactive clients; no duplicate for same UTC date."],
            ["Admin Ops Digest v1", "Confirm at least one admin profile; run workflow manually; run again.", "admin_notifications row per admin with title Daily ops digest and deterministic date link; no duplicate for same date."],
        ],
        [1.35, 3.2, 2.5],
    )
    add_heading(doc, "Appendix E: AI Validation Checklist", 2)
    add_basic_table(
        doc,
        "Table A.9: AI validation checklist",
        ["Checklist item", "Expected evidence", "Status"],
        [
            ["Pose detection works on device", "MediaPipe debug values show pose count, landmarks, visibility, and inference time.", "Partially validated on Android development client."],
            ["Segmentation mask is usable", "Person-mask coverage, class grid, connected component, and scan-line debug values are reviewed.", "Needs broader validation."],
            ["Measurements are compared with tape", "Chest, waist, hip, and shoulder tape values are recorded and compared with AI drafts.", "Protocol defined; dataset must be expanded."],
            ["Confidence is honest", "Low-quality scans reduce confidence and produce warnings.", "Needs continued refinement."],
            ["LLM outputs are scoped", "Prompts use only selected data and instruct the model not to invent medical facts.", "Implemented in workflow design."],
            ["AI results remain reviewable", "Users and coaches can interpret or override AI assistance.", "Design principle defined."],
        ],
        [1.7, 3.0, 1.8],
    )
    add_heading(doc, "Appendix F: Remaining Work Before Final Submission", 2)
    add_basic_table(
        doc,
        "Table A.10: Remaining work",
        ["Area", "Task", "Reason"],
        [
            ["Report formatting", "Update table of contents, list of figures, list of tables, page numbers, university names, supervisors, and final screenshots.", "Required for institutional submission."],
            ["Visual evidence", "Replace generated diagram drafts with final UML exports and real application screenshots where required.", "Improves credibility and alignment with the example report style."],
            ["Mobile QA", "Run full Android and iOS tests, including auth, workouts, progress, marketplace, coach flows, notifications, and AI measurement.", "Confirms behavior on real devices."],
            ["AI validation", "Collect a larger set of tape measurements and photo captures under controlled conditions.", "Needed to quantify accuracy and improve confidence scoring."],
            ["Deployment", "Finalize EAS builds, admin hosting, Supabase environment variables, Edge Functions, and monitoring.", "Required before production release."],
            ["Security review", "Confirm RLS policies, admin-role checks, service-role boundaries, storage permissions, and prompt-data exposure.", "Protects personal and operational data."],
        ],
        [1.35, 3.0, 2.1],
    )


def main():
    doc = Document()
    setup_styles(doc)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    cover(doc)
    front_matter(doc)
    introduction(doc)
    chapter1(doc)
    chapter2(doc)
    chapters3_to5(doc)
    chapter6(doc)
    final_pages(doc)
    finalize(doc)
    doc.core_properties.title = "GoFit Graduation Internship Report Draft"
    doc.core_properties.subject = "Filled PFE report draft for GoFit"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
