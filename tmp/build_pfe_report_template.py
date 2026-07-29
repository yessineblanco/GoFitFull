from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "output/documents/PFE_Graduation_Internship_Report_Template.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 90, 90)
LIGHT_FILL = "F2F4F7"
BOX_FILL = "F7F9FC"
BORDER = "B8C4D6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tcw = tc_pr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tc_pr.append(tcw)
            tcw.set(qn("w:w"), str(int(width * 1440)))
            tcw.set(qn("w:type"), "dxa")


def style_table(table, header=True):
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            if header and r_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.5)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)
        sec.header_distance = Inches(0.49)
        sec.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_placeholder_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = GRAY
    return p


def add_instruction(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(label + ": ")
    r1.bold = True
    r1.font.color.rgb = DARK_BLUE
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.color.rgb = GRAY


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9.5)
    return p


def add_figure_placeholder(doc, caption, note, height_lines=4):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.3])
    cell = table.cell(0, 0)
    set_cell_shading(cell, BOX_FILL)
    set_cell_border(cell)
    set_cell_margins(cell, top=180, bottom=180, start=180, end=180)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[DIAGRAM / FIGURE PLACEHOLDER]")
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    for _ in range(height_lines):
        p.add_run("\n")
    p.add_run(note).italic = True
    add_caption(doc, caption)


def add_basic_table(doc, caption, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    if widths:
        set_table_width(table, widths)
    style_table(table)
    add_caption(doc, caption)
    return table


def add_static_list(doc, title, entries):
    add_heading(doc, title, 1)
    add_placeholder_para(doc, "Update page numbers after you fill the report. This static list is provided as a clean starting structure.")
    table = doc.add_table(rows=0, cols=3)
    set_table_width(table, [5.4, 0.55, 0.35])
    for label, page in entries:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = "." * 24
        row[2].text = page
    style_table(table, header=False)


def add_front_matter(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    r = p.add_run("[UNIVERSITY / INSTITUTE NAME]")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("[DEPARTMENT / PROGRAM]").bold = True

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(72)
    title.paragraph_format.space_after = Pt(16)
    run = title.add_run("Graduation Internship Report")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("GoFit - Fitness Mobile Application and Administration Platform")
    run.bold = True
    run.font.size = Pt(18)

    add_figure_placeholder(doc, "Figure 0.1: Cover page logo placeholders", "Insert university and company logos here.", height_lines=2)

    meta = doc.add_table(rows=5, cols=2)
    set_table_width(meta, [2.3, 4.0])
    rows = [
        ("Prepared by", "[Student full name]"),
        ("Academic supervisor", "[Academic supervisor name]"),
        ("Professional supervisor", "[Company supervisor name]"),
        ("Host organization", "GoFit / [Host company name]"),
        ("Academic year", "[YYYY - YYYY]"),
    ]
    for i, (a, b) in enumerate(rows):
        meta.rows[i].cells[0].text = a
        meta.rows[i].cells[1].text = b
    style_table(meta, header=False)
    doc.add_page_break()

    add_heading(doc, "Validation / Signature Page", 1)
    add_placeholder_para(doc, "Use this page for official validation, jury signatures, supervisor approvals, internship dates, and administrative stamps required by your institution.")
    add_basic_table(
        doc,
        "Table 0.1: Validation signatures",
        ["Role", "Name", "Signature", "Date"],
        [
            ["Academic supervisor", "[Name]", "", ""],
            ["Professional supervisor", "[Name]", "", ""],
            ["Jury president", "[Name]", "", ""],
            ["Reviewer / examiner", "[Name]", "", ""],
        ],
        [1.8, 1.8, 1.5, 1.2],
    )
    doc.add_page_break()

    add_heading(doc, "Acknowledgements", 1)
    add_placeholder_para(doc, "Write a brief, personal but professional acknowledgement of the people who supported the internship and report. Mention the host organization, supervisors, teaching staff, and anyone who contributed directly to the project.")
    doc.add_page_break()

    add_heading(doc, "Dedication", 1)
    add_placeholder_para(doc, "Write a short dedication to family, friends, mentors, or others you wish to honor. Keep the tone sincere and concise.")
    doc.add_page_break()

    toc = [
        ("General Introduction", "1"),
        ("1 Project Study", "2"),
        ("2 Sprint 0", "8"),
        ("3 Sprint 1", "15"),
        ("4 Sprint 2", "25"),
        ("5 Sprint 3", "35"),
        ("6 Deployment and Closing Phase", "45"),
        ("General Conclusion", "52"),
        ("Bibliography", "53"),
        ("Appendices", "54"),
    ]
    add_static_list(doc, "Table of Contents", toc)
    doc.add_page_break()
    figs = [
        ("Figure 1.1: Company logo", "3"),
        ("Figure 1.2: Scrum flow", "7"),
        ("Figure 2.1: Global use case diagram", "14"),
        ("Figure 2.2: Gantt chart", "15"),
        ("Figure 3.1: Sprint 1 class diagram", "17"),
        ("Figure 3.2: Refined use case diagram of Sprint 1", "18"),
        ("Figure 3.3: Sequence diagram", "20"),
        ("Figure 3.4: Burn down chart Sprint 1", "24"),
        ("Figure 6.1: Physical architecture diagram", "47"),
        ("Figure 6.2: Logical architecture diagram", "49"),
        ("Figure 6.3: Deployment diagram", "50"),
    ]
    add_static_list(doc, "List of Figures", figs)
    doc.add_page_break()
    tabs = [
        ("Table 1.1: Comparison between traditional and agile methodologies", "6"),
        ("Table 1.2: Comparison between Scrum and Kanban", "6"),
        ("Table 2.1: Functional requirements", "9"),
        ("Table 2.2: Non-functional requirements", "10"),
        ("Table 2.3: Scrum team", "11"),
        ("Table 2.4: Product backlog", "12"),
        ("Table 3.1: Sprint 1 backlog", "16"),
        ("Table 3.2: Use case textual description", "19"),
        ("Table 6.1: Hardware/material environment", "46"),
        ("Table 6.2: Software environment", "46"),
    ]
    add_static_list(doc, "List of Tables", tabs)
    doc.add_page_break()


def requirements_tables(doc):
    add_basic_table(
        doc,
        "Table 2.1: Functional requirements",
        ["ID", "Requirement", "Actor", "Priority", "Status"],
        [
            ["FR-01", "The user can create an account, complete a profile, and access personalized fitness features.", "Mobile user", "High", "[Planned/Done]"],
            ["FR-02", "The user can plan workouts, browse exercises, log sessions, and follow progress over time.", "Mobile user", "High", "[Planned/Done]"],
            ["FR-03", "The user can track nutrition, meals, and related health indicators.", "Mobile user", "Medium", "[Planned/Done]"],
            ["FR-04", "A coach can create a profile, manage clients, propose programs, and communicate with users.", "Coach", "Medium", "[Planned/Done]"],
            ["FR-05", "An administrator can manage users, coaches, exercises, workouts, and platform content.", "Admin", "High", "[Planned/Done]"],
        ],
        [0.7, 2.8, 1.2, 0.9, 0.9],
    )
    add_basic_table(
        doc,
        "Table 2.2: Non-functional requirements",
        ["ID", "Category", "Requirement", "Measurement / Constraint"],
        [
            ["NFR-01", "Security", "Supabase authentication, role-based access, and row-level security must protect user, coach, and admin data.", "[State RLS/auth rules]"],
            ["NFR-02", "Performance", "Mobile screens and admin pages should load quickly enough for daily use and repeated tracking.", "[Target metric]"],
            ["NFR-03", "Usability", "Workout, nutrition, progress, coach, and admin flows should remain clear for non-technical users.", "[Evaluation method]"],
            ["NFR-04", "Maintainability", "The codebase should separate screens, services, stores, database access, and reusable components.", "[Constraint]"],
        ],
        [0.7, 1.2, 3.0, 1.6],
    )


def sprint_backlog_table(doc, sprint_no):
    sprint_rows = {
        3: [
            ["S1-01", "Implement authentication, onboarding, and user profile setup.", "High", "[Points/hours]", "[Name]", "[To do]"],
            ["S1-02", "Build workout and exercise browsing/planning foundations in the mobile app.", "High", "[Points/hours]", "[Name]", "[To do]"],
            ["S1-03", "Connect mobile screens to Supabase services and validate core data flows.", "Medium", "[Points/hours]", "[Name]", "[To do]"],
        ],
        4: [
            ["S2-01", "Implement nutrition logging and saved meal flows.", "High", "[Points/hours]", "[Name]", "[To do]"],
            ["S2-02", "Build progress tracking views, charts, photos, or body measurement support.", "High", "[Points/hours]", "[Name]", "[To do]"],
            ["S2-03", "Add tests and validation for workout, nutrition, and progress services.", "Medium", "[Points/hours]", "[Name]", "[To do]"],
        ],
        5: [
            ["S3-01", "Implement coach onboarding, marketplace browsing, and coach profile flows.", "High", "[Points/hours]", "[Name]", "[To do]"],
            ["S3-02", "Develop coach-client programs, bookings, check-ins, or communication features.", "Medium", "[Points/hours]", "[Name]", "[To do]"],
            ["S3-03", "Build or finalize the Next.js admin panel for users, coaches, exercises, workouts, and analytics.", "High", "[Points/hours]", "[Name]", "[To do]"],
        ],
    }
    add_basic_table(
        doc,
        f"Table {sprint_no}.1: Sprint {sprint_no - 2 if sprint_no >= 3 else 1} backlog",
        ["ID", "User story / Task", "Priority", "Estimate", "Responsible", "Status"],
        sprint_rows.get(sprint_no, [
            ["S-01", "[As an actor, I want to...]", "[High]", "[Points/hours]", "[Name]", "[To do]"],
            ["S-02", "[Implementation or design task]", "[Medium]", "[Points/hours]", "[Name]", "[To do]"],
            ["S-03", "[Testing, documentation, or integration task]", "[Medium]", "[Points/hours]", "[Name]", "[To do]"],
        ]),
        [0.55, 2.55, 0.75, 0.85, 1.0, 0.8],
    )


def use_case_table(doc, caption):
    add_basic_table(
        doc,
        caption,
        ["Field", "Description"],
        [
            ["Use case", "[Name of the use case]"],
            ["Primary actor", "[Actor who starts the interaction]"],
            ["Preconditions", "[Conditions that must be true before execution]"],
            ["Main scenario", "[Number the main steps of the interaction]"],
            ["Alternative scenarios", "[Describe exceptions or optional paths]"],
            ["Postconditions", "[Expected system state after successful execution]"],
        ],
        [1.7, 4.8],
    )


def add_sprint_chapter(doc, chapter_no, title):
    add_heading(doc, f"{chapter_no} {title}", 1)
    add_placeholder_para(doc, f"Open this chapter with a short paragraph explaining the sprint goal, selected backlog scope, and business value delivered during {title}.")
    add_heading(doc, f"{chapter_no}.1 Needs Assessment", 2)
    add_instruction(doc, "What to write", "Explain the needs selected for this sprint, why they were prioritized, and how they contribute to the global project objective.")
    add_heading(doc, f"{chapter_no}.1.1 Sprint Backlog", 3)
    sprint_backlog_table(doc, chapter_no)
    add_heading(doc, f"{chapter_no}.1.2 Sprint Class Diagram", 3)
    add_figure_placeholder(doc, f"Figure {chapter_no}.1: {title} class diagram", "Insert the UML class diagram for the sprint scope.", height_lines=5)

    add_heading(doc, f"{chapter_no}.2 Functional Specification of Requirements", 2)
    add_instruction(doc, "What to write", "Describe the selected use cases at a refined level. Clarify actors, preconditions, nominal scenarios, alternatives, and expected results.")
    add_heading(doc, f"{chapter_no}.2.1 Refined Use Case Diagram of {title}", 3)
    add_figure_placeholder(doc, f"Figure {chapter_no}.2: Refined use case diagram of {title}", "Insert the sprint-level use case diagram.", height_lines=5)
    add_heading(doc, f"{chapter_no}.2.2 Use Case Analysis", 3)
    use_case_table(doc, f"Table {chapter_no}.2: Use case textual description")
    add_figure_placeholder(doc, f"Figure {chapter_no}.3: Sequence diagram for a key use case", "Insert a sequence diagram showing actor-system interactions for one important use case.", height_lines=5)

    add_heading(doc, f"{chapter_no}.3 Realisation", 2)
    add_instruction(doc, "What to write", "Present the implemented interfaces, services, database changes, APIs, or modules. Add screenshots with captions and explain the functional result.")
    add_figure_placeholder(doc, f"Figure {chapter_no}.4: Realisation screenshot placeholder", "Insert a representative screen, API result, or implemented feature view.", height_lines=4)

    test_label = "Sprint Tests" if chapter_no == 3 else "Tests"
    add_heading(doc, f"{chapter_no}.4 {test_label}", 2)
    add_instruction(doc, "What to write", "Summarize unit, integration, end-to-end, or acceptance tests. State what was tested, expected result, actual result, and whether defects remain.")
    add_basic_table(
        doc,
        f"Table {chapter_no}.3: Test cases",
        ["Test ID", "Scenario", "Expected result", "Actual result", "Status"],
        [
            ["T-01", "[Describe test scenario]", "[Expected behavior]", "[Observed behavior]", "[Pass/Fail]"],
            ["T-02", "[Describe test scenario]", "[Expected behavior]", "[Observed behavior]", "[Pass/Fail]"],
        ],
        [0.75, 1.8, 1.55, 1.55, 0.85],
    )

    add_heading(doc, f"{chapter_no}.5 Sprint Burn Down Chart", 2)
    add_figure_placeholder(doc, f"Figure {chapter_no}.5: Burn down chart for {title}", "Insert a burn down chart showing remaining work across sprint days.", height_lines=4)
    add_heading(doc, "Conclusion", 2)
    add_placeholder_para(doc, f"Conclude the chapter by summarizing what {title} delivered, what was validated, and what will be addressed in the following sprint.")
    doc.add_page_break()


def add_body(doc):
    add_heading(doc, "General Introduction", 1)
    add_instruction(doc, "General context", "Introduce the digital fitness domain and the need for unified tools that combine workout planning, nutrition tracking, progress monitoring, coach-client collaboration, and administration.")
    add_instruction(doc, "Project objective", "Present GoFit as the internship project: a fitness platform composed of an Expo/React Native mobile application, a Supabase backend, and a Next.js administration interface.")
    add_instruction(doc, "Internship context", "Briefly present the host organization or project environment, the internship duration, supervision context, and your role in the GoFit implementation.")
    add_instruction(doc, "Report organization", "End with a chapter-by-chapter roadmap: project study, Sprint 0, three sprint chapters, deployment and closing phase, then conclusion.")
    doc.add_page_break()

    add_heading(doc, "1 Project Study", 1)
    sections = [
        ("1.1 Project Context", "Describe the business, academic, and technical context that led to GoFit. Explain the fitness application domain, user expectations, coach-client needs, and the operational need for a connected mobile and web platform."),
        ("1.2 Host Organization Presentation", "Introduce the host organization or project owner and explain its relation to GoFit, including the environment in which the mobile app, admin panel, and backend were developed."),
        ("1.2.1 Presentation of the Company", "Present the company's activity area, services, target clients, structure, and relevant departments."),
        ("1.2.2 Company Logo", "Insert the company logo and give a short caption. Do not over-explain the logo unless it is relevant to the report."),
        ("1.3 Problem Statement", "Define the central problem GoFit addresses: fragmented workout tracking, nutrition follow-up, progress analysis, coach communication, booking, and platform administration."),
        ("1.4 Study of Existing Conditions", "Describe existing fitness applications, manual coaching workflows, separate nutrition trackers, spreadsheet-based progress tracking, and disconnected admin tools."),
        ("1.5 Limits and Inconveniences of Existing Solution", "Analyze shortcomings such as fragmented user journeys, weak coach-client connectivity, limited customization, manual progress analysis, separate payment or booking tools, and lack of centralized administration."),
        ("1.6 Proposed Solution", "Present GoFit at a high level: an Expo/React Native mobile application connected to a Supabase backend, complemented by a Next.js administration panel and coach marketplace features."),
        ("1.7 Adopted Development Methodology and Modeling Language", "Justify the project management and modeling approach used throughout the internship."),
        ("1.7.1 Comparison Between Traditional and Agile Methodologies", "Compare sequential and iterative approaches, then explain why the chosen approach fits the project constraints."),
    ]
    for title, text in sections:
        add_heading(doc, title, 2 if title.count(".") == 1 else 3)
        add_placeholder_para(doc, text)
        if title == "1.2.2 Company Logo":
            add_figure_placeholder(doc, "Figure 1.1: Company logo", "Insert the official company logo.", height_lines=2)
        if title == "1.7.1 Comparison Between Traditional and Agile Methodologies":
            add_basic_table(
                doc,
                "Table 1.1: Comparison between traditional and agile methodologies",
                ["Criterion", "Traditional approach", "Agile approach"],
                [
                    ["Planning", "[Describe upfront planning]", "[Describe iterative planning]"],
                    ["Change management", "[Describe change control]", "[Describe adaptability]"],
                    ["Delivery", "[Describe final delivery]", "[Describe incremental delivery]"],
                ],
                [1.4, 2.45, 2.45],
            )
    add_heading(doc, "1.7.2 Comparison Between Scrum and Kanban", 3)
    add_placeholder_para(doc, "Compare Scrum and Kanban using criteria such as roles, events, planning rhythm, visual management, and suitability for the internship project.")
    add_basic_table(
        doc,
        "Table 1.2: Comparison between Scrum and Kanban",
        ["Criterion", "Scrum", "Kanban"],
        [
            ["Cadence", "[Sprint-based iterations]", "[Continuous flow]"],
            ["Roles", "[Product Owner, Scrum Master, Development Team]", "[No mandatory roles]"],
            ["Best fit", "[Projects with planned increments]", "[Operational flow and support work]"],
        ],
        [1.4, 2.45, 2.45],
    )
    add_heading(doc, "1.7.3 Scrum", 3)
    add_placeholder_para(doc, "Define Scrum briefly and explain how its artifacts and events will structure the project work.")
    add_figure_placeholder(doc, "Figure 1.2: Scrum flow", "Insert a Scrum process flow showing product backlog, sprint planning, sprint, review, and retrospective.", height_lines=4)
    add_heading(doc, "1.7.4 Modeling Language / UML", 3)
    add_placeholder_para(doc, "Explain why UML diagrams are used to model requirements, structure, behavior, and deployment views.")
    add_heading(doc, "1.8 Conclusion", 2)
    add_placeholder_para(doc, "Summarize the chapter and transition to Sprint 0, where requirements and project organization are prepared.")
    doc.add_page_break()

    add_heading(doc, "2 Sprint 0", 1)
    add_placeholder_para(doc, "Introduce Sprint 0 as the preparation phase: requirement discovery, actor identification, backlog creation, planning, architecture orientation, and initial project management setup.")
    add_heading(doc, "2.1 Needs Assessment", 2)
    add_heading(doc, "2.1.1 Identification of Actors", 3)
    add_placeholder_para(doc, "List all actors who interact with the system. Distinguish primary users, administrators, external systems, and indirect stakeholders.")
    add_basic_table(
        doc,
        "Table 2.0: Actor identification",
        ["Actor", "Role in the system", "Main responsibilities"],
        [
            ["Mobile user", "Primary", "Uses GoFit to manage profile, workouts, nutrition, progress, and coach interactions."],
            ["Coach", "Primary", "Creates a coach profile, manages clients, programs, bookings, communication, and check-ins."],
            ["Administrator", "Primary", "Uses the web admin panel to manage users, coaches, exercises, workouts, analytics, and content."],
            ["Supabase services", "External system", "Provides authentication, PostgreSQL database, storage, and backend security rules."],
        ],
        [1.5, 1.6, 3.4],
    )
    add_heading(doc, "2.1.2 Functional Requirements", 3)
    requirements_tables(doc)
    add_heading(doc, "2.1.3 Non-Functional Requirements", 3)
    add_placeholder_para(doc, "Complement the requirements tables with a short explanation of the quality constraints that guide implementation decisions.")
    add_heading(doc, "2.2 Project Structure and Management with Scrum", 2)
    add_heading(doc, "2.2.1 Scrum Team", 3)
    add_basic_table(
        doc,
        "Table 2.3: Scrum team",
        ["Scrum role", "Assigned person", "Responsibilities"],
        [
            ["Product Owner", "[Name]", "[Defines priorities and validates business value]"],
            ["Scrum Master", "[Name]", "[Facilitates Scrum events and removes blockers]"],
            ["Development Team", "[Names]", "[Designs, implements, tests, and documents the solution]"],
            ["Stakeholders", "[Names/roles]", "[Provide feedback and validation]"],
        ],
        [1.4, 1.6, 3.5],
    )
    add_heading(doc, "2.2.2 Product Backlog", 3)
    add_basic_table(
        doc,
        "Table 2.4: Product backlog",
        ["ID", "Epic / User story", "Priority", "Estimate", "Sprint"],
        [
            ["PB-01", "As a mobile user, I want to authenticate and manage my profile so that my fitness experience is personalized.", "High", "[Points]", "Sprint 1"],
            ["PB-02", "As a mobile user, I want to create and follow workout plans so that I can organize my training.", "High", "[Points]", "Sprint 1"],
            ["PB-03", "As a mobile user, I want to track nutrition and progress so that I can measure my evolution.", "Medium", "[Points]", "Sprint 2"],
            ["PB-04", "As a coach, I want to manage clients, programs, bookings, and communication so that I can deliver personalized support.", "Medium", "[Points]", "Sprint 3"],
            ["PB-05", "As an administrator, I want to manage users, coaches, exercises, workouts, and analytics from a web panel.", "High", "[Points]", "Sprint 3"],
        ],
        [0.65, 3.15, 0.85, 0.85, 1.0],
    )
    add_heading(doc, "2.2.3 Sprint Planning", 3)
    add_placeholder_para(doc, "Describe sprint duration, selection criteria, planned increments, review dates, and how backlog items are distributed across sprints.")
    add_basic_table(
        doc,
        "Table 2.5: Sprint planning",
        ["Sprint", "Objective", "Main backlog items", "Start", "End"],
        [
            ["Sprint 1", "[Objective]", "[PB IDs]", "[Date]", "[Date]"],
            ["Sprint 2", "[Objective]", "[PB IDs]", "[Date]", "[Date]"],
            ["Sprint 3", "[Objective]", "[PB IDs]", "[Date]", "[Date]"],
        ],
        [0.9, 1.8, 1.8, 1.0, 1.0],
    )
    add_figure_placeholder(doc, "Figure 2.1: Gantt chart", "Insert a Gantt chart showing phases, sprints, and major milestones.", height_lines=4)
    add_heading(doc, "2.2.4 Global Use Case Diagram", 3)
    add_figure_placeholder(doc, "Figure 2.2: Global use case diagram", "Insert the global UML use case diagram for all actors and major system functions.", height_lines=6)
    add_heading(doc, "2.3 Project Management", 2)
    add_placeholder_para(doc, "Explain tools and practices used to manage GoFit tasks, source code, documentation, communication, database changes, mobile builds, admin-panel work, and sprint progress tracking.")
    add_heading(doc, "Conclusion", 2)
    add_placeholder_para(doc, "Summarize the preparation work and introduce Sprint 1 as the first implementation increment.")
    doc.add_page_break()

    add_sprint_chapter(doc, 3, "Sprint 1")
    add_sprint_chapter(doc, 4, "Sprint 2")
    add_sprint_chapter(doc, 5, "Sprint 3")

    add_heading(doc, "6 Deployment and Closing Phase", 1)
    add_placeholder_para(doc, "Present the final GoFit technical environment, architecture, deployment strategy, and validation results for the mobile app, backend, and administration panel.")
    add_heading(doc, "6.1 Working Environment", 2)
    add_heading(doc, "6.1.1 Material Environment", 3)
    add_basic_table(
        doc,
        "Table 6.1: Hardware/material environment",
        ["Resource", "Specification", "Purpose"],
        [
            ["Development machine", "[CPU, RAM, OS]", "[Development and testing]"],
            ["Mobile device / emulator", "[Model or emulator]", "[Application testing]"],
            ["Server / hosting resource", "[Plan or specification]", "[Deployment]"],
        ],
        [1.7, 2.5, 2.3],
    )
    add_heading(doc, "6.1.2 Software Environment", 3)
    add_basic_table(
        doc,
        "Table 6.2: Software environment",
        ["Software / Tool", "Version", "Use"],
        [
            ["Visual Studio Code", "[Version]", "Development environment"],
            ["Expo / React Native", "[SDK / RN version]", "Mobile application"],
            ["TypeScript", "[Version]", "Typed application code"],
            ["Next.js", "[Version]", "Administration panel"],
            ["Supabase", "[Project/version]", "Authentication, PostgreSQL database, storage, and security policies"],
            ["Testing tools", "[Jest/Playwright/manual protocol]", "Verification"],
            ["Deployment platform", "[EAS/Vercel/Supabase]", "Mobile, web, and backend deployment"],
        ],
        [2.2, 1.2, 3.1],
    )
    add_heading(doc, "6.2 Architecture and Technological Choices", 2)
    add_heading(doc, "6.2.1 Physical Architecture Diagram", 3)
    add_figure_placeholder(doc, "Figure 6.1: Physical architecture diagram", "Insert the infrastructure-level architecture showing devices, servers, services, and networks.", height_lines=5)
    add_heading(doc, "6.2.2 Technologies", 3)
    add_placeholder_para(doc, "Justify each GoFit technology: Expo/React Native for the mobile app, TypeScript for reliability, Next.js for the admin panel, Supabase for authentication/database/storage, and the chosen testing and deployment tools.")
    add_heading(doc, "6.2.3 Architecture Pattern", 3)
    add_placeholder_para(doc, "Explain the chosen architecture pattern, such as MVC, layered architecture, client-server, clean architecture, microservices, or another pattern used by the project.")
    add_heading(doc, "6.2.4 Logical Architecture Diagram", 3)
    add_figure_placeholder(doc, "Figure 6.2: Logical architecture diagram", "Insert the logical architecture with modules, layers, data flows, and service responsibilities.", height_lines=5)
    add_heading(doc, "6.3 Deployment", 2)
    add_heading(doc, "6.3.1 Deployment Diagram", 3)
    add_figure_placeholder(doc, "Figure 6.3: Deployment diagram", "Insert the UML deployment diagram with nodes, artifacts, and communication links.", height_lines=5)
    add_heading(doc, "6.3.2 CI/CD or Hosting", 3)
    add_placeholder_para(doc, "Describe the build, test, and release process. Mention hosting provider, environment variables, database migration strategy, and rollback approach if applicable.")
    add_heading(doc, "6.3.3 Test Results", 3)
    add_basic_table(
        doc,
        "Table 6.3: Final test results",
        ["Test type", "Scope", "Result", "Evidence / Notes"],
        [
            ["Unit tests", "[Modules covered]", "[Pass/Fail]", "[Link or screenshot reference]"],
            ["Integration tests", "[Flows covered]", "[Pass/Fail]", "[Link or screenshot reference]"],
            ["User acceptance", "[Scenarios covered]", "[Pass/Fail]", "[Reviewer feedback]"],
        ],
        [1.3, 1.8, 1.1, 2.3],
    )
    doc.add_page_break()

    add_heading(doc, "General Conclusion", 1)
    add_placeholder_para(doc, "Summarize the internship project, the final solution, achieved objectives, main technical and organizational lessons, limitations, and possible future improvements.")
    add_heading(doc, "Bibliography", 1)
    add_placeholder_para(doc, "List books, articles, official documentation, standards, tutorials, and tools consulted. Use the citation style required by your institution.")
    add_basic_table(
        doc,
        "Table B.1: Bibliography placeholder",
        ["Reference ID", "Source", "Type", "Access date / Notes"],
        [["[1]", "[Author or organization, title, URL if applicable]", "[Documentation/book/article]", "[Date]"]],
        [0.8, 3.5, 1.2, 1.0],
    )
    add_heading(doc, "Appendices", 1)
    add_placeholder_para(doc, "Add supporting material only if needed: extended diagrams, API examples, test reports, installation steps, additional screenshots, or administrative documents.")


def finalize(doc):
    for section in doc.sections:
        footer = section.footer
        if footer.paragraphs:
            p = footer.paragraphs[0]
        else:
            p = footer.add_paragraph()
        add_page_number(p)
        for p in section.header.paragraphs:
            p.text = ""

    props = doc.core_properties
    props.title = "PFE Graduation Internship Report Template"
    props.subject = "Academic report template"
    props.author = "Codex"
    props.comments = "Original fill-in template generated from requested structure; no source report content copied."


def main():
    doc = Document()
    setup_styles(doc)
    add_front_matter(doc)
    add_body(doc)
    finalize(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
