import json
import os
import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_gofit_filled_pfe_report as report_builder


ROOT = Path.cwd()
AI_WORKOUT_REPORT = ROOT / "output" / "documents" / "PFE_GoFit_Polished_Academic_Draft_Accessible_AI_Workout.docx"
BI_REPORT = ROOT / "output" / "documents" / "PFE_GoFit_Polished_Academic_Draft_Accessible_BI.docx"
BASE_REPORT = ROOT / "output" / "documents" / "PFE_GoFit_Polished_Academic_Draft_Accessible.docx"
REPORT = AI_WORKOUT_REPORT if AI_WORKOUT_REPORT.exists() else BI_REPORT if BI_REPORT.exists() else BASE_REPORT
OUT_JSON = ROOT / "output" / "documents" / "gofit_report_feature_coverage_audit.json"
OUT_MD = ROOT / "output" / "documents" / "gofit_report_feature_coverage_audit.md"


def norm(value):
    text = str(value or "").replace("\\", "/").lower()
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def norm_key(value):
    return re.sub(r"[^a-z0-9]+", "", norm(value))


def read_docx(docx_path):
    doc = Document(str(docx_path))
    paragraphs = [p.text for p in doc.paragraphs]
    table_rows = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        table_rows.append(rows)
    text = "\n".join(paragraphs + ["\t".join(row) for rows in table_rows for row in rows])
    return doc, text, table_rows


def load_package_dependencies(path):
    data = json.loads(path.read_text(encoding="utf-8"))
    deps = data.get("dependencies", {})
    dev_deps = data.get("devDependencies", {})
    return sorted(deps), sorted(dev_deps)


def package_aliases(package_name):
    lower = package_name.lower()
    aliases = {lower, lower.replace("@", "").replace("/", " "), lower.split("/")[-1].replace("-", " ")}
    special = {
        "@expo-google-fonts/barlow": ["barlow", "expo fonts", "fonts"],
        "@expo/vector-icons": ["vector icons", "icons"],
        "@hookform/resolvers": ["hookform resolvers", "react hook form"],
        "@livekit/react-native": ["livekit"],
        "@livekit/react-native-expo-plugin": ["livekit", "expo plugin"],
        "@livekit/react-native-webrtc": ["livekit", "webrtc"],
        "@react-native-async-storage/async-storage": ["async storage"],
        "@react-native-community/datetimepicker": ["datetime picker", "date picker", "calendar"],
        "@react-native-masked-view/masked-view": ["masked view"],
        "@react-navigation/bottom-tabs": ["react navigation", "bottom tabs"],
        "@react-navigation/native": ["react navigation"],
        "@react-navigation/stack": ["react navigation", "stack"],
        "@supabase/supabase-js": ["supabase javascript client", "supabase js", "supabase"],
        "expo": ["expo sdk", "expo"],
        "expo-av": ["audio", "video", "media"],
        "expo-blur": ["blur"],
        "expo-build-properties": ["eas build", "build properties"],
        "expo-calendar": ["calendar"],
        "expo-camera": ["camera"],
        "expo-dev-client": ["expo dev client"],
        "expo-document-picker": ["document picker"],
        "expo-file-system": ["file system"],
        "expo-font": ["fonts"],
        "expo-haptics": ["expo haptics", "haptics"],
        "expo-health-connect": ["health connect"],
        "expo-image": ["expo image", "image"],
        "expo-image-manipulator": ["image manipulator"],
        "expo-image-picker": ["image picker"],
        "expo-linear-gradient": ["linear gradient"],
        "expo-localization": ["localization"],
        "expo-location": ["location"],
        "expo-notifications": ["expo notifications", "notifications"],
        "expo-secure-store": ["secure store"],
        "expo-splash-screen": ["splash screen"],
        "expo-status-bar": ["status bar"],
        "expo-video": ["video"],
        "i18next": ["i18next", "localization"],
        "jpeg-js": ["jpeg", "image processing"],
        "lucide-react-native": ["lucide-react-native", "lucide"],
        "react": ["react"],
        "react-dom": ["react"],
        "react-hook-form": ["react hook form"],
        "react-i18next": ["react-i18next", "localization"],
        "react-native": ["react native"],
        "react-native-chart-kit": ["chart kit", "charts"],
        "react-native-fast-tflite": ["fast-tflite", "tflite"],
        "react-native-gesture-handler": ["gesture handler"],
        "react-native-health-connect": ["health connect"],
        "react-native-incall-manager": ["incall manager", "video"],
        "react-native-nitro-modules": ["nitro modules", "native modules"],
        "react-native-reanimated": ["reanimated"],
        "react-native-safe-area-context": ["safe area"],
        "react-native-screens": ["screens"],
        "react-native-svg": ["svg"],
        "react-native-vector-icons": ["vector icons"],
        "react-native-web": ["react-native-web", "web"],
        "react-native-worklets": ["worklets"],
        "react-native-worklets-core": ["worklets"],
        "webrtc-adapter": ["webrtc"],
        "zod": ["zod"],
        "zustand": ["zustand"],
        "@aws-sdk/client-s3": ["aws s3", "s3 client"],
        "@radix-ui/react-alert-dialog": ["radix", "alert dialog"],
        "@radix-ui/react-avatar": ["radix", "avatar"],
        "@radix-ui/react-checkbox": ["radix", "checkbox"],
        "@radix-ui/react-collapsible": ["radix", "collapsible"],
        "@radix-ui/react-dialog": ["radix", "dialog"],
        "@radix-ui/react-dropdown-menu": ["radix", "dropdown"],
        "@radix-ui/react-popover": ["radix", "popover"],
        "@radix-ui/react-progress": ["radix", "progress"],
        "@radix-ui/react-scroll-area": ["radix", "scroll area"],
        "@radix-ui/react-select": ["radix", "select"],
        "@radix-ui/react-slot": ["radix", "slot"],
        "@radix-ui/react-switch": ["radix", "switch"],
        "@radix-ui/react-toast": ["radix", "toast"],
        "@supabase/ssr": ["supabase ssr", "supabase"],
        "@tanstack/react-table": ["tanstack", "react table"],
        "class-variance-authority": ["class utilities", "component variants"],
        "clsx": ["class utilities"],
        "cmdk": ["cmdk"],
        "lucide-react": ["lucide react", "lucide"],
        "next": ["next.js", "next"],
        "next-themes": ["next-themes", "theme"],
        "react-day-picker": ["day picker", "date picker"],
        "recharts": ["recharts"],
        "sonner": ["sonner"],
        "tailwind-merge": ["class utilities", "tailwind"],
        "tailwindcss-animate": ["tailwind", "animations"],
    }
    aliases.update(special.get(lower, []))
    return sorted(aliases, key=len, reverse=True)


def audit_dependencies(report_text):
    report_lower = norm(report_text)
    results = []
    package_sets = [
        ("mobile", ROOT / "GoFitMobile" / "package.json"),
        ("admin", ROOT / "admin-panel" / "package.json"),
    ]
    for app, path in package_sets:
        deps, dev_deps = load_package_dependencies(path)
        for name in deps:
            aliases = package_aliases(name)
            exact = norm(name) in report_lower
            family = any(norm(alias) in report_lower for alias in aliases)
            results.append(
                {
                    "app": app,
                    "package": name,
                    "type": "dependency",
                    "exact_mentioned": exact,
                    "covered_by_family_or_alias": family,
                    "aliases_checked": aliases[:6],
                }
            )
        for name in dev_deps:
            aliases = package_aliases(name)
            exact = norm(name) in report_lower
            family = any(norm(alias) in report_lower for alias in aliases)
            results.append(
                {
                    "app": app,
                    "package": name,
                    "type": "devDependency",
                    "exact_mentioned": exact,
                    "covered_by_family_or_alias": family,
                    "aliases_checked": aliases[:6],
                }
            )
    return results


def find_feature_inventory_tables(tables):
    inventory = []
    expected_header = ["No.", "Feature", "Implementation evidence", "Libraries / backend", "Purpose"]
    for table in tables:
        if table and table[0] == expected_header:
            inventory.append(table)
    return inventory


def audit_feature_inventory_tables(tables):
    features = report_builder.parse_feature_inventory()
    groups = ["Mobile client", "Mobile coach", "Admin panel", "Backend and infrastructure"]
    actual_tables = find_feature_inventory_tables(tables)
    expected_by_group = {
        group: report_builder.feature_inventory_table_rows(features, group)
        for group in groups
    }
    comparison = []
    for idx, group in enumerate(groups):
        expected_rows = expected_by_group[group]
        actual_rows = actual_tables[idx][1:] if idx < len(actual_tables) else []
        expected_features = [row[1] for row in expected_rows]
        actual_features = [row[1] for row in actual_rows]
        comparison.append(
            {
                "group": group,
                "expected_rows": len(expected_rows),
                "actual_rows": len(actual_rows),
                "missing_feature_titles": [name for name in expected_features if name not in actual_features],
                "unexpected_feature_titles": [name for name in actual_features if name not in expected_features],
            }
        )
    return features, comparison


def report_term_checks(report_text):
    lower = norm(report_text)
    terms = [
        "AI Session Prep v1",
        "Booking Reminders v1",
        "Check-in Reminders v1",
        "Coach Daily Digest v1",
        "Admin Ops Digest v1",
        "n8n",
        "MediaPipe Pose Landmarker",
        "MediaPipe Image Segmenter",
        "MoveNet",
        "TFLite",
        "Groq",
        "LiveKit",
        "Expo Haptics",
        "haptics_enabled",
        "Codex",
        "MCP",
        "Supabase",
        "Row-level security",
        "React Hook Form",
        "Zod",
        "TanStack",
        "Radix",
        "Reanimated",
        "Gesture Handler",
        "ai-workout-recommendation",
        "RecommendedWorkouts",
        "workoutRecommendationService",
        "Adaptive workout",
        "llama-3.3-70b-versatile",
        "daily_readiness",
        "computeAdaptiveContext",
        "volumeAdjustment",
        "coach companion",
        "response_format",
        "Business Intelligence",
        "Advanced BI",
        "bi_finance_daily",
        "bi_user_lifecycle_daily",
        "bi_coach_ops_daily",
        "bi_client_health_daily",
        "/api/bi/export",
        "/api/bi/saved-views",
        "/api/bi/snapshot",
        "/api/bi/scheduled-digests",
        "saved BI views",
        "threshold alerts",
        "scheduled digests",
        "CSV exports",
    ]
    return {term: lower.count(norm(term)) for term in terms}


def project_artifact_counts(report_text):
    lower = norm(report_text)
    patterns = {
        "mobile_screens": ["GoFitMobile/src/screens"],
        "mobile_navigation": ["GoFitMobile/src/navigation"],
        "admin_pages": ["admin-panel/app"],
        "supabase_functions": ["supabase/functions"],
        "database_migrations": ["database/migrations"],
        "n8n_workflows": ["docs/automation/n8n/workflows"],
    }
    rows = {}
    for label, roots in patterns.items():
        files = []
        for root in roots:
            root_path = ROOT / root
            if not root_path.exists():
                continue
            for dirpath, _, filenames in os.walk(root_path):
                for filename in filenames:
                    if filename.endswith((".ts", ".tsx", ".sql", ".json")):
                        path = Path(dirpath, filename)
                        rel = path.relative_to(ROOT).as_posix()
                        files.append(rel)
        covered = [rel for rel in files if norm(rel) in lower]
        rows[label] = {
            "total_files": len(files),
            "explicit_path_mentions_in_report": len(covered),
            "sample_unmentioned_paths": [rel for rel in files if rel not in covered][:15],
        }
    return rows


def haptics_counts(report_text):
    report_lower = norm(report_text)
    files = set()
    calls = 0
    skip_dirs = {
        ".expo",
        ".git",
        ".gradle",
        ".next",
        "android",
        "build",
        "dist",
        "ios",
        "node_modules",
    }
    for dirpath, dirnames, filenames in os.walk(ROOT / "GoFitMobile"):
        dirnames[:] = [name for name in dirnames if name not in skip_dirs]
        for filename in filenames:
            if not filename.endswith((".ts", ".tsx")):
                continue
            path = Path(dirpath, filename)
            try:
                text = path.read_text(encoding="utf-8", errors="replace")
            except OSError:
                continue
            if "expo-haptics" in text or "Haptics." in text or "haptics_enabled" in text:
                files.add(path.relative_to(ROOT).as_posix())
                calls += text.count("Haptics.")
    return {
        "source_files_with_haptics": len(files),
        "source_haptics_calls": calls,
        "report_mentions_haptics": report_lower.count("haptic"),
        "sample_haptics_files": sorted(files)[:20],
    }


def n8n_workflow_audit(report_text):
    lower = norm(report_text)
    key_text = norm_key(report_text)
    workflow_dir = ROOT / "docs" / "automation" / "n8n" / "workflows"
    rows = []
    for path in sorted(workflow_dir.glob("*.json")):
        data = json.loads(path.read_text(encoding="utf-8"))
        name = data.get("name") or path.stem
        name_without_prefix = re.sub(r"^gofit\s+", "", norm(name), flags=re.IGNORECASE)
        stem_label = path.stem.replace("-", " ")
        rows.append(
            {
                "file": path.relative_to(ROOT).as_posix(),
                "workflow_name": name,
                "mentioned_by_name": norm(name) in lower or norm_key(name) in key_text or norm_key(name_without_prefix) in key_text,
                "mentioned_by_file_stem": norm(stem_label) in lower or norm_key(stem_label) in key_text,
            }
        )
    return rows


def verdict(feature_comparison, dependency_results, term_counts, n8n_rows):
    missing_features = sum(len(row["missing_feature_titles"]) for row in feature_comparison)
    unexpected_features = sum(len(row["unexpected_feature_titles"]) for row in feature_comparison)
    missing_n8n = [row for row in n8n_rows if not row["mentioned_by_name"] and not row["mentioned_by_file_stem"]]
    missing_core_terms = [term for term, count in term_counts.items() if count == 0]
    dependency_family_missing = [
        row for row in dependency_results
        if row["type"] == "dependency" and not row["covered_by_family_or_alias"]
    ]
    return {
        "feature_inventory_status": "PASS" if missing_features == 0 and unexpected_features == 0 else "FAIL",
        "missing_feature_count": missing_features,
        "unexpected_feature_count": unexpected_features,
        "n8n_status": "PASS" if not missing_n8n else "FAIL",
        "missing_n8n_workflows": missing_n8n,
        "core_term_status": "PASS" if not missing_core_terms else "WARN",
        "missing_core_terms": missing_core_terms,
        "dependency_family_status": "PASS" if not dependency_family_missing else "WARN",
        "dependency_family_missing_count": len(dependency_family_missing),
        "dependency_family_missing": dependency_family_missing,
    }


def write_markdown(data):
    v = data["verdict"]
    lines = [
        "# GoFit Report Feature Coverage Audit",
        "",
        f"Report audited: `{REPORT.as_posix()}`",
        "",
        "## Verdict",
        "",
        f"- Feature inventory: **{v['feature_inventory_status']}** ({v['missing_feature_count']} missing, {v['unexpected_feature_count']} unexpected).",
        f"- n8n workflows: **{v['n8n_status']}**.",
        f"- Core feature/tool terms: **{v['core_term_status']}**.",
        f"- Runtime dependency family coverage: **{v['dependency_family_status']}**.",
        "",
        "## Feature Inventory Tables",
        "",
    ]
    for row in data["feature_table_comparison"]:
        lines.append(f"- {row['group']}: expected {row['expected_rows']}, report has {row['actual_rows']}.")
        if row["missing_feature_titles"]:
            lines.append(f"  Missing: {', '.join(row['missing_feature_titles'][:20])}")
        if row["unexpected_feature_titles"]:
            lines.append(f"  Unexpected: {', '.join(row['unexpected_feature_titles'][:20])}")

    lines += [
        "",
        "## Key Term Coverage",
        "",
    ]
    for term, count in data["term_counts"].items():
        lines.append(f"- {term}: {count}")

    lines += [
        "",
        "## Dependency Coverage",
        "",
        "Dependency coverage is evaluated at family/alias level because the report describes libraries academically rather than as a raw package-lock dump.",
    ]
    dep_counter = Counter(
        (row["app"], row["type"], row["covered_by_family_or_alias"])
        for row in data["dependencies"]
    )
    for key, count in sorted(dep_counter.items()):
        lines.append(f"- {key[0]} {key[1]} family-covered={key[2]}: {count}")
    missing = [row for row in data["dependencies"] if row["type"] == "dependency" and not row["covered_by_family_or_alias"]]
    if missing:
        lines.append("")
        lines.append("Runtime dependencies not covered even by family/alias:")
        for row in missing:
            lines.append(f"- {row['app']}: {row['package']}")

    lines += [
        "",
        "## n8n Workflows",
        "",
    ]
    for row in data["n8n_workflows"]:
        status = "covered" if row["mentioned_by_name"] or row["mentioned_by_file_stem"] else "missing"
        lines.append(f"- {row['workflow_name']} (`{row['file']}`): {status}")

    lines += [
        "",
        "## Haptics Probe",
        "",
        f"- Source files using haptics: {data['haptics']['source_files_with_haptics']}",
        f"- Haptics calls found: {data['haptics']['source_haptics_calls']}",
        f"- Report mentions of haptic/haptics: {data['haptics']['report_mentions_haptics']}",
        "",
        "## Project Artifact Path Mentions",
        "",
        "This is a secondary sanity check. A feature can be covered even when every source path is not printed in the report.",
    ]
    for label, row in data["artifact_counts"].items():
        lines.append(
            f"- {label}: {row['explicit_path_mentions_in_report']} explicit path mentions out of {row['total_files']} files."
        )

    OUT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main():
    if not REPORT.exists():
        raise SystemExit(f"Missing report: {REPORT}")
    _, report_text, tables = read_docx(REPORT)
    features, feature_table_comparison = audit_feature_inventory_tables(tables)
    dependencies = audit_dependencies(report_text)
    term_counts = report_term_checks(report_text)
    n8n_rows = n8n_workflow_audit(report_text)
    data = {
        "report": str(REPORT),
        "feature_source": "FEATURES.md",
        "parsed_feature_rows": len(features),
        "feature_group_counts": Counter(report_builder.feature_group(row["area"]) for row in features),
        "feature_table_comparison": feature_table_comparison,
        "term_counts": term_counts,
        "dependencies": dependencies,
        "n8n_workflows": n8n_rows,
        "haptics": haptics_counts(report_text),
        "artifact_counts": project_artifact_counts(report_text),
    }
    data["feature_group_counts"] = dict(data["feature_group_counts"])
    data["verdict"] = verdict(feature_table_comparison, dependencies, term_counts, n8n_rows)
    OUT_JSON.write_text(json.dumps(data, indent=2), encoding="utf-8")
    write_markdown(data)
    print(json.dumps(data["verdict"], indent=2))
    print(str(OUT_MD))
    print(str(OUT_JSON))


if __name__ == "__main__":
    main()
